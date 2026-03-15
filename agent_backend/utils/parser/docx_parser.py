from typing import List, Dict
from docx import Document


def parse_docx(file_path: str) -> List[Dict]:
    print("[方法调试] 进入 parse_docx | 局部变量:", locals())
    print(f"[ParserDebug] parse_docx called: file_path={file_path}")
    doc = Document(file_path)
    print(f"[ParserDebug] parse_docx opened: paragraph_count={len(doc.paragraphs)}")
    chunks: List[Dict] = []
    idx = 1
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        chunks.append({"chunk_id": f"d_{idx}", "page": None, "text": text})
        idx += 1
    print(f"[ParserDebug] parse_docx success: chunk_count={len(chunks)}")
    return chunks
