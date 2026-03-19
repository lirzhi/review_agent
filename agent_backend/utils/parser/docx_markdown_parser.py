#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
通用 DOCX / DOC 解析器：Word -> JSON（按章节粗略分块，content 为 Markdown）

增强点：
1. 页眉页脚过滤
   - 过滤页码、日期、重复页眉/页脚、机构名等
2. 更稳健的目录识别与跳过
   - 识别“目录/contents”标题
   - 识别点线目录、目录页码、目录块
3. 跨页表格 / 连续表格处理
   - 识别前后相邻且表头一致的表格并自动拼接
4. 合并单元格处理
   - 通过 OOXML gridSpan / vMerge 尽可能展开横向/纵向合并单元格
5. 附件表单类 doc/docx 更强结构抽取
   - 针对“附1-1 / 附件 1 / 汇总表 / 批分析数据 / 信息表”等表单型附件增强识别
   - 可将表格前的标题、注释、说明挂接到表格元数据中
6. content 保序输出
   - 文本 / 表格 / 图片按 Word 中原始顺序输出为 Markdown

依赖：
    pip install python-docx lxml pillow

可选：
    系统安装 libreoffice/soffice 后，可自动将 .doc 转为 .docx

示例：
    python generic_docx_to_markdown_json.py input.docx
    python generic_docx_to_markdown_json.py input.doc --embed-images true
"""

from __future__ import annotations

import argparse
import base64
import copy
import json
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Tuple

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


CN_NUM = "一二三四五六七八九十百千零"


# ============================================================
# 基础工具
# ============================================================

def str2bool(v: str) -> bool:
    return str(v).strip().lower() in {"1", "true", "yes", "y", "on"}


def normalize_text(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    s = s.replace("\u3000", " ").replace("\xa0", " ")
    s = s.replace("［", "[").replace("］", "]")
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace("：", ":")
    s = s.replace("．", ".")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\s+\n", "\n", s)
    s = re.sub(r"\n\s+", "\n", s)
    return s.strip()


def strip_trailing_colon(text: str) -> str:
    return re.sub(r"[:：]\s*$", "", normalize_text(text)).strip()


def escape_md(text: str) -> str:
    return str(text).replace("\n", "<br>").strip().replace("|", "\\|")


def safe_filename(name: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name)
    name = re.sub(r"\s+", "_", name).strip("_")
    return name or "image"


def markdown_table(columns: List[str], data: List[List[str]]) -> str:
    if not columns:
        return ""
    rows = []
    rows.append("| " + " | ".join(escape_md(c) for c in columns) + " |")
    rows.append("| " + " | ".join("---" for _ in columns) + " |")
    for row in data:
        row = row + [""] * (len(columns) - len(row))
        rows.append("| " + " | ".join(escape_md(c) for c in row[: len(columns)]) + " |")
    return "\n".join(rows)


def image_bytes_to_base64(data: bytes) -> str:
    return base64.b64encode(data).decode("utf-8")


def mostly_empty_row(row: List[str]) -> bool:
    return sum(1 for x in row if normalize_text(x)) <= 1


def is_all_empty_rows(rows: List[List[str]]) -> bool:
    return all(sum(1 for x in r if normalize_text(x)) == 0 for r in rows)


# ============================================================
# 标题识别
# ============================================================

@dataclass
class HeadingInfo:
    level: int
    title: str
    number: str = ""
    kind: str = "generic"


def normalize_style_name(style_name: Optional[str]) -> str:
    if not style_name:
        return ""
    return normalize_text(style_name).lower()


def heading_level_from_style(style_name: Optional[str]) -> Optional[int]:
    s = normalize_style_name(style_name)
    if not s:
        return None
    m = re.search(r"(?:heading|标题)\s*([1-9])", s)
    if m:
        return int(m.group(1))
    if s in {"title", "标题"}:
        return 1
    return None


def extract_heading_info(text: str) -> Optional[HeadingInfo]:
    t = strip_trailing_colon(text)
    if not t:
        return None

    m = re.match(r"^(3\.2\.[SP]\.(?:\d+)(?:\.\d+)*)\s+(.+)$", t, re.I)
    if m:
        depth = m.group(1).count(".") - 2
        return HeadingInfo(level=min(6, 2 + depth), title=t, number=m.group(1), kind="ctd")
    m = re.match(r"^(3\.2\.[SP]\.(?:\d+)(?:\.\d+)*)$", t, re.I)
    if m:
        depth = m.group(1).count(".") - 2
        return HeadingInfo(level=min(6, 2 + depth), title=t, number=m.group(1), kind="ctd")

    m = re.match(rf"^([{CN_NUM}]+)[、.]\s*(.+)$", t)
    if m:
        return HeadingInfo(level=1, title=t, number=m.group(1), kind="cn1")

    m = re.match(rf"^\(([{CN_NUM}]+)\)\s*(.+)$", t)
    if m:
        return HeadingInfo(level=2, title=t, number=m.group(1), kind="cn2")

    m = re.match(r"^(\d+)[、.]\s*(.+)$", t)
    if m:
        return HeadingInfo(level=3, title=t, number=m.group(1), kind="arabic1")

    m = re.match(r"^(\d+)\s+(.+)$", t)
    if m and len(m.group(2)) <= 60:
        return HeadingInfo(level=3, title=t, number=m.group(1), kind="arabic1_space")

    m = re.match(r"^(\d+(?:\.\d+){1,5})\s+(.+)$", t)
    if m:
        lvl = min(6, 3 + m.group(1).count("."))
        return HeadingInfo(level=lvl, title=t, number=m.group(1), kind="arabic_multi")
    m = re.match(r"^(\d+(?:\.\d+){1,5})$", t)
    if m:
        lvl = min(6, 3 + m.group(1).count("."))
        return HeadingInfo(level=lvl, title=t, number=m.group(1), kind="arabic_multi")

    if re.match(r"^(附件|附录|附)\s*[-— ]*[A-Za-z0-9一二三四五六七八九十]+[:：]?\s*.*$", t):
        return HeadingInfo(level=1, title=t, kind="appendix")
    if re.match(r"^(附\d+(?:-\d+)?)\s*$", t):
        return HeadingInfo(level=1, title=t, kind="appendix")
    if re.match(r"^[\[【]?(参考文献|名词解释|名词术语|术语|申报要求|申报资料要求|沟通交流|注册检验|著者)[\]】]?$", t):
        return HeadingInfo(level=1, title=t, kind="tail")
    return None


def is_probable_heading(text: str, style_name: Optional[str]) -> Optional[HeadingInfo]:
    t = strip_trailing_colon(text)
    if not t or len(t) > 150:
        return None

    lvl = heading_level_from_style(style_name)
    if lvl is not None:
        return HeadingInfo(level=lvl, title=t, kind="style")

    info = extract_heading_info(t)
    if info is None:
        return None

    if t.endswith(("。", "；", ";", "，", ",")):
        return None
    return info


def is_appendix_like_title(text: str) -> bool:
    t = strip_trailing_colon(text)
    return bool(
        re.match(r"^(附件|附录|附)\s*[-— ]*[A-Za-z0-9一二三四五六七八九十]+", t)
        or re.match(r"^(附\d+(?:-\d+)?)\s*$", t)
        or re.search(r"(汇总表|信息表|批分析数据|杂质谱分析|工艺流程图|生产工艺|主要仪器设备)", t)
    )


# ============================================================
# 页眉页脚 / 目录识别
# ============================================================

HEADER_FOOTER_PATTERNS = [
    re.compile(r"^第?\s*\d+\s*页$"),
    re.compile(r"^\d+\s*/\s*\d+$"),
    re.compile(r"^\d+$"),
    re.compile(r"^\d{4}\s*年\s*\d{1,2}\s*月$"),
    re.compile(r"^国家药品监督管理局药品审评中心$"),
    re.compile(r"^药品审评中心$"),
]

TOC_HINT_PATTERNS = [
    re.compile(r"^目录$"),
    re.compile(r"^contents$", re.I),
]

TOC_LINE_PATTERNS = [
    re.compile(r"^.+\.{2,}\s*\d+\s*$"),
    re.compile(rf"^[{CN_NUM}]+[、.]\s*.+\.{2,}\s*\d+\s*$"),
    re.compile(rf"^\([{CN_NUM}]+\)\s*.+\.{2,}\s*\d+\s*$"),
    re.compile(r"^\d+(?:\.\d+)*\s+.+\.{2,}\s*\d+\s*$"),
]


def is_header_or_footer_text(text: str) -> bool:
    t = normalize_text(text)
    if not t:
        return True
    return any(p.match(t) for p in HEADER_FOOTER_PATTERNS)


def is_toc_hint(text: str) -> bool:
    t = strip_trailing_colon(text)
    return any(p.match(t) for p in TOC_HINT_PATTERNS)


def is_toc_line(text: str) -> bool:
    t = normalize_text(text)
    if any(p.match(t) for p in TOC_LINE_PATTERNS):
        return True
    if re.match(r"^.+\s+\d+$", t) and len(t) <= 80 and extract_heading_info(t) is None:
        # 如 “附件 1：... 30”
        return True
    return False


def should_skip_toc_block(texts: List[str]) -> bool:
    cleaned = [normalize_text(t) for t in texts if normalize_text(t)]
    if not cleaned:
        return False
    if any(is_toc_hint(t) for t in cleaned[:5]):
        return True
    toc_like = sum(1 for t in cleaned if is_toc_line(t))
    return toc_like >= max(3, len(cleaned) // 2)


# ============================================================
# Markdown 文本压缩
# ============================================================

def markdownize_text_line(text: str) -> str:
    text = normalize_text(text)
    if not text:
        return ""

    m = re.match(r"^\((\d+)\)\s*(.+)$", text)
    if m:
        return f"1. {m.group(2).strip()}"

    m = re.match(rf"^([{CN_NUM}]+[、.])\s*(.+)$", text)
    if m:
        return f"- {m.group(2).strip()}"

    m = re.match(r"^([\-•·●])\s*(.+)$", text)
    if m:
        return f"- {m.group(2).strip()}"

    return text


def compact_markdown_parts(parts: List[Dict[str, str]]) -> str:
    blocks: List[str] = []
    text_buffer: List[str] = []

    def flush_text_buffer() -> None:
        nonlocal text_buffer, blocks
        if not text_buffer:
            return
        current_para: List[str] = []
        for line in text_buffer:
            line = markdownize_text_line(line)
            if not line:
                if current_para:
                    blocks.append("\n".join(current_para))
                    current_para = []
                continue
            if line.startswith("- ") or re.match(r"^\d+\.\s+", line):
                if current_para:
                    blocks.append(" ".join(current_para))
                    current_para = []
                blocks.append(line)
            else:
                current_para.append(line)
        if current_para:
            blocks.append(" ".join(current_para))
        text_buffer = []

    for part in parts:
        kind = part["kind"]
        value = part["value"].strip()
        if not value:
            continue
        if kind == "text":
            text_buffer.append(value)
        else:
            flush_text_buffer()
            blocks.append(value)

    flush_text_buffer()
    text = "\n\n".join(blocks)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


# ============================================================
# 章节树
# ============================================================

@dataclass
class SectionNode:
    id: str
    title: str
    level: int
    content_parts: List[Dict[str, str]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    children: List["SectionNode"] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "section_id": self.id,
            "section_title": self.title,
            "level": self.level,
            "content": compact_markdown_parts(self.content_parts),
            "tables": self.tables,
            "images": self.images,
            "children_sections": [c.to_dict() for c in self.children],
        }


class SectionTreeBuilder:
    def __init__(self) -> None:
        self.root = SectionNode(id="root", title="document", level=0)
        self.stack: List[SectionNode] = [self.root]
        self.counter = 0

    def add_heading(self, title: str, level: int) -> SectionNode:
        self.counter += 1
        node = SectionNode(id=f"sec_{self.counter:04d}", title=title, level=level)
        while len(self.stack) > 1 and self.stack[-1].level >= level:
            self.stack.pop()
        self.stack[-1].children.append(node)
        self.stack.append(node)
        return node

    def current_node(self) -> SectionNode:
        return self.stack[-1]

    def to_sections(self) -> List[Dict[str, Any]]:
        return [c.to_dict() for c in self.root.children]


# ============================================================
# DOC / DOCX 打开
# ============================================================

def convert_doc_to_docx(doc_path: Path) -> Path:
    out_dir = Path(tempfile.mkdtemp(prefix="doc_to_docx_"))
    cmd = [
        "soffice",
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(out_dir),
        str(doc_path),
    ]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except FileNotFoundError as e:
        raise RuntimeError("检测到 .doc 文件，但系统未安装 soffice/libreoffice，无法自动转换为 .docx") from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f".doc 转 .docx 失败: {e.stderr.decode('utf-8', errors='ignore')}") from e

    candidates = list(out_dir.glob("*.docx"))
    if not candidates:
        raise RuntimeError(".doc 转 .docx 后未找到输出文件")
    return candidates[0]


def load_document(input_path: Path) -> Tuple[Document, Path]:
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        return Document(str(input_path)), input_path
    if suffix == ".doc":
        converted = convert_doc_to_docx(input_path)
        return Document(str(converted)), converted
    raise ValueError(f"不支持的文件类型: {suffix}")


# ============================================================
# 遍历块级元素
# ============================================================

def iter_block_items(parent: _Document) -> Iterator[Any]:
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# ============================================================
# 段落 / 图片 / 表格提取
# ============================================================

def paragraph_text(paragraph: Paragraph) -> str:
    return normalize_text(paragraph.text)


def paragraph_images(paragraph: Paragraph, doc: Document, image_dir: Optional[Path], embed_images: bool, img_counter_start: int) -> Tuple[List[Dict[str, Any]], int]:
    images: List[Dict[str, Any]] = []
    counter = img_counter_start
    blips = paragraph._p.xpath('.//a:blip')
    for blip in blips:
        rid = blip.get(qn('r:embed'))
        if not rid:
            continue
        rel = doc.part.related_parts.get(rid)
        if rel is None:
            continue
        data = rel.blob
        counter += 1

        ext = Path(getattr(rel, 'partname', f'image_{counter}.bin')).suffix.lower() or '.bin'
        image_id = f"img_{counter:04d}"
        entry: Dict[str, Any] = {
            "id": image_id,
            "filename": f"{image_id}{ext}",
            "content_type": getattr(rel, 'content_type', ''),
        }

        if embed_images:
            entry["image_base64"] = image_bytes_to_base64(data)
            md_ref = f"embedded://{image_id}"
        else:
            if image_dir is not None:
                image_dir.mkdir(parents=True, exist_ok=True)
                out_path = image_dir / safe_filename(entry["filename"])
                out_path.write_bytes(data)
                entry["image_path"] = str(out_path)
                md_ref = out_path.name
            else:
                md_ref = entry["filename"]

        entry["markdown_ref"] = md_ref
        images.append(entry)
    return images, counter


def get_gridspan(tc) -> int:
    tcPr = tc.tcPr
    if tcPr is None:
        return 1
    el = tcPr.find(qn('w:gridSpan'))
    if el is None:
        return 1
    try:
        return max(1, int(el.get(qn('w:val')) or 1))
    except Exception:
        return 1


def get_vmerge_state(tc) -> Optional[str]:
    tcPr = tc.tcPr
    if tcPr is None:
        return None
    el = tcPr.find(qn('w:vMerge'))
    if el is None:
        return None
    val = el.get(qn('w:val'))
    return val or 'continue'


def extract_table_rows_with_merges(table: Table) -> List[List[str]]:
    rows: List[List[str]] = []
    vertical_memory: Dict[int, str] = {}

    for tr in table._tbl.tr_lst:
        cur_row: List[str] = []
        col_idx = 0
        for tc in tr.tc_lst:
            text = normalize_text("\n".join(p.text for p in tc.p_lst))
            grid_span = get_gridspan(tc)
            vmerge = get_vmerge_state(tc)

            if vmerge == 'continue':
                # 纵向合并的续行，继承上方内容
                inherited = vertical_memory.get(col_idx, "")
                for _ in range(grid_span):
                    cur_row.append(inherited)
                    col_idx += 1
                continue

            for _ in range(grid_span):
                cur_row.append(text)
                if vmerge == 'restart':
                    vertical_memory[col_idx] = text
                elif col_idx in vertical_memory and vmerge is None:
                    vertical_memory.pop(col_idx, None)
                col_idx += 1
        rows.append(cur_row)

    max_cols = max((len(r) for r in rows), default=0)
    return [r + [""] * (max_cols - len(r)) for r in rows]


def table_to_rows(table: Table) -> List[List[str]]:
    rows = extract_table_rows_with_merges(table)
    if not rows:
        rows = []
        for row in table.rows:
            values = [normalize_text(cell.text) for cell in row.cells]
            rows.append(values)
    max_cols = max((len(r) for r in rows), default=0)
    return [r + [""] * (max_cols - len(r)) for r in rows]


def find_table_header_end(rows: List[List[str]]) -> int:
    if not rows:
        return 0
    if len(rows) == 1:
        return 1

    # 表单/附件常有多行表头，前 1-3 行密集且后续明显进入数据/说明
    limit = min(3, len(rows))
    header_end = 1
    for i in range(limit):
        row = rows[i]
        nonempty = sum(1 for x in row if normalize_text(x))
        if nonempty >= max(2, len(row) // 2):
            header_end = i + 1
        else:
            break
    return header_end


def infer_table_title_from_context(context_paragraphs: List[str]) -> str:
    candidates = [strip_trailing_colon(x) for x in context_paragraphs if strip_trailing_colon(x)]
    if not candidates:
        return ""
    for t in reversed(candidates[-3:]):
        if is_appendix_like_title(t) or len(t) <= 60:
            return t
    return candidates[-1]


def classify_table_kind(title: str, rows: List[List[str]]) -> str:
    title_n = normalize_text(title)
    joined = " ".join(normalize_text(x) for row in rows[:4] for x in row if normalize_text(x))
    if re.search(r"(批分析数据|杂质谱分析|信息汇总表|汇总表|信息表|工艺流程图|生产工艺|主要仪器设备)", title_n + " " + joined):
        return "appendix_form"
    return "table"


def table_to_frame(table: Table, table_idx: int, context_paragraphs: Optional[List[str]] = None) -> Dict[str, Any]:
    rows = table_to_rows(table)
    rows = [r for r in rows if any(normalize_text(x) for x in r)]
    title = infer_table_title_from_context(context_paragraphs or [])

    if not rows:
        return {
            "id": f"tbl_{table_idx:04d}",
            "type": "table",
            "title": title,
            "columns": [],
            "data": [],
            "markdown": "",
        }

    header_end = find_table_header_end(rows)
    header_rows = rows[:header_end]
    data_rows = rows[header_end:] if header_end < len(rows) else []

    if header_rows:
        max_cols = max(len(r) for r in rows)
        columns: List[str] = []
        for c in range(max_cols):
            parts = []
            for hr in header_rows:
                v = normalize_text(hr[c] if c < len(hr) else "")
                if v and (not parts or v != parts[-1]):
                    parts.append(v)
            columns.append(" / ".join(parts).strip(" /"))
    else:
        columns = [f"col_{i+1}" for i in range(max(len(r) for r in rows), 0)]

    if sum(1 for x in columns if normalize_text(x)) == 0:
        columns = [f"col_{i+1}" for i in range(max(len(r) for r in rows), 0)]
        data_rows = rows

    notes: List[str] = []
    actual_data: List[List[str]] = []
    for r in data_rows:
        joined = normalize_text(" ".join(x for x in r if normalize_text(x)))
        if len(r) <= 2 and joined and (joined.startswith("*") or joined.startswith("注") or joined.startswith("备注")):
            notes.append(joined)
        else:
            actual_data.append(r)

    frame = {
        "id": f"tbl_{table_idx:04d}",
        "type": classify_table_kind(title, rows),
        "title": title,
        "header_rows": header_rows,
        "columns": columns,
        "data": actual_data,
        "notes": notes,
        "markdown": markdown_table(columns, actual_data),
    }
    return frame


def tables_can_merge(prev_frame: Dict[str, Any], cur_frame: Dict[str, Any]) -> bool:
    if not prev_frame or not cur_frame:
        return False
    if prev_frame.get("type") != cur_frame.get("type"):
        return False
    prev_cols = [normalize_text(x) for x in prev_frame.get("columns", [])]
    cur_cols = [normalize_text(x) for x in cur_frame.get("columns", [])]
    if not prev_cols or not cur_cols:
        return False
    if prev_cols == cur_cols:
        return True
    # 表头高度一致且前若干列一致
    if len(prev_cols) == len(cur_cols):
        same = sum(1 for a, b in zip(prev_cols, cur_cols) if a == b)
        if same >= max(2, len(prev_cols) - 1):
            return True
    return False


def merge_table_frames(prev_frame: Dict[str, Any], cur_frame: Dict[str, Any]) -> Dict[str, Any]:
    merged = copy.deepcopy(prev_frame)
    merged["data"] = prev_frame.get("data", []) + cur_frame.get("data", [])
    merged["notes"] = prev_frame.get("notes", []) + cur_frame.get("notes", [])
    merged["continued_from_previous"] = True
    merged["markdown"] = markdown_table(merged.get("columns", []), merged.get("data", []))
    return merged


# ============================================================
# 主逻辑
# ============================================================

def parse_docx_to_markdown_json(
    input_path: Path,
    title: Optional[str] = None,
    embed_images: bool = False,
    skip_toc: bool = True,
    filter_header_footer: bool = True,
    merge_continuous_tables: bool = True,
) -> Dict[str, Any]:
    doc, real_docx_path = load_document(input_path)
    builder = SectionTreeBuilder()
    image_dir = input_path.parent / f"{input_path.stem}_images" if not embed_images else None

    doc_title = title
    paragraph_idx = 0
    image_counter = 0
    table_counter = 0

    recent_texts: List[str] = []
    pending_context_paragraphs: List[str] = []
    skip_toc_mode = False

    prev_table_node: Optional[SectionNode] = None
    prev_table_id: Optional[str] = None

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            paragraph_idx += 1
            txt = paragraph_text(block)
            style_name = block.style.name if block.style is not None else ""

            if txt and not doc_title and paragraph_idx <= 3 and len(txt) <= 100:
                doc_title = txt

            if filter_header_footer and txt and is_header_or_footer_text(txt):
                recent_texts.append(txt)
                recent_texts = recent_texts[-20:]
                continue

            if skip_toc and txt:
                if is_toc_hint(txt):
                    skip_toc_mode = True
                    recent_texts.append(txt)
                    recent_texts = recent_texts[-20:]
                    continue
                if skip_toc_mode:
                    if is_toc_line(txt) or not txt:
                        recent_texts.append(txt)
                        recent_texts = recent_texts[-20:]
                        continue
                    else:
                        skip_toc_mode = False

            heading = is_probable_heading(txt, style_name) if txt else None
            if heading is not None:
                builder.add_heading(heading.title, heading.level)
                pending_context_paragraphs = [txt]
                prev_table_node = None
                prev_table_id = None
            else:
                if txt:
                    node = builder.current_node()
                    node.content_parts.append({"kind": "text", "value": txt})
                    pending_context_paragraphs.append(txt)
                    pending_context_paragraphs = pending_context_paragraphs[-5:]
                    prev_table_node = None
                    prev_table_id = None

            para_images, image_counter = paragraph_images(
                block, doc, image_dir=image_dir, embed_images=embed_images, img_counter_start=image_counter
            )
            if para_images:
                node = builder.current_node()
                for img in para_images:
                    node.images.append(img)
                    alt = img.get("filename", img["id"])
                    ref = img.get("markdown_ref", "")
                    node.content_parts.append({"kind": "image", "value": f"![{alt}]({ref})"})
                prev_table_node = None
                prev_table_id = None

            recent_texts.append(txt)
            recent_texts = recent_texts[-20:]

        elif isinstance(block, Table):
            table_counter += 1
            node = builder.current_node()
            frame = table_to_frame(block, table_counter, context_paragraphs=pending_context_paragraphs)

            if merge_continuous_tables and prev_table_node is node and node.tables and tables_can_merge(node.tables[-1], frame):
                node.tables[-1] = merge_table_frames(node.tables[-1], frame)
                # 回写 content_parts 中最后一个 table markdown
                for j in range(len(node.content_parts) - 1, -1, -1):
                    if node.content_parts[j]["kind"] == "table":
                        node.content_parts[j] = {"kind": "table", "value": node.tables[-1]["markdown"]}
                        break
            else:
                node.tables.append(frame)
                if frame.get("title") and frame["type"] == "appendix_form":
                    node.content_parts.append({"kind": "text", "value": frame["title"]})
                if frame.get("markdown"):
                    node.content_parts.append({"kind": "table", "value": frame["markdown"]})
                if frame.get("notes"):
                    for note in frame["notes"]:
                        node.content_parts.append({"kind": "text", "value": note})

            prev_table_node = node
            prev_table_id = frame["id"]
            pending_context_paragraphs = pending_context_paragraphs[-3:]

    sections = builder.to_sections()

    if not sections:
        fallback = SectionNode(id="sec_0001", title="全文", level=1)
        pending_context: List[str] = []
        image_counter = 0
        table_counter = 0
        prev_fallback_table: Optional[Dict[str, Any]] = None
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                txt = paragraph_text(block)
                if filter_header_footer and txt and is_header_or_footer_text(txt):
                    continue
                if skip_toc and (is_toc_hint(txt) or is_toc_line(txt)):
                    continue
                if txt:
                    fallback.content_parts.append({"kind": "text", "value": txt})
                    pending_context.append(txt)
                    pending_context = pending_context[-5:]
                para_images, image_counter = paragraph_images(
                    block, doc, image_dir=image_dir, embed_images=embed_images, img_counter_start=image_counter
                )
                for img in para_images:
                    fallback.images.append(img)
                    alt = img.get("filename", img["id"])
                    ref = img.get("markdown_ref", "")
                    fallback.content_parts.append({"kind": "image", "value": f"![{alt}]({ref})"})
            elif isinstance(block, Table):
                table_counter += 1
                frame = table_to_frame(block, table_counter, context_paragraphs=pending_context)
                if merge_continuous_tables and prev_fallback_table is not None and tables_can_merge(prev_fallback_table, frame):
                    prev_fallback_table = merge_table_frames(prev_fallback_table, frame)
                    fallback.tables[-1] = prev_fallback_table
                    for j in range(len(fallback.content_parts) - 1, -1, -1):
                        if fallback.content_parts[j]["kind"] == "table":
                            fallback.content_parts[j] = {"kind": "table", "value": prev_fallback_table["markdown"]}
                            break
                else:
                    fallback.tables.append(frame)
                    prev_fallback_table = frame
                    if frame.get("title") and frame["type"] == "appendix_form":
                        fallback.content_parts.append({"kind": "text", "value": frame["title"]})
                    if frame.get("markdown"):
                        fallback.content_parts.append({"kind": "table", "value": frame["markdown"]})
                    if frame.get("notes"):
                        for note in frame["notes"]:
                            fallback.content_parts.append({"kind": "text", "value": note})
        sections = [fallback.to_dict()]

    return {
        "title": doc_title or input_path.stem,
        "source_file": str(input_path),
        "resolved_docx": str(real_docx_path),
        "structure_type": "generic_heading_based_markdown_json_for_docx_v2",
        "parser": {
            "heading_strategy": "word_heading_style_plus_regex_fallback",
            "chunking_strategy": "rough_section_chunking_by_detected_headings",
            "skip_toc": skip_toc,
            "filter_header_footer": filter_header_footer,
            "merge_continuous_tables": merge_continuous_tables,
            "embed_images": embed_images,
        },
        "sections": sections,
    }


# ============================================================
# CLI
# ============================================================

def derive_default_output(input_path: Path) -> Path:
    return input_path.parent / f"{input_path.stem}_parsed_markdown.json"


def main() -> None:
    parser = argparse.ArgumentParser(description="通用 DOCX / DOC -> 章节 Markdown JSON")
    parser.add_argument("input", help="输入 Word 文件路径（.docx / .doc）")
    parser.add_argument("--output-json", help="输出 JSON 路径")
    parser.add_argument("--title", default=None, help="文档标题")
    parser.add_argument("--embed-images", default="false", help="是否将图片以 base64 内嵌到 JSON：true/false")
    parser.add_argument("--skip-toc", default="true", help="是否跳过目录：true/false")
    parser.add_argument("--filter-header-footer", default="true", help="是否过滤页眉页脚：true/false")
    parser.add_argument("--merge-continuous-tables", default="true", help="是否合并连续表格：true/false")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        raise FileNotFoundError(f"输入文件不存在: {input_path}")

    output_json = Path(args.output_json).expanduser().resolve() if args.output_json else derive_default_output(input_path)

    result = parse_docx_to_markdown_json(
        input_path=input_path,
        title=args.title,
        embed_images=str2bool(args.embed_images),
        skip_toc=str2bool(args.skip_toc),
        filter_header_footer=str2bool(args.filter_header_footer),
        merge_continuous_tables=str2bool(args.merge_continuous_tables),
    )

    output_json.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[OK] JSON: {output_json}")


if __name__ == "__main__":
    main()
