import json
import os
import re
import threading
import uuid
import zipfile
from copy import deepcopy
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from sqlalchemy import and_, desc, func
from werkzeug.utils import secure_filename

from agent.agent_backend.agents.pre_review_agent.pre_review_state import SectionReviewPacket
from agent.agent_backend.agents.review.consistency_agent import ConsistencyAgent
from agent.agent_backend.agents.review.feedback_agent import FeedbackAgent
from agent.agent_backend.agents.review.lead_reviewer_agent import LeadReviewerAgent
from agent.agent_backend.agents.review.planner_reviewer_agent import PlannerReviewerAgent
from agent.agent_backend.agents.review.qa_agent import QAAgent
from agent.agent_backend.agents.review.run_review_workflow import RunReviewWorkflow
from agent.agent_backend.agents.review.submission_parser_agent import SubmissionParserAgent
from agent.agent_backend.common_tools.builtin.memory_tool import MemoryTool
from agent.agent_backend.config.settings import settings
from agent.agent_backend.database.mysql.db_model import (
    FileInfo,
    PreReviewFeedback,
    PreReviewProject,
    PreReviewRun,
    PreReviewSectionConcern,
    PreReviewPromptRule,
    PreReviewSectionConclusion,
    PreReviewSectionOutput,
    PreReviewFeedbackAnalysisResult,
    PreReviewPatchRegistry,
    PreReviewSubmissionSectionContent,
    PreReviewSubmissionFile,
    PreReviewExperienceMemory,
    PreReviewSectionTrace,
)
from agent.agent_backend.database.mysql.mysql_conn import MysqlConnection
from agent.agent_backend.feedback.optimization.prompt_version_registry import PromptVersionRegistry
from agent.agent_backend.infrastructure.repositories.pre_review_repository import (
    PreReviewRepository,
    SectionConclusionRecord,
    SectionTraceRecord,
)
from agent.agent_backend.agentic_rl.reward_functions import feedback_metrics
from agent.agent_backend.services.ctd_section_service import CTDSectionService
from agent.agent_backend.services.knowledge_service import KnowledgeService
from agent.agent_backend.services.pharmacopeia_service import PharmacopeiaService
from agent.agent_backend.services.pre_review_prompt_rule_service import PreReviewPromptRuleService
from agent.agent_backend.utils.file_util import ensure_dir_exists
from agent.agent_backend.memory.storage.vector_store import VectorStore
from agent.agent_backend.utils.parser import ParserManager
from agent.agent_backend.utils.parser.ctd_api_markdown_parser import parse_ctd_api_pdf_to_payload
from agent.agent_backend.utils.parser.coarse_ctd_submission_parser import parse_coarse_ctd_submission
from agent.agent_backend.utils.parser.strict_ctd_submission_parser import parse_strict_ctd_submission
from agent.agent_backend.utils.parser.submission_pdf_markdown_parser import parse_submission_pdf_to_payload


SUBMISSION_UPLOAD_DIR = settings.submission_upload_dir
SUBMISSION_PARSED_DIR = settings.submission_parse_dir
SUBMISSION_EDIT_DIR = str(Path(__file__).resolve().parents[1] / "data" / "submission_edits")
REPORT_DIR = settings.report_dir
RAW_DATA_DIR = str(Path(__file__).resolve().parents[1] / "data" / "raw_data")
ensure_dir_exists(SUBMISSION_UPLOAD_DIR)
ensure_dir_exists(SUBMISSION_PARSED_DIR)
ensure_dir_exists(SUBMISSION_EDIT_DIR)
ensure_dir_exists(REPORT_DIR)

_SERVICE_DEBUG_ONCE_KEYS: set[str] = set()

SINGLE_SECTION_PRE_REVIEW_V2 = "single_section_pre_review_v2"
RETRIEVAL_SOURCE_GUIDANCE = "指导原则"
RETRIEVAL_SOURCE_ICH = "ICH"
RETRIEVAL_SOURCE_REGULATION = "法律法规"
RETRIEVAL_SOURCE_PHARMACOPOEIA = "药典数据"
RETRIEVAL_SOURCE_EXPERIENCE = "历史经验"

_RETRIEVAL_SOURCE_ALIASES = {
    RETRIEVAL_SOURCE_GUIDANCE: {
        RETRIEVAL_SOURCE_GUIDANCE,
        "guidance",
        "guidance_principles",
        "guideline",
        "guidelines",
    },
    RETRIEVAL_SOURCE_ICH: {
        RETRIEVAL_SOURCE_ICH,
        "ich",
    },
    RETRIEVAL_SOURCE_REGULATION: {
        RETRIEVAL_SOURCE_REGULATION,
        "法规",
        "法律",
        "regulation",
        "regulations",
        "law",
        "law_and_regulation",
    },
    RETRIEVAL_SOURCE_PHARMACOPOEIA: {
        RETRIEVAL_SOURCE_PHARMACOPOEIA,
        "药典",
        "pharmacopeia",
        "pharmacopoeia",
    },
    RETRIEVAL_SOURCE_EXPERIENCE: {
        RETRIEVAL_SOURCE_EXPERIENCE,
        "experience",
        "historical_experience",
    },
}

RETRIEVAL_FALSE_NEGATIVE_ERRORS = {
    "query_miss",
    "historical_experience_missing",
}

RETRIEVAL_FALSE_POSITIVE_ERRORS = {
    "retrieval_scope_error",
    "retrieval_ranking_error",
}


def _service_debug_once(key: str, message: str) -> None:
    if key in _SERVICE_DEBUG_ONCE_KEYS:
        return
    _SERVICE_DEBUG_ONCE_KEYS.add(key)
    print(message)


def _safe_ratio(numerator: float, denominator: float) -> float:
    if denominator <= 0:
        return 0.0
    return float(numerator) / float(denominator)


def _safe_f1(precision: float, recall: float) -> float:
    if precision + recall <= 0:
        return 0.0
    return 2.0 * precision * recall / (precision + recall)


class PreReviewService:
    def __init__(self):
        _service_debug_once("PreReviewService.__init__", "[DEBUG] enter PreReviewService.__init__ | core: {}")
        self.db_conn = MysqlConnection()
        self.memory_tool = MemoryTool()
        self.planner_reviewer_agent = PlannerReviewerAgent()
        self.feedback_agent = FeedbackAgent()
        self.consistency_agent = ConsistencyAgent()
        self.qa_agent = QAAgent()
        self.lead_reviewer_agent = LeadReviewerAgent()
        self.submission_parser_agent = SubmissionParserAgent()
        self.run_review_workflow = RunReviewWorkflow(
            consistency_agent=self.consistency_agent,
            qa_agent=self.qa_agent,
            lead_reviewer_agent=self.lead_reviewer_agent,
        )
        self.knowledge = KnowledgeService()
        self.pharmacopeia = PharmacopeiaService()
        self.ctd_sections = CTDSectionService(raw_data_dir=RAW_DATA_DIR)
        self.prompt_registry = PromptVersionRegistry()
        self.prompt_rule_service = PreReviewPromptRuleService()
        self.submission_vector_store = VectorStore()
        self._submission_payload_cache: Dict[str, Any] = {}

    @staticmethod
    def _now() -> datetime:
        return datetime.now()

    @staticmethod
    def _new_project_id() -> str:
        print("[DEBUG] enter PreReviewService._new_project_id | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        return f"prj_{uuid.uuid4().hex[:12]}"

    @staticmethod
    def _new_run_id() -> str:
        print("[DEBUG] enter PreReviewService._new_run_id | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        return f"run_{uuid.uuid4().hex[:12]}"

    @staticmethod
    def _new_submission_doc_id() -> str:
        print("[DEBUG] enter PreReviewService._new_submission_doc_id | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        return f"sub_{uuid.uuid4().hex[:16]}"

    @staticmethod
    def _safe_display_name(file_name: str) -> str:
        print("[DEBUG] enter PreReviewService._safe_display_name | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        return os.path.basename(str(file_name or "")).strip()

    @staticmethod
    def _submission_storage_name(doc_id: str, display_name: str) -> str:
        print("[DEBUG] enter PreReviewService._submission_storage_name | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        ext = ""
        if "." in display_name:
            ext = display_name.rsplit(".", 1)[-1].lower().strip()
        safe = secure_filename(display_name) or "submission"
        if ext and "." not in safe:
            safe = f"{safe}.{ext}"
        return f"{doc_id}_{safe}"

    @staticmethod
    def _submission_payload_cache_key(project_id: str, doc_id: str) -> str:
        return f"{str(project_id or '').strip()}::{str(doc_id or '').strip()}"

    @staticmethod
    def _is_zip_file_name(file_name: str) -> bool:
        return str(file_name or "").lower().endswith(".zip")

    @staticmethod
    def _decode_zip_name(raw_name: str) -> str:
        try:
            repaired = raw_name.encode("cp437").decode("gb18030")
            return repaired
        except Exception:
            return raw_name

    @staticmethod
    def _sanitize_member_name(raw_name: str) -> str:
        normalized = str(raw_name or "").replace("\\", "/").strip().strip("/")
        parts = [p for p in normalized.split("/") if p not in {"", ".", ".."}]
        return "/".join(parts)

    @staticmethod
    def _preview_text_blocks(blocks: List[str]) -> str:
        merged = "\n\n".join([str(x).strip() for x in blocks if str(x).strip()]).strip()
        return merged

    @staticmethod
    def _safe_json_list(value: Any) -> str:
        return json.dumps(value if isinstance(value, list) else [], ensure_ascii=False)

    @staticmethod
    def _dedupe_text_list(values: List[Any]) -> List[str]:
        out: List[str] = []
        seen = set()
        for item in values or []:
            value = str(item or "").strip()
            if not value or value in seen:
                continue
            seen.add(value)
            out.append(value)
        return out

    @staticmethod
    def _submission_edit_path(doc_id: str) -> str:
        return os.path.join(SUBMISSION_EDIT_DIR, f"{doc_id}.txt")

    def _load_submission_edit(self, doc_id: str) -> str:
        path = self._submission_edit_path(doc_id)
        if not os.path.exists(path):
            return ""
        try:
            with open(path, "r", encoding="utf-8") as fh:
                return fh.read()
        except Exception:
            return ""

    def _save_submission_edit(self, doc_id: str, content: str) -> str:
        path = self._submission_edit_path(doc_id)
        with open(path, "w", encoding="utf-8") as fh:
            fh.write(str(content or ""))
        return path


    def _load_manual_concern_map(self, session, project_id: str) -> Dict[str, List[str]]:
        rows = (
            session.query(PreReviewSectionConcern)
            .filter(PreReviewSectionConcern.project_id == project_id)
            .all()
        )
        out: Dict[str, List[str]] = {}
        for row in rows:
            out[str(row.section_id)] = self._dedupe_text_list(self._parse_json_list(row.concern_points))
        return out

    def _merge_catalog_with_manual_concerns(self, catalog: Dict[str, Any], manual_map: Dict[str, List[str]]) -> Dict[str, Any]:
        chapter_structure = deepcopy(catalog.get("chapter_structure", []))

        def attach(nodes: List[Dict[str, Any]]):
            for node in nodes or []:
                sid = str(node.get("section_id", "")).strip()
                merged = self._dedupe_text_list(list(node.get("concern_points") or []) + manual_map.get(sid, []))
                node["concern_points"] = merged
                attach(node.get("children_sections") or [])

        attach(chapter_structure)
        flat_sections = self.ctd_sections._flatten_nodes(chapter_structure)
        return {
            "chapter_structure": chapter_structure,
            "flat_sections": flat_sections,
            "section_map": {item["section_id"]: item for item in flat_sections},
        }

    @staticmethod
    def _is_ctd_structure_project(project: PreReviewProject) -> bool:
        return str(getattr(project, "registration_scope", "") or "").strip() == "鍖栬嵂"

    @staticmethod
    def _branch_root_from_section_id(section_id: str) -> str:
        value = str(section_id or "").strip()
        if value.startswith("3.2.S"):
            return "3.2.S"
        if value.startswith("3.2.P"):
            return "3.2.P"
        return ""

    def _infer_submission_branch_root(
        self,
        material_category: str,
        explicit_section_id: str,
        section_meta: Optional[Dict[str, Any]] = None,
        display_name: str = "",
        relative_path: str = "",
    ) -> str:
        for candidate in [
            explicit_section_id,
            str((section_meta or {}).get("section_id", "") or ""),
            relative_path,
            display_name,
        ]:
            root = self._branch_root_from_section_id(candidate)
            if root:
                return root
        material_text = str(material_category or "")
        lowered = material_text.strip().lower()
        if "api" in lowered or "鍘熸枡" in material_text:
            return "3.2.S"
        if "fpp" in lowered or "鍒跺墏" in material_text:
            return "3.2.P"
        return ""

    def _extract_section_content_rows_from_payload(
        self,
        payload: Any,
        fallback_section_id: str = "",
        fallback_section_name: str = "",
    ) -> List[Dict[str, Any]]:
        sections: List[Dict[str, Any]] = []
        if isinstance(payload, dict) and isinstance(payload.get("sections"), list):
            sections = [item for item in payload.get("sections", []) if isinstance(item, dict)]
        elif isinstance(payload, dict) and isinstance(payload.get("review_units"), list):
            for item in payload.get("review_units", []):
                if not isinstance(item, dict):
                    continue
                sections.append(
                    {
                        "section_id": item.get("section_id") or item.get("chunk_id"),
                        "code": item.get("section_code") or item.get("section_id") or item.get("chunk_id"),
                        "section_name": item.get("section_name") or item.get("title"),
                        "content": item.get("text") or item.get("content") or "",
                    }
                )
        elif isinstance(payload, list):
            merged_text = self._preview_text_blocks(
                [str(item.get("text", "")).strip() for item in payload if isinstance(item, dict)]
            )
            if fallback_section_id and merged_text:
                sections = [
                    {
                        "section_id": fallback_section_id,
                        "code": fallback_section_id,
                        "section_name": fallback_section_name or fallback_section_id,
                        "content": merged_text,
                    }
                ]

        out: List[Dict[str, Any]] = []
        seen = set()
        for item in sections:
            section_id = str(item.get("section_id", "") or "").strip()
            if not section_id:
                continue
            content = str(item.get("content") or item.get("text") or "").strip()
            if not content:
                continue
            key = (section_id, content)
            if key in seen:
                continue
            seen.add(key)
            out.append(
                {
                    "section_id": section_id,
                    "section_code": str(item.get("code") or item.get("section_code") or section_id).strip() or section_id,
                    "section_name": str(item.get("section_name") or item.get("title") or section_id).strip() or section_id,
                    "content": content,
                    "content_preview": self._preview(content, 320),
                }
            )
        return out

    def _sync_submission_section_content_rows(
        self,
        session,
        file_row: PreReviewSubmissionFile,
        payload: Any,
    ) -> None:
        session.query(PreReviewSubmissionSectionContent).filter(
            PreReviewSubmissionSectionContent.doc_id == file_row.doc_id
        ).delete(synchronize_session=False)

        section_rows = self._extract_section_content_rows_from_payload(
            payload=payload,
            fallback_section_id=str(getattr(file_row, "section_id", "") or ""),
            fallback_section_name=str(getattr(file_row, "section_name", "") or ""),
        )
        payload_type = str(payload.get("structure_type", "")).strip() if isinstance(payload, dict) else ""
        if payload_type == "ctd_api_markdown_json":
            parser_name = "ctd_api_markdown"
        elif payload_type == "ctd_fixed_outline_coarse_payload_v1":
            parser_name = "ctd_outline_coarse"
        elif payload_type == "generic_heading_based_markdown_json_v2":
            parser_name = "submission_pdf_markdown"
        else:
            parser_name = "strict_ctd"
        now = self._now()
        for item in section_rows:
            session.add(
                PreReviewSubmissionSectionContent(
                    doc_id=file_row.doc_id,
                    project_id=file_row.project_id,
                    section_id=str(item.get("section_id", "")).strip(),
                    section_code=str(item.get("section_code", "")).strip(),
                    section_name=str(item.get("section_name", "")).strip(),
                    content=str(item.get("content", "")).strip(),
                    content_preview=str(item.get("content_preview", "")).strip(),
                    source_parser=parser_name if str(item.get("section_id", "")).startswith("3.2.") else "generic",
                    create_time=now,
                    update_time=now,
                )
            )

    def _build_submission_vector_units(
        self,
        payload: Any,
        fallback_section_id: str = "",
        fallback_section_name: str = "",
    ) -> List[Dict[str, Any]]:
        units = payload.get("review_units", []) if isinstance(payload, dict) else []
        if isinstance(units, list) and units:
            out: List[Dict[str, Any]] = []
            for item in units:
                if not isinstance(item, dict):
                    continue
                out.extend(self._split_submission_vector_unit(item))
            return out
        section_rows = self._extract_section_content_rows_from_payload(
            payload=payload,
            fallback_section_id=fallback_section_id,
            fallback_section_name=fallback_section_name,
        )
        out: List[Dict[str, Any]] = []
        for idx, item in enumerate(section_rows, start=1):
            text = str(item.get("content", "")).strip()
            if not text:
                continue
            section_id = str(item.get("section_id", "")).strip()
            section_code = str(item.get("section_code", "")).strip() or section_id
            section_name = str(item.get("section_name", "")).strip() or section_code
            out.append(
                {
                    "chunk_id": f"{section_id}_content",
                    "section_id": section_id,
                    "section_code": section_code,
                    "section_name": section_name,
                    "page": None,
                    "page_start": None,
                    "page_end": None,
                    "char_count": len(text),
                    "text": text,
                    "unit_type": "section_content",
                    "source_section_codes": [section_code],
                    "title_path": [section_name],
                    "unit_order": idx,
                }
            )
        return out

    def _index_submission_payload(
        self,
        file_row: PreReviewSubmissionFile,
        payload: Any,
        force_reindex: bool = False,
    ) -> Tuple[bool, str, int]:
        units = self._build_submission_vector_units(
            payload=payload,
            fallback_section_id=str(getattr(file_row, "section_id", "") or ""),
            fallback_section_name=str(getattr(file_row, "section_name", "") or ""),
        )
        if not units:
            return True, "no review units to index", 0
        try:
            doc_id = str(file_row.doc_id or "").strip()
            if not force_reindex and self.submission_vector_store.has_doc(doc_id):
                return True, "already indexed", len(units)
            self.submission_vector_store.delete_by_doc(doc_id)
            for idx, unit in enumerate(units, start=1):
                text = str(unit.get("text", "")).strip()
                if not text:
                    continue
                chunk_id = str(unit.get("chunk_id") or f"{doc_id}_chunk_{idx}").strip()
                vec_id = f"{doc_id}:{chunk_id}"
                metadata = {
                    "doc_id": doc_id,
                    "chunk_id": chunk_id,
                    "chunk_order": int(unit.get("unit_order") or idx),
                    "item_type": "submission_chunk",
                    "classification": "submission_material",
                    "project_id": str(getattr(file_row, "project_id", "") or ""),
                    "source_domain": "pre_review_submission",
                    "material_category": str(getattr(file_row, "material_category", "") or "other"),
                    "section_id": str(unit.get("section_id", "") or ""),
                    "section_code": str(unit.get("section_code", "") or ""),
                    "section_name": str(unit.get("section_name", "") or ""),
                    "unit_type": str(unit.get("unit_type", "") or ""),
                    "page": unit.get("page"),
                    "page_start": unit.get("page_start"),
                    "page_end": unit.get("page_end"),
                    "title_path": list(unit.get("title_path") or []),
                    "source_parser": str(payload.get("structure_type", "")) if isinstance(payload, dict) else "",
                    "summary": self._preview(text, 240),
                    "keywords": self._dedupe_text_list(
                        [
                            str(unit.get("section_code", "")).strip(),
                            str(unit.get("section_name", "")).strip(),
                            str(getattr(file_row, "material_category", "") or "").strip(),
                        ]
                    ),
                }
                self.submission_vector_store.add_text(vec_id=vec_id, text=text, metadata=metadata)
            return True, "ok", len(units)
        except Exception as exc:
            return False, f"submission vector indexing failed: {str(exc)}", 0

    @staticmethod
    def _rule_label(rule_item: Dict[str, Any]) -> str:
        """
        Semantic retrieval currently returns doc_id/content/score fields.
        Build a stable human-readable label for linkage display.
        """
        print("[DEBUG] enter PreReviewService._rule_label | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        doc_id = str(rule_item.get("doc_id", "")).strip()
        classification = str(rule_item.get("classification", "")).strip()
        score = rule_item.get("score", None)
        parts = [p for p in [doc_id, classification] if p]
        label = "/".join(parts) if parts else "rule"
        if score is not None:
            try:
                label = f"{label}(score={float(score):.3f})"
            except Exception:
                pass
        return label

    @staticmethod
    def _preview(value: Any, max_len: int = 180) -> str:
        text = str(value or "").replace("\n", " ").strip()
        return text if len(text) <= max_len else f"{text[:max_len]}..."

    @staticmethod
    def _parse_json_list(raw_value: Any) -> List[str]:
        if isinstance(raw_value, list):
            return [str(x).strip() for x in raw_value if str(x).strip()]
        text = str(raw_value or "").strip()
        if not text:
            return []
        try:
            data = json.loads(text)
            if isinstance(data, list):
                return [str(x).strip() for x in data if str(x).strip()]
        except Exception:
            pass
        return [text]

    @staticmethod
    def _code_sort_key(code: str) -> Tuple:
        parts = str(code or "").split(".")
        key = []
        for p in parts:
            if p.isdigit():
                key.append((0, int(p)))
            else:
                key.append((1, p))
        return tuple(key)

    def _order_review_units(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        return sorted(
            chunks or [],
            key=lambda x: (
                int(x.get("unit_order") or 10**9),
                self._code_sort_key(str(x.get("section_code", ""))),
                int(x.get("page_start") or x.get("page") or 10**9),
                str(x.get("section_id") or x.get("chunk_id") or ""),
            ),
        )

    def _build_section_query(
        self,
        section_name: str,
        section_code: str,
        text: str,
        title_path: List[Any],
        concern_points: Optional[List[Any]] = None,
        section_summary: Optional[Dict[str, Any]] = None,
    ) -> str:
        path = " > ".join([str(x).strip() for x in (title_path or []) if str(x).strip()])
        query_parts = [f"??: {section_code} {section_name}".strip()]
        if path:
            query_parts.append(f"??: {path}")
        concerns = [str(x).strip() for x in (concern_points or []) if str(x).strip()]
        if concerns:
            query_parts.append(f"???: {'?'.join(concerns[:8])}")
        if isinstance(section_summary, dict):
            summary_text = str(section_summary.get("structured_summary", "") or "").strip()
            if summary_text:
                query_parts.append(f"???: {self._preview(summary_text, 240)}")
            key_facts = [str(x).strip() for x in section_summary.get("key_facts", []) if str(x).strip()] if isinstance(section_summary.get("key_facts", []), list) else []
            if key_facts:
                query_parts.append(f"???: {'; '.join(key_facts[:8])}")
        query_parts.append(f"??: {self._preview(text, 260)}")
        return "\n".join(query_parts)

    def _merge_retrieval_queries_with_focus_points(
        self,
        query_list: List[Any],
        focus_points: Optional[List[Any]] = None,
        max_queries: int = 6,
    ) -> List[str]:
        merged: List[str] = []
        seen = set()
        for item in self._normalize_text_list(query_list):
            text = self._sanitize_retrieval_query(item)
            if not text or text in seen:
                continue
            seen.add(text)
            merged.append(text)
            if len(merged) >= max_queries:
                return merged
        for point in self._normalize_text_list(focus_points or []):
            compact_point = self._sanitize_retrieval_query(point, max_len=64)
            if not compact_point:
                continue
            variants = [compact_point]
            for base_query in merged[:2]:
                combined = self._sanitize_retrieval_query(f"{base_query} 关注点 {compact_point}".strip(), max_len=120)
                if combined:
                    variants.append(combined)
            for variant in variants:
                text = str(variant or "").strip()
                if not text or text in seen:
                    continue
                seen.add(text)
                merged.append(text)
                if len(merged) >= max_queries:
                    return merged
        return merged


    def _normalize_finding_dicts(self, findings: List[Any]) -> List[Dict[str, Any]]:
        normalized: List[Dict[str, Any]] = []
        for item in findings or []:
            if isinstance(item, dict):
                title = str(item.get("title", "") or "").strip()
                if not title:
                    continue
                normalized.append(
                    {
                        "title": title,
                        "problem_type": str(item.get("problem_type", "") or "").strip(),
                        "severity": str(item.get("severity", "low") or "low").strip().lower(),
                        "evidence": str(item.get("evidence", "") or "").strip(),
                        "recommendation": str(item.get("recommendation", "") or "").strip(),
                        "metadata": item.get("metadata", {}) if isinstance(item.get("metadata", {}), dict) else {},
                    }
                )
                continue
            title = str(item or "").strip()
            if title:
                normalized.append(
                    {
                        "title": title,
                        "problem_type": "",
                        "severity": "low",
                        "evidence": "",
                        "recommendation": "",
                        "metadata": {},
                    }
                )
        return normalized

    def _findings_to_titles(self, findings: List[Any]) -> List[str]:
        return [item.get("title", "") for item in self._normalize_finding_dicts(findings) if item.get("title", "")]

    def _render_findings_summary(self, findings: List[Any]) -> str:
        titles = self._findings_to_titles(findings)
        return "；".join(titles) if titles else "未发现明显问题"

    @staticmethod
    def _emit_progress(progress_callback, stage: str, message: str, **payload: Any) -> None:
        if not callable(progress_callback):
            return
        data = {"stage": str(stage or "").strip(), "message": str(message or "").strip()}
        data.update(payload or {})
        try:
            progress_callback(data)
        except Exception:
            return

    def _build_paragraph_anchors(self, text: str, section_id: str, section_code: str) -> List[Dict[str, Any]]:
        normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
        raw_blocks = [item.strip() for item in re.split(r"\n{2,}", normalized) if item.strip()]
        if not raw_blocks and normalized.strip():
            raw_blocks = [normalized.strip()]
        anchors: List[Dict[str, Any]] = []
        cursor = 0
        for index, paragraph in enumerate(raw_blocks, start=1):
            start = normalized.find(paragraph, cursor)
            if start < 0:
                start = cursor
            end = start + len(paragraph)
            cursor = end
            anchors.append(
                {
                    "anchor_id": f"{section_id}:p{index}",
                    "anchor_label": f"{section_code or section_id}-P{index}",
                    "paragraph_index": index,
                    "text": paragraph,
                    "span_start": start,
                    "span_end": end,
                    "char_count": len(paragraph),
                }
            )
        return anchors

    def _bind_findings_to_paragraphs(
        self,
        findings: List[Dict[str, Any]],
        paragraph_blocks: List[Dict[str, Any]],
        section_id: str,
        section_code: str,
    ) -> List[Dict[str, Any]]:
        if not findings:
            return []
        anchors = paragraph_blocks or []
        bound: List[Dict[str, Any]] = []
        for finding in findings:
            metadata = finding.get("metadata", {}) if isinstance(finding.get("metadata", {}), dict) else {}
            metadata = dict(metadata)
            metadata["section_id"] = section_id
            metadata["section_code"] = section_code
            evidence_text = " ".join(
                [
                    str(finding.get("title", "") or ""),
                    str(finding.get("evidence", "") or ""),
                    str(finding.get("recommendation", "") or ""),
                ]
            ).strip()
            tokens = [
                token.strip()
                for token in re.split(r"[\s,锛屻€傦紱;銆?锛?)锛堬級]+", evidence_text)
                if token.strip() and len(token.strip()) >= 2
            ][:8]
            matched: List[Dict[str, Any]] = []
            for anchor in anchors:
                paragraph_text = str(anchor.get("text", "") or "")
                if not paragraph_text:
                    continue
                if tokens and any(token in paragraph_text for token in tokens):
                    matched.append(
                        {
                            "anchor_id": anchor.get("anchor_id", ""),
                            "anchor_label": anchor.get("anchor_label", ""),
                            "paragraph_index": anchor.get("paragraph_index"),
                            "span_start": anchor.get("span_start"),
                            "span_end": anchor.get("span_end"),
                            "snippet": self._preview(paragraph_text, 160),
                        }
                    )
                if len(matched) >= 3:
                    break
            metadata["anchors"] = matched
            bound.append({**finding, "metadata": metadata})
        return bound

    def _classify_risk_level(self, findings: List[Any], score: float) -> str:
        normalized = self._normalize_finding_dicts(findings)
        text = " ".join(
            [
                " ".join(
                    [
                        item.get("title", ""),
                        item.get("problem_type", ""),
                        item.get("evidence", ""),
                        item.get("recommendation", ""),
                    ]
                )
                for item in normalized
            ]
        ).lower()
        severities = {item.get("severity", "low") for item in normalized}
        high_markers = ["绂佸繉", "涓ラ噸", "鑷村懡", "black box", "contraindication", "fatal", "high risk", "critical"]
        medium_markers = ["鍓傞噺", "涓嶈壇鍙嶅簲", "鐩戞祴", "鐩镐簰浣滅敤", "warning", "adverse", "interaction", "monitor"]
        if "high" in severities:
            return "high"
        if "medium" in severities and score >= 0.2:
            return "medium"
        if any(k in text for k in high_markers):
            return "high"
        if score >= 0.65:
            return "high"
        if any(k in text for k in medium_markers):
            return "medium"
        if score >= 0.25:
            return "medium"
        return "low"

    def _apply_active_prompt_version(self, run_config: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        merged = dict(run_config or {})
        prompt_config = merged.get("prompt_config", {}) if isinstance(merged.get("prompt_config", {}), dict) else {}
        if prompt_config.get("prompt_bundle_path") and prompt_config.get("prompt_version_id"):
            merged["prompt_config"] = prompt_config
            return merged
        active_prompt_config = self.prompt_registry.resolve_active_prompt_config()
        if not active_prompt_config:
            return merged
        prompt_config = dict(prompt_config)
        prompt_config.update(active_prompt_config)
        merged["prompt_config"] = prompt_config
        return merged

    def _compose_runtime_prompt_config(
        self,
        project_id: str,
        task_type: str,
        base_prompt_config: Optional[Dict[str, Any]] = None,
        section_id: str = "",
        section_name: str = "",
        review_domain: str = "",
        product_type: str = "",
        registration_class: str = "",
    ) -> Dict[str, Any]:
        session = self.db_conn.get_session()
        try:
            return self.prompt_rule_service.compose_prompt_config(
                session=session,
                project_id=project_id,
                base_prompt_config=base_prompt_config or {},
                task_type=task_type,
                section_id=section_id,
                section_name=section_name,
                review_domain=review_domain,
                product_type=product_type,
                registration_class=registration_class,
            )
        finally:
            session.close()

    def _build_coordination_payload(
        self,
        project_id: str,
        run_id: str,
        source_doc_id: str,
        section_meta: Dict[str, Any],
        related_result: Dict[str, Any],
    ) -> Dict[str, Any]:
        hits = related_result.get("list", []) if isinstance(related_result, dict) else []
        grouped_docs = related_result.get("grouped_docs", []) if isinstance(related_result, dict) else []
        hit_digest = []
        for h in hits[:5]:
            if not isinstance(h, dict):
                continue
            hit_digest.append(
                {
                    "doc_id": h.get("doc_id", ""),
                    "chunk_id": h.get("chunk_id", ""),
                    "score": h.get("score", 0.0),
                    "classification": h.get("classification", ""),
                    "summary": self._preview(h.get("summary", "")),
                    "content_preview": self._preview(h.get("content", "")),
                }
            )
        grouped_digest = []
        for g in grouped_docs[:3]:
            if not isinstance(g, dict):
                continue
            grouped_digest.append(
                {
                    "doc_id": g.get("doc_id", ""),
                    "doc_summary": self._preview(g.get("doc_summary", "")),
                    "doc_keywords": g.get("doc_keywords", []),
                    "matched_count": len(g.get("matched_hits", []) or []),
                    "related_chunk_count": len(g.get("related_chunks", []) or []),
                }
            )
        return {
            "coordination_version": "v1",
            "project_id": project_id,
            "run_id": run_id,
            "source_doc_id": source_doc_id,
            "section_meta": section_meta,
            "retrieval": {
                "query": section_meta.get("query", ""),
                "hit_count": len(hits),
                "grouped_doc_count": len(grouped_docs),
                "hits": hit_digest,
                "grouped_docs": grouped_digest,
            },
            "memory_strategy": {
                "metadata_filter": {"project_id": project_id},
                "types": ["episodic", "semantic", "working"],
                "top_k": 8,
            },
        }

    def _seed_submission_structure_memory(
        self,
        project_id: str,
        source_doc_id: str,
        review_units: List[Dict[str, Any]],
    ) -> None:
        """
        Inject section skeleton into semantic memory to stabilize cross-section consistency.
        """
        for unit in review_units or []:
            if not isinstance(unit, dict):
                continue
            section_id = str(unit.get("section_id") or unit.get("chunk_id") or "").strip()
            if not section_id:
                continue
            section_code = str(unit.get("section_code", "")).strip()
            section_name = str(unit.get("section_name", "")).strip()
            title_path = unit.get("title_path", [])
            if not isinstance(title_path, list):
                title_path = []
            skeleton = {
                "section_id": section_id,
                "section_code": section_code,
                "section_name": section_name,
                "parent_code": str(unit.get("parent_code", "")).strip(),
                "unit_type": str(unit.get("unit_type", "")).strip(),
                "unit_order": unit.get("unit_order"),
                "title_path": [str(x).strip() for x in title_path if str(x).strip()],
                "page_start": unit.get("page_start"),
                "page_end": unit.get("page_end"),
            }
            self.memory_tool.remember(
                key=f"outline:{source_doc_id}:{section_id}",
                value=json.dumps(skeleton, ensure_ascii=False),
                memory_type="semantic",
                metadata={
                    "source": "submission_outline",
                    "project_id": project_id,
                    "doc_id": source_doc_id,
                    "section_id": section_id,
                    },
                )

    def _run_post_review_agents(
        self,
        project: PreReviewProject,
        run_id: str,
        source_doc_id: str,
        section_results: List[Dict[str, Any]],
        run_config: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        prompt_config = (run_config or {}).get("prompt_config", {}) if isinstance((run_config or {}).get("prompt_config", {}), dict) else {}
        return self.run_review_workflow.run(
            project_meta={
                "project_id": project.project_id,
                "project_name": project.project_name,
                "registration_scope": getattr(project, "registration_scope", "") or "",
                "source_doc_id": source_doc_id,
                "run_id": run_id,
            },
            section_results=section_results,
            prompt_config=prompt_config,
        )

    def _build_section_review_packet(
        self,
        project_id: str,
        run_id: str,
        source_doc_id: str,
        chunk: Dict[str, Any],
        section_meta: Dict[str, Any],
        related_result: Dict[str, Any],
        coordination_payload: Dict[str, Any],
        run_config: Optional[Dict[str, Any]] = None,
    ) -> SectionReviewPacket:
        related_rules = related_result.get("list", []) if isinstance(related_result, dict) else []
        title_path = chunk.get("title_path", []) if isinstance(chunk.get("title_path", []), list) else []
        return SectionReviewPacket(
            project_id=project_id,
            run_id=run_id,
            doc_id=source_doc_id,
            section_id=str(section_meta.get("section_id", "")),
            section_code=str(section_meta.get("section_code", "")),
            section_title=str(section_meta.get("section_name", "")),
            section_text=str(chunk.get("text", "") or ""),
            title_path=[str(x).strip() for x in title_path if str(x).strip()],
            page_start=chunk.get("page_start"),
            page_end=chunk.get("page_end"),
            unit_type=str(chunk.get("unit_type", "") or ""),
            retrieval_context=related_result if isinstance(related_result, dict) else {},
            related_rules=related_rules if isinstance(related_rules, list) else [],
            coordination_payload=coordination_payload,
            memory_metadata_filter={"project_id": project_id},
            run_config=dict(run_config or {}),
        )

    @staticmethod
    def _normalize_text_list(values: Any) -> List[str]:
        out: List[str] = []
        if not isinstance(values, list):
            return out
        seen = set()
        for item in values:
            text = str(item or "").strip()
            if not text or text in seen:
                continue
            seen.add(text)
            out.append(text)
        return out

    def _resolve_chunk_focus_points(
        self,
        project_id: str,
        section_id: str,
        chunk: Optional[Dict[str, Any]] = None,
        run_config: Optional[Dict[str, Any]] = None,
    ) -> List[str]:
        merged = self._dedupe_text_list(
            self._normalize_text_list((chunk or {}).get("concern_points", []))
            + self._normalize_text_list((chunk or {}).get("focus_points", []))
        )
        override_map = (run_config or {}).get("section_focus_overrides", {}) if isinstance(run_config or {}, dict) else {}
        if isinstance(override_map, dict):
            merged = self._dedupe_text_list(merged + self._normalize_text_list(override_map.get(section_id, [])))
        if merged:
            return merged
        if not project_id or not section_id:
            return merged
        session = self.db_conn.get_session()
        try:
            row = (
                session.query(PreReviewSectionConcern)
                .filter(
                    and_(
                        PreReviewSectionConcern.project_id == project_id,
                        PreReviewSectionConcern.section_id == section_id,
                    )
                )
                .first()
            )
            if row is None:
                return merged
            return self._dedupe_text_list(merged + self._parse_json_list(getattr(row, "concern_points", "") or ""))
        finally:
            session.close()

    def _split_submission_vector_unit(
        self,
        unit: Dict[str, Any],
        max_chars: int = 3200,
        overlap: int = 320,
    ) -> List[Dict[str, Any]]:
        text = str(unit.get("text", "")).strip()
        if not text:
            return []
        if len(text) <= max_chars:
            return [dict(unit)]
        overlap = max(0, min(overlap, max_chars // 2))
        chunk_rows: List[Dict[str, Any]] = []
        start = 0
        part_no = 1
        text_len = len(text)
        base_chunk_id = str(unit.get("chunk_id") or unit.get("section_id") or "submission_chunk").strip() or "submission_chunk"
        base_order = int(unit.get("unit_order") or 0)
        while start < text_len:
            end = min(text_len, start + max_chars)
            if end < text_len:
                split_at = max(
                    text.rfind("\n", start, end),
                    text.rfind("。", start, end),
                    text.rfind("；", start, end),
                    text.rfind(";", start, end),
                )
                if split_at > start + max_chars // 2:
                    end = split_at + 1
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk_unit = dict(unit)
                chunk_unit["chunk_id"] = f"{base_chunk_id}_part_{part_no}"
                chunk_unit["unit_type"] = str(unit.get("unit_type") or "submission_chunk")
                chunk_unit["unit_order"] = base_order * 100 + part_no if base_order else part_no
                chunk_unit["text"] = chunk_text
                chunk_unit["char_count"] = len(chunk_text)
                chunk_unit["source_chunk_id"] = base_chunk_id
                chunk_unit["chunk_part"] = part_no
                chunk_rows.append(chunk_unit)
                part_no += 1
            if end >= text_len:
                break
            start = max(0, end - overlap)
        return chunk_rows

    @staticmethod
    def _compact_text(text: Any, max_len: int = 120) -> str:
        value = re.sub(r"\s+", " ", str(text or "").strip())
        if not value:
            return ""
        if len(value) <= max_len:
            return value
        for sep in ["。", "；", ";", "，", ",", " ", "\n"]:
            idx = value.find(sep)
            if 0 < idx <= max_len:
                value = value[:idx]
                break
        return value[:max_len].strip(" ,;，；。")

    def _sanitize_retrieval_query(self, query: Any, max_len: int = 96) -> str:
        text = self._compact_text(query, max_len=max_len)
        if not text:
            return ""
        text = text.lstrip(" ._-:/\\")
        for pattern in [r"^\d+(\.\d+)+\s*", r"^第[一二三四五六七八九十百]+[章节部分]\s*"]:
            text = re.sub(pattern, "", text).strip()
        if re.fullmatch(r"[A-Za-z]?(?:\.\d+)+(?:\.\d+)?", text) or re.fullmatch(r"(?:\d+|[A-Za-z])(?:\.\d+)+", text):
            return ""
        return text


    @staticmethod
    def _infer_registration_class(project: PreReviewProject) -> str:
        for value in [
            getattr(project, "registration_leaf", ""),
            getattr(project, "registration_scope", ""),
            getattr(project, "registration_description", ""),
        ]:
            text = str(value or "").strip()
            if text:
                return PreReviewService._compact_text(text, max_len=64)
        return "药品"

    @staticmethod
    def _infer_review_domain(project: PreReviewProject) -> str:
        for value in [
            getattr(project, "registration_scope", ""),
            getattr(project, "registration_leaf", ""),
            getattr(project, "registration_description", ""),
        ]:
            text = str(value or "").strip()
            if not text:
                continue
            if "中药" in text:
                return "中药"
            if "化学" in text:
                return "化学药"
            if "生物" in text:
                return "生物制品"
        return "药学"

    @staticmethod
    def _resolve_pre_review_mode(run_config: Optional[Dict[str, Any]] = None) -> str:
        config = run_config if isinstance(run_config, dict) else {}
        strategy = str(config.get("strategy", "") or "").strip()
        workflow_mode = str(config.get("workflow_mode", "") or "").strip()
        if strategy == SINGLE_SECTION_PRE_REVIEW_V2 or workflow_mode == SINGLE_SECTION_PRE_REVIEW_V2:
            return SINGLE_SECTION_PRE_REVIEW_V2
        return "legacy"

    @staticmethod
    def _canonicalize_source_type(source_type: str) -> str:
        text = str(source_type or "").strip()
        if not text:
            return ""
        normalized = text.lower()
        for canonical, aliases in _RETRIEVAL_SOURCE_ALIASES.items():
            if text == canonical:
                return canonical
            if text in aliases or normalized in {str(x).lower() for x in aliases}:
                return canonical
        return text

    @staticmethod
    def _pharmacopeia_affect_range(review_domain: str) -> str:
        return "中药" if "中药" in str(review_domain or "") else "化学药"

    @staticmethod
    def _infer_product_type(project: Optional[PreReviewProject], section_name: str = "") -> str:
        for value in [
            getattr(project, "registration_leaf", "") if project is not None else "",
            getattr(project, "registration_description", "") if project is not None else "",
            section_name,
        ]:
            text = str(value or "").strip()
            if not text:
                continue
            if "注射" in text:
                return "注射剂"
            if "片" in text:
                return "片剂"
            if "胶囊" in text:
                return "胶囊剂"
            if "原料药" in text:
                return "原料药"
        return "化学药"

    def _load_historical_experience(
        self,
        session,
        section_id: str,
        product_type: str,
        limit: int = 12,
    ) -> List[Dict[str, Any]]:
        query = session.query(PreReviewExperienceMemory).filter(PreReviewExperienceMemory.status == "active")
        rows = query.order_by(PreReviewExperienceMemory.id.desc()).all()
        out: List[Dict[str, Any]] = []
        for row in rows:
            scope_type = str(getattr(row, "scope_type", "") or "").strip()
            scope_key = str(getattr(row, "scope_key", "") or "").strip()
            if scope_type == "section" and scope_key != section_id:
                continue
            if scope_type == "product_type" and scope_key != product_type:
                continue
            payload = {}
            try:
                payload = json.loads(getattr(row, "payload_json", "") or "{}")
            except Exception:
                payload = {}
            out.append(
                {
                    "experience_id": getattr(row, "experience_id", ""),
                    "experience_type": getattr(row, "experience_type", ""),
                    "content": getattr(row, "content", ""),
                    "applicable_sections": [section_id] if scope_type == "section" else [],
                    "scope_type": scope_type,
                    "scope_key": scope_key,
                    "trigger_conditions": self._parse_json_list(getattr(row, "trigger_conditions", "") or ""),
                    "payload": payload if isinstance(payload, dict) else {},
                }
            )
            if len(out) >= limit:
                break
        return out

    @staticmethod
    def _classification_from_source_type(source_type: str) -> str:
        canonical = PreReviewService._canonicalize_source_type(source_type)
        mapping = {
            RETRIEVAL_SOURCE_GUIDANCE: RETRIEVAL_SOURCE_GUIDANCE,
            RETRIEVAL_SOURCE_ICH: RETRIEVAL_SOURCE_ICH,
            RETRIEVAL_SOURCE_REGULATION: RETRIEVAL_SOURCE_REGULATION,
        }
        return mapping.get(canonical, "")

    @staticmethod
    def _retrieval_source_quota(max_materials: int) -> Dict[str, int]:
        total = max(5, int(max_materials or 12))
        quotas = {
            RETRIEVAL_SOURCE_GUIDANCE: 3,
            RETRIEVAL_SOURCE_ICH: 2,
            RETRIEVAL_SOURCE_REGULATION: 2,
            RETRIEVAL_SOURCE_PHARMACOPOEIA: 2,
            RETRIEVAL_SOURCE_EXPERIENCE: 1,
        }
        base_total = sum(quotas.values())
        if total <= base_total:
            scale = total / base_total
            scaled = {key: max(1, int(round(value * scale))) for key, value in quotas.items()}
            ordered_keys = [
                RETRIEVAL_SOURCE_GUIDANCE,
                RETRIEVAL_SOURCE_ICH,
                RETRIEVAL_SOURCE_REGULATION,
                RETRIEVAL_SOURCE_PHARMACOPOEIA,
                RETRIEVAL_SOURCE_EXPERIENCE,
            ]
            while sum(scaled.values()) > total:
                for key in ordered_keys:
                    if sum(scaled.values()) <= total:
                        break
                    if scaled[key] > 1:
                        scaled[key] -= 1
            return scaled
        quotas[RETRIEVAL_SOURCE_GUIDANCE] += total - base_total
        return quotas

    @staticmethod
    def _score_retrieval_material(item: Dict[str, Any], query_terms: List[str]) -> float:
        source_type = str(item.get("source_type", "") or "").strip()
        title = str(item.get("title", "") or "").strip()
        content = str(item.get("content", "") or "").strip()
        metadata = item.get("metadata", {}) if isinstance(item.get("metadata", {}), dict) else {}
        text = f"{title}\n{content}\n{metadata.get('section_path_text', '')}".strip().lower()
        score = 0.0
        for term in query_terms:
            if term and term.lower() in text:
                score += 1.0
        score += {
            RETRIEVAL_SOURCE_GUIDANCE: 0.50,
            RETRIEVAL_SOURCE_ICH: 0.40,
            RETRIEVAL_SOURCE_REGULATION: 0.35,
            RETRIEVAL_SOURCE_PHARMACOPOEIA: 0.30,
            RETRIEVAL_SOURCE_EXPERIENCE: 0.20,
        }.get(source_type, 0.0)
        if title:
            score += min(len(title) / 200.0, 0.2)
        return score

    def _finalize_retrieval_materials(
        self,
        materials: List[Dict[str, Any]],
        query_list: List[str],
        max_materials: int = 12,
    ) -> List[Dict[str, Any]]:
        if not materials:
            return []
        quotas = self._retrieval_source_quota(max_materials=max_materials)
        query_terms: List[str] = []
        for query in query_list:
            for term in re.findall(r"[A-Za-z0-9_./\\-]+|[\u4e00-\u9fff]{2,}", str(query or "")):
                text = str(term or "").strip()
                if text and text not in query_terms:
                    query_terms.append(text)

        grouped: Dict[str, List[Dict[str, Any]]] = {}
        for item in materials:
            if not isinstance(item, dict):
                continue
            normalized = dict(item)
            normalized["source_type"] = self._canonicalize_source_type(item.get("source_type", ""))
            normalized["_score"] = self._score_retrieval_material(normalized, query_terms)
            grouped.setdefault(str(normalized.get("source_type", "") or ""), []).append(normalized)

        ordered_sources = [
            RETRIEVAL_SOURCE_GUIDANCE,
            RETRIEVAL_SOURCE_ICH,
            RETRIEVAL_SOURCE_REGULATION,
            RETRIEVAL_SOURCE_PHARMACOPOEIA,
            RETRIEVAL_SOURCE_EXPERIENCE,
        ]
        selected: List[Dict[str, Any]] = []
        selected_ids = set()
        for source_type in ordered_sources:
            bucket = sorted(
                grouped.get(source_type, []),
                key=lambda x: (
                    -float(x.get("_score", 0.0) or 0.0),
                    str(x.get("title", "") or ""),
                    str(x.get("evidence_id", "") or ""),
                ),
            )
            for item in bucket[: quotas.get(source_type, 0)]:
                evidence_id = str(item.get("evidence_id", "") or "").strip()
                if not evidence_id or evidence_id in selected_ids:
                    continue
                selected_ids.add(evidence_id)
                selected.append(item)

        if len(selected) < max_materials:
            leftovers: List[Dict[str, Any]] = []
            for bucket in grouped.values():
                leftovers.extend(bucket)
            leftovers = sorted(
                leftovers,
                key=lambda x: (
                    -float(x.get("_score", 0.0) or 0.0),
                    str(x.get("title", "") or ""),
                    str(x.get("evidence_id", "") or ""),
                ),
            )
            for item in leftovers:
                evidence_id = str(item.get("evidence_id", "") or "").strip()
                if not evidence_id or evidence_id in selected_ids:
                    continue
                selected_ids.add(evidence_id)
                selected.append(item)
                if len(selected) >= max_materials:
                    break

        finalized: List[Dict[str, Any]] = []
        for item in selected[:max_materials]:
            row = dict(item)
            row.pop("_score", None)
            finalized.append(row)
        return finalized

    def _execute_retrieval_plan(
        self,
        planner_result: Dict[str, Any],
        historical_experience: List[Dict[str, Any]],
        review_domain: str = "",
        focus_points: Optional[List[Any]] = None,
        section_id: str = "",
        section_name: str = "",
        progress_callback=None,
        top_k_per_query: int = 3,
        max_materials: int = 12,
    ) -> List[Dict[str, Any]]:
        def _source_titles(items: List[Dict[str, Any]], current_source: str, limit: int = 5) -> List[str]:
            titles: List[str] = []
            for item in items:
                if str(item.get("source_type", "") or "") != current_source:
                    continue
                title = str(item.get("title", "") or "").strip()
                if not title or title in titles:
                    continue
                titles.append(title)
                if len(titles) >= limit:
                    break
            return titles

        materials: List[Dict[str, Any]] = []
        seen = set()
        retrieval_plan = planner_result.get("retrieval_plan", []) if isinstance(planner_result.get("retrieval_plan", []), list) else []
        normalized_focus_points = self._normalize_text_list(focus_points or [])
        query_list = self._merge_retrieval_queries_with_focus_points(
            planner_result.get("query_list", []),
            normalized_focus_points,
        )
        if not retrieval_plan:
            retrieval_plan = [
                {"source_type": RETRIEVAL_SOURCE_GUIDANCE, "purpose": "基础规范检索", "query_subset": query_list[:2]},
                {"source_type": RETRIEVAL_SOURCE_ICH, "purpose": "ICH 技术要求检索", "query_subset": query_list[:1]},
                {"source_type": RETRIEVAL_SOURCE_REGULATION, "purpose": "法律法规依据检索", "query_subset": query_list[:1]},
                {"source_type": RETRIEVAL_SOURCE_PHARMACOPOEIA, "purpose": "药典结构化数据检索", "query_subset": query_list[:1]},
                {"source_type": RETRIEVAL_SOURCE_EXPERIENCE, "purpose": "历史经验提醒", "query_subset": query_list[:1]},
            ]
        for step in retrieval_plan:
            if not isinstance(step, dict):
                continue
            source_type = self._canonicalize_source_type(step.get("source_type", ""))
            self._emit_progress(
                progress_callback,
                "retrieval_source_start",
                f"开始检索{source_type or '未知来源'}",
                section_id=section_id,
                section_name=section_name,
                source_type=source_type,
                purpose=str(step.get("purpose", "") or "").strip(),
            )
            if source_type == RETRIEVAL_SOURCE_EXPERIENCE:
                for item in historical_experience[:3]:
                    evidence_id = f"exp:{item.get('experience_id', uuid.uuid4().hex[:8])}"
                    if evidence_id in seen:
                        continue
                    seen.add(evidence_id)
                    materials.append(
                        {
                            "source_type": RETRIEVAL_SOURCE_EXPERIENCE,
                            "title": str(item.get("experience_type", "") or RETRIEVAL_SOURCE_EXPERIENCE),
                            "content": str(item.get("content", "") or ""),
                            "evidence_id": evidence_id,
                            "metadata": {
                                "experience_id": str(item.get("experience_id", "") or ""),
                                "scope_type": str(item.get("scope_type", "") or ""),
                                "scope_key": str(item.get("scope_key", "") or ""),
                            },
                        }
                    )
                self._emit_progress(
                    progress_callback,
                    "retrieval_source_done",
                    f"{source_type or '历史经验'} 检索完成",
                    section_id=section_id,
                    section_name=section_name,
                    source_type=source_type,
                    hit_count=len([x for x in materials if str(x.get("source_type", "") or "") == RETRIEVAL_SOURCE_EXPERIENCE]),
                    hit_titles=_source_titles(materials, RETRIEVAL_SOURCE_EXPERIENCE),
                )
                continue
            step_queries = self._merge_retrieval_queries_with_focus_points(
                step.get("query_subset", []) or query_list[:2],
                normalized_focus_points,
                max_queries=4,
            ) or query_list[:2]
            sanitized_step_queries: List[str] = []
            seen_queries = set()
            for raw_query in step_queries:
                cleaned = self._sanitize_retrieval_query(raw_query, max_len=120)
                if not cleaned or cleaned in seen_queries:
                    continue
                seen_queries.add(cleaned)
                sanitized_step_queries.append(cleaned)
            if not sanitized_step_queries:
                fallback_query = self._sanitize_retrieval_query(section_name, max_len=64)
                if fallback_query:
                    sanitized_step_queries = [fallback_query]
            if not sanitized_step_queries:
                self._emit_progress(
                    progress_callback,
                    "retrieval_source_done",
                    f"{source_type or '未知来源'} 检索跳过：无有效 query",
                    section_id=section_id,
                    section_name=section_name,
                    source_type=source_type,
                    hit_count=len([x for x in materials if str(x.get("source_type", "") or "") == source_type]),
                    hit_titles=_source_titles(materials, source_type),
                )
                continue
            step_queries = sanitized_step_queries
            if source_type == RETRIEVAL_SOURCE_PHARMACOPOEIA:
                affect_range = self._pharmacopeia_affect_range(review_domain)
                for query in step_queries[:3]:
                    self._emit_progress(
                        progress_callback,
                        "retrieval_query",
                        f"正在检索{source_type}：{query}",
                        section_id=section_id,
                        section_name=section_name,
                        source_type=source_type,
                        query=query,
                    )
                    result_rows = self.pharmacopeia.search_entries(query=query, affect_range=affect_range, top_k=top_k_per_query)
                    for hit in result_rows:
                        evidence_id = f"pharm:{hit.get('entry_id', '')}"
                        if not str(evidence_id).strip() or evidence_id in seen:
                            continue
                        seen.add(evidence_id)
                        materials.append(
                            {
                                "source_type": RETRIEVAL_SOURCE_PHARMACOPOEIA,
                                "title": str(hit.get("drug_name", "") or "鑽吀鏉＄洰").strip(),
                                "content": str(hit.get("retrieval_text", "") or "").strip(),
                                "evidence_id": evidence_id,
                                "metadata": {
                                    "entry_id": str(hit.get("entry_id", "") or ""),
                                    "drug_name": str(hit.get("drug_name", "") or ""),
                                    "affect_range": str(hit.get("affect_range", "") or ""),
                                    "source_file_name": str(hit.get("source_file_name", "") or ""),
                                },
                            }
                        )
                self._emit_progress(
                    progress_callback,
                    "retrieval_source_done",
                    f"{source_type} 检索完成",
                    section_id=section_id,
                    section_name=section_name,
                    source_type=source_type,
                    hit_count=len([x for x in materials if str(x.get("source_type", "") or "") == RETRIEVAL_SOURCE_PHARMACOPOEIA]),
                    hit_titles=_source_titles(materials, RETRIEVAL_SOURCE_PHARMACOPOEIA),
                )
                continue
            classification = self._classification_from_source_type(source_type)
            if not classification:
                continue
            for query in step_queries[:3]:
                self._emit_progress(
                    progress_callback,
                    "retrieval_query",
                    f"正在检索{source_type}：{query}",
                    section_id=section_id,
                    section_name=section_name,
                    source_type=source_type,
                    query=query,
                )
                result = self.knowledge.semantic_query(
                    query=query,
                    top_k=top_k_per_query,
                    classification=classification,
                    min_score=0.0,
                )
                for hit in result.get("list", []) if isinstance(result, dict) else []:
                    evidence_id = f"{hit.get('doc_id', '')}:{hit.get('chunk_id', '')}"
                    if not str(evidence_id).strip() or evidence_id in seen:
                        continue
                    seen.add(evidence_id)
                    materials.append(
                        {
                            "source_type": source_type or self._canonicalize_source_type(hit.get("classification", "")) or RETRIEVAL_SOURCE_GUIDANCE,
                            "title": str(hit.get("section_name", "") or hit.get("doc_id", "") or "").strip(),
                            "content": str(hit.get("content", "") or "").strip(),
                            "evidence_id": evidence_id,
                            "metadata": {
                                "doc_id": str(hit.get("doc_id", "") or ""),
                                "chunk_id": str(hit.get("chunk_id", "") or ""),
                                "classification": str(hit.get("classification", "") or ""),
                                "section_path_text": str(hit.get("section_path_text", "") or ""),
                            },
                        }
                    )
            self._emit_progress(
                progress_callback,
                "retrieval_source_done",
                f"{source_type} 检索完成",
                section_id=section_id,
                section_name=section_name,
                source_type=source_type,
                hit_count=len([x for x in materials if str(x.get("source_type", "") or "") == source_type]),
                hit_titles=_source_titles(materials, source_type),
            )
        return self._finalize_retrieval_materials(
            materials=materials,
            query_list=query_list,
            max_materials=max_materials,
        )

    @staticmethod
    def _map_review_result_to_legacy(review_result: Dict[str, Any]) -> Tuple[str, List[Dict[str, Any]], List[str], str]:
        section_summary = str(review_result.get("section_summary", "") or "").strip()
        supported_points = review_result.get("supported_points", []) if isinstance(review_result.get("supported_points", []), list) else []
        missing_points = review_result.get("missing_points", []) if isinstance(review_result.get("missing_points", []), list) else []
        unsupported_points = review_result.get("unsupported_points", []) if isinstance(review_result.get("unsupported_points", []), list) else []
        questions = review_result.get("questions", []) if isinstance(review_result.get("questions", []), list) else []
        highlighted = []
        for item in questions:
            if not isinstance(item, dict):
                continue
            highlighted.append(
                {
                    "title": str(item.get("issue", "") or "").strip(),
                    "problem_type": "chapter_question",
                    "severity": "medium" if unsupported_points else "low",
                    "evidence": str(item.get("basis", "") or "").strip(),
                    "recommendation": str(item.get("requested_action", "") or "").strip(),
                    "metadata": {},
                }
            )
        if not highlighted:
            for item in unsupported_points[:3] + missing_points[:3]:
                text = str(item or "").strip()
                if text:
                    highlighted.append(
                        {
                            "title": text,
                            "problem_type": "missing_or_unsupported",
                            "severity": "medium",
                            "evidence": "",
                            "recommendation": "",
                            "metadata": {},
                        }
                    )
        linked_rules = [str(x).strip() for x in review_result.get("evidence_refs", []) if str(x).strip()]
        risk_points = review_result.get("risk_points", []) if isinstance(review_result.get("risk_points", []), list) else []
        risk_level = "low"
        if unsupported_points or len(risk_points) >= 3:
            risk_level = "high"
        elif missing_points or risk_points:
            risk_level = "medium"
        conclusion = section_summary or "；".join([str(x).strip() for x in supported_points[:2] + missing_points[:2] if str(x).strip()])
        return conclusion[:1000], highlighted, linked_rules, risk_level

    def _review_single_chunk(
        self,
        project: PreReviewProject,
        project_id: str,
        run_id: str,
        source_doc_id: str,
        chunk: Dict[str, Any],
        previous_section_meta: Optional[Dict[str, Any]] = None,
        run_config: Optional[Dict[str, Any]] = None,
        progress_callback=None,
    ) -> Dict[str, Any]:
        """Review one parsed section chunk and return structured review artifacts."""
        previous_section_meta = previous_section_meta or {}
        workflow_mode = self._resolve_pre_review_mode(run_config)
        section_id = str(chunk.get("section_id") or chunk.get("chunk_id", ""))
        section_code = str(chunk.get("section_code") or section_id)
        section_name = str(chunk.get("section_name") or chunk.get("title") or section_code or f"section-{section_id}")
        text = str(chunk.get("text", "")).strip()
        if not text:
            return {"success": False, "message": "empty section text", "section_id": section_id}
        self._emit_progress(
            progress_callback,
            "section_start",
            f"开始处理章节 {section_code} {section_name}".strip(),
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
            unit_order=chunk.get("unit_order"),
        )
        paragraph_blocks = self._build_paragraph_anchors(text=text, section_id=section_id, section_code=section_code)
        base_prompt_config = (run_config or {}).get("prompt_config", {}) if isinstance((run_config or {}).get("prompt_config", {}), dict) else {}
        registration_class = self._infer_registration_class(project)
        review_domain = self._infer_review_domain(project)
        product_type = self._infer_product_type(project, section_name=section_name)
        focus_points = self._resolve_chunk_focus_points(
            project_id=project_id,
            section_id=section_id,
            chunk=chunk if isinstance(chunk, dict) else {},
            run_config=run_config if isinstance(run_config, dict) else {},
        )
        exp_session = self.db_conn.get_session()
        try:
            historical_experience = self._load_historical_experience(
                session=exp_session,
                section_id=section_id,
                product_type=product_type,
            )
        finally:
            exp_session.close()
        planner_input = {
            "task_id": f"{run_id}:{section_id}",
            "application_id": project_id,
            "section_id": section_id,
            "section_name": section_name,
            "registration_class": registration_class,
            "review_domain": review_domain,
            "product_type": product_type,
            "raw_text": self._compact_text(f"{section_name} {text}".strip(), max_len=240),
            "focus_points": focus_points,
            "historical_experience": historical_experience,
            "metadata": {
                "source_file_ids": [str(item.get("doc_id", "") or "").strip() for item in chunk.get("attached_files", []) if isinstance(item, dict)],
                "language": "zh-CN",
            },
        }
        planner_prompt_config = self._compose_runtime_prompt_config(
            project_id=project_id,
            task_type="planner",
            base_prompt_config=base_prompt_config,
            section_id=section_id,
            section_name=section_name,
            review_domain=review_domain,
            product_type=product_type,
            registration_class=registration_class,
        )
        self._emit_progress(
            progress_callback,
            "planner_start",
            f"开始规划章节 {section_code} 的检索计划".strip(),
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
            focus_points=focus_points,
        )
        planner_result = self.planner_reviewer_agent.plan(planner_input, prompt_config=planner_prompt_config)
        self._emit_progress(
            progress_callback,
            "planner_done",
            f"章节 {section_code} 检索计划已生成".strip(),
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
            query_list=self._normalize_text_list(planner_result.get("query_list", [])),
            retrieval_plan=planner_result.get("retrieval_plan", []),
        )
        self._emit_progress(
            progress_callback,
            "retrieval_start",
            f"开始检索章节 {section_code} 相关资料".strip(),
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
        )
        retrieved_materials = self._execute_retrieval_plan(
            planner_result,
            historical_experience,
            review_domain=review_domain,
            focus_points=focus_points,
            section_id=section_id,
            section_name=section_name,
            progress_callback=progress_callback,
        )
        self._emit_progress(
            progress_callback,
            "retrieval_done",
            f"章节 {section_code} 检索完成".strip(),
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
            hit_count=len(retrieved_materials),
            source_types=self._normalize_text_list([item.get("source_type", "") for item in retrieved_materials if isinstance(item, dict)]),
        )
        review_input = dict(planner_input)
        review_input["raw_text"] = text
        review_input["retrieved_materials"] = retrieved_materials
        reviewer_prompt_config = self._compose_runtime_prompt_config(
            project_id=project_id,
            task_type="reviewer",
            base_prompt_config=base_prompt_config,
            section_id=section_id,
            section_name=section_name,
            review_domain=review_domain,
            product_type=product_type,
            registration_class=registration_class,
        )
        self._emit_progress(
            progress_callback,
            "reviewer_start",
            f"开始生成章节 {section_code} 的预审结论".strip(),
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
        )
        review_result = self.planner_reviewer_agent.review(review_input, prompt_config=reviewer_prompt_config)
        self._emit_progress(
            progress_callback,
            "reviewer_done",
            f"章节 {section_code} 预审完成".strip(),
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
            conclusion=str(review_result.get("pre_review_conclusion", "") or ""),
        )
        section_summary = {
            "structured_summary": str(review_result.get("section_summary", "") or "").strip(),
            "key_facts": self._normalize_text_list(review_result.get("fact_basis", {}).get("explicit_in_text", []) if isinstance(review_result.get("fact_basis", {}), dict) else []),
            "missing_items": self._normalize_text_list(review_result.get("missing_points", [])),
            "draft_risks": self._normalize_text_list(review_result.get("risk_points", [])),
            "source_files": [str(item.get("file_name", "") or "").strip() for item in chunk.get("attached_files", []) if isinstance(item, dict)],
        }
        section_meta = {
            "section_id": section_id,
            "section_code": section_code,
            "section_name": section_name,
            "query": " | ".join(self._normalize_text_list(planner_result.get("query_list", []))[:3]),
            "text_preview": self._preview(text, max_len=220),
            "char_count": len(text),
            "page": chunk.get("page"),
            "page_start": chunk.get("page_start"),
            "page_end": chunk.get("page_end"),
            "unit_order": chunk.get("unit_order"),
            "unit_type": chunk.get("unit_type"),
            "parent_section_id": chunk.get("parent_section_id"),
            "parent_code": chunk.get("parent_code"),
            "title_path": chunk.get("title_path", []),
            "concern_points": focus_points,
            "attached_files": chunk.get("attached_files", []),
            "paragraph_blocks": paragraph_blocks,
            "section_summary": section_summary,
            "previous_section_id": previous_section_meta.get("section_id", ""),
            "previous_conclusion_preview": previous_section_meta.get("conclusion_preview", ""),
        }
        coordination_payload = {
            "planner_result": planner_result,
            "retrieval": {
                "query_count": len(self._normalize_text_list(planner_result.get("query_list", []))),
                "focus_points": focus_points,
                "focus_point_count": len(focus_points),
                "effective_queries": self._merge_retrieval_queries_with_focus_points(
                    planner_result.get("query_list", []),
                    focus_points,
                ),
                "hit_count": len(retrieved_materials),
                "grouped_doc_count": len({str(x.get('evidence_id', '')).split(':')[0] for x in retrieved_materials if str(x.get('evidence_id', '')).strip()}),
                "retrieved_materials": retrieved_materials,
                "sources": self._normalize_text_list([item.get("source_type", "") for item in retrieved_materials if isinstance(item, dict)]),
            },
        }
        section_review_packet = self._build_section_review_packet(
            project_id=project_id,
            run_id=run_id,
            source_doc_id=source_doc_id,
            chunk=chunk,
            section_meta=section_meta,
            related_result={"list": retrieved_materials},
            coordination_payload=coordination_payload,
            run_config=run_config,
        )
        conclusion, findings, linked_rules, risk_level = self._map_review_result_to_legacy(review_result)
        findings = self._bind_findings_to_paragraphs(
            findings=findings,
            paragraph_blocks=paragraph_blocks,
            section_id=section_id,
            section_code=section_code,
        )
        score_map = {
            "supported": 0.9,
            "partially_supported": 0.65,
            "unsupported": 0.25,
            "insufficient_information": 0.4,
        }
        score = float(score_map.get(str(review_result.get("pre_review_conclusion", "") or ""), 0.5))

        self.memory_tool.remember(
            key=f"recent:{source_doc_id}:{section_id}",
            value=conclusion,
            memory_type="working",
            metadata={
                "source": "recent_finding",
                "project_id": project_id,
                "doc_id": source_doc_id,
                "section_id": section_id,
            },
        )

        conclusion_row = SectionConclusionRecord(
            run_id=run_id,
            section_id=section_id,
            section_name=f"{section_code} {section_name}".strip(),
            conclusion=conclusion,
            highlighted_issues=findings,
            linked_rules=linked_rules,
            risk_level=risk_level,
            create_time=self._now(),
        ).to_entity()
        trace_payload = {
            "trace_schema": "chapter_review_v1",
            "section_packet": {
                "section_id": section_id,
                "section_code": section_code,
                "section_title": section_name,
                "page_start": chunk.get("page_start"),
                "page_end": chunk.get("page_end"),
                "title_path": chunk.get("title_path", []),
                "paragraph_blocks": paragraph_blocks,
                "raw_text": text,
                "focus_points": self._normalize_text_list(focus_points),
            },
            "section_review_packet": section_review_packet.to_dict(),
            "coordination": coordination_payload,
            "memory": {
                "hit_count": len(historical_experience),
                "context_preview": self._preview(json.dumps(historical_experience[:3], ensure_ascii=False), max_len=600),
                "hits_by_type": {"鍘嗗彶缁忛獙": len(historical_experience)},
            },
            "summary_agent": section_summary,
            "agent": {
                "strategy": "planner_review_single_v2" if workflow_mode == SINGLE_SECTION_PRE_REVIEW_V2 else "planner_review_v1",
                "roles": ["planner", "pre_review"],
                "findings": findings,
                "findings_count": len(findings),
                "score": score,
            },
            "planner_result": planner_result,
            "retrieved_materials": retrieved_materials,
            "standardized_output": review_result,
            "prompt_rules": {
                "planner": ((planner_prompt_config.get("prompt_bundle", {}) if isinstance(planner_prompt_config.get("prompt_bundle", {}), dict) else {}).get("active_rules", {}) or {}).get("planner", []),
                "reviewer": ((reviewer_prompt_config.get("prompt_bundle", {}) if isinstance(reviewer_prompt_config.get("prompt_bundle", {}), dict) else {}).get("active_rules", {}) or {}).get("reviewer", []),
            },
            "trace": {
                "registration_class": registration_class,
                "review_domain": review_domain,
                "product_type": product_type,
            },
        }
        trace_row = SectionTraceRecord(
            run_id=run_id,
            section_id=section_id,
            trace_payload=trace_payload,
            create_time=self._now(),
        ).to_entity()
        return {
            "success": True,
            "section_id": section_id,
            "section_meta": section_meta,
            "findings": findings,
            "score": score,
            "linked_rules": linked_rules,
            "risk_level": risk_level,
            "section_summary": section_summary,
            "planner_result": planner_result,
            "review_result": review_result,
            "conclusion_row": conclusion_row,
            "trace_row": trace_row,
            "trace_payload": trace_payload,
            "previous_section_meta": {
                "section_id": section_id,
                "conclusion_preview": self._preview(conclusion),
            },
        }

    def create_project(
        self,
        project_name: str,
        description: str = "",
        owner: str = "",
        registration_scope: str = "",
        registration_path: Optional[List[str]] = None,
        registration_leaf: str = "",
        registration_description: str = "",
    ) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.create_project | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        if not project_name:
            return False, "project_name is required", None
        session = self.db_conn.get_session()
        try:
            entity = PreReviewProject(
                project_id=self._new_project_id(),
                project_name=project_name,
                description=description,
                registration_scope=str(registration_scope or "").strip(),
                registration_path=json.dumps(registration_path or [], ensure_ascii=False),
                registration_leaf=str(registration_leaf or "").strip(),
                registration_description=str(registration_description or "").strip(),
                status="created",
                progress=0.0,
                owner=owner,
                create_time=self._now(),
                update_time=self._now(),
                is_deleted=False,
            )
            session.add(entity)
            session.flush()
            self.prompt_rule_service.ensure_project_rules(session, entity.project_id)
            session.commit()
            return True, "project created", {
                "project_id": entity.project_id,
                "project_name": entity.project_name,
                "registration_scope": entity.registration_scope,
                "registration_path": registration_path or [],
                "registration_leaf": entity.registration_leaf,
                "registration_description": entity.registration_description,
            }
        except Exception as e:
            session.rollback()
            return False, f"create project failed: {str(e)}", None
        finally:
            session.close()

    def delete_project(self, project_id: str) -> Tuple[bool, str]:
        print("[DEBUG] enter PreReviewService.delete_project | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            row = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if row is None:
                return False, "project not found"
            self._purge_project_submission_files(session=session, project_id=project_id)
            self._purge_project_runtime_records(session=session, project_id=project_id)
            row.is_deleted = True
            row.status = "archived"
            row.update_time = self._now()
            session.commit()
            return True, "project deleted"
        except Exception as e:
            session.rollback()
            return False, f"delete project failed: {str(e)}"
        finally:
            session.close()

    def _purge_project_submission_files(self, session, project_id: str) -> None:
        rows = (
            session.query(PreReviewSubmissionFile)
            .filter(PreReviewSubmissionFile.project_id == project_id)
            .all()
        )
        for row in rows:
            doc_id = str(getattr(row, "doc_id", "") or "").strip()
            file_path = str(getattr(row, "file_path", "") or "").strip()
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception:
                    pass
            parsed_path = os.path.join(SUBMISSION_PARSED_DIR, f"{doc_id}.json")
            if doc_id and os.path.exists(parsed_path):
                try:
                    os.remove(parsed_path)
                except Exception:
                    pass
            edit_path = self._submission_edit_path(doc_id)
            if doc_id and os.path.exists(edit_path):
                try:
                    os.remove(edit_path)
                except Exception:
                    pass
            if doc_id:
                try:
                    self.submission_vector_store.delete_by_doc(doc_id)
                except Exception:
                    pass
                self._submission_payload_cache.pop(self._submission_payload_cache_key(project_id, doc_id), None)
            row.is_deleted = True
            row.is_chunked = False
            row.chunk_ids = ""
            row.chunk_size = 0

        session.query(PreReviewSubmissionSectionContent).filter(
            PreReviewSubmissionSectionContent.project_id == project_id
        ).delete(synchronize_session=False)

        shadow_rows = (
            session.query(FileInfo)
            .filter(
                and_(
                    FileInfo.classification == "submission_material",
                    FileInfo.affect_range == "pre_review",
                    FileInfo.doc_id.in_([str(getattr(item, "doc_id", "") or "") for item in rows if str(getattr(item, "doc_id", "") or "").strip()]),
                )
            )
            .all()
        )
        for shadow in shadow_rows:
            shadow.is_deleted = True
            shadow.is_chunked = False
            shadow.chunk_ids = ""
            shadow.chunk_size = 0

    def _purge_project_runtime_records(self, session, project_id: str) -> None:
        run_ids = [
            str(item.run_id or "").strip()
            for item in session.query(PreReviewRun.run_id).filter(PreReviewRun.project_id == project_id).all()
            if str(item.run_id or "").strip()
        ]
        if run_ids:
            session.query(PreReviewSectionConclusion).filter(
                PreReviewSectionConclusion.run_id.in_(run_ids)
            ).delete(synchronize_session=False)
            session.query(PreReviewSectionTrace).filter(
                PreReviewSectionTrace.run_id.in_(run_ids)
            ).delete(synchronize_session=False)
            session.query(PreReviewFeedback).filter(
                PreReviewFeedback.run_id.in_(run_ids)
            ).delete(synchronize_session=False)
            session.query(PreReviewSectionOutput).filter(
                PreReviewSectionOutput.run_id.in_(run_ids)
            ).delete(synchronize_session=False)
            session.query(PreReviewFeedbackAnalysisResult).filter(
                PreReviewFeedbackAnalysisResult.run_id.in_(run_ids)
            ).delete(synchronize_session=False)
            session.query(PreReviewPatchRegistry).filter(
                PreReviewPatchRegistry.run_id.in_(run_ids)
            ).delete(synchronize_session=False)
            session.query(PreReviewRun).filter(
                PreReviewRun.run_id.in_(run_ids)
            ).delete(synchronize_session=False)

        session.query(PreReviewSectionConcern).filter(
            PreReviewSectionConcern.project_id == project_id
        ).delete(synchronize_session=False)
        session.query(PreReviewPromptRule).filter(
            PreReviewPromptRule.project_id == project_id
        ).delete(synchronize_session=False)
        session.query(PreReviewExperienceMemory).filter(
            PreReviewExperienceMemory.scope_key == project_id
        ).delete(synchronize_session=False)

    def list_projects(self, page: int = 1, page_size: int = 10, status: str = "") -> Dict[str, Any]:
        print("[DEBUG] enter PreReviewService.list_projects | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            query = session.query(PreReviewProject).filter(PreReviewProject.is_deleted == 0)
            if status:
                query = query.filter(PreReviewProject.status == status)
            total = query.count()
            page = max(1, int(page))
            page_size = max(1, int(page_size))
            rows = (
                query.order_by(desc(PreReviewProject.id))
                .offset((page - 1) * page_size)
                .limit(page_size)
                .all()
            )
            return {
                "list": [
                    {
                        **{
                            "project_id": r.project_id,
                            "project_name": r.project_name,
                            "description": r.description,
                            "registration_scope": r.registration_scope,
                            "registration_path": self._parse_json_list(r.registration_path),
                            "registration_leaf": r.registration_leaf,
                            "registration_description": r.registration_description,
                            "status": r.status,
                            "progress": r.progress,
                            "owner": r.owner,
                            "create_time": r.create_time.strftime("%Y-%m-%d %H:%M:%S"),
                            "update_time": r.update_time.strftime("%Y-%m-%d %H:%M:%S"),
                        },
                        **(lambda ver_stats: {
                            "version_count": int(ver_stats[1] or 0),
                            "current_version": int(ver_stats[0] or 0),
                        })(
                            session.query(
                                func.max(PreReviewRun.version_no),
                                func.count(PreReviewRun.id),
                            )
                            .filter(PreReviewRun.project_id == r.project_id)
                            .first()
                            or (0, 0)
                        ),
                    }
                    for r in rows
                ],
                "page": page,
                "page_size": page_size,
                "total": total,
                "total_pages": (total + page_size - 1) // page_size,
            }
        finally:
            session.close()

    def get_project_detail(self, project_id: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.get_project_detail | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            row = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if row is None:
                return False, "project not found", None
            return True, "success", {
                "project_id": row.project_id,
                "project_name": row.project_name,
                "description": row.description,
                "registration_scope": row.registration_scope,
                "registration_path": self._parse_json_list(row.registration_path),
                "registration_leaf": row.registration_leaf,
                "registration_description": row.registration_description,
                "status": row.status,
                "progress": row.progress,
                "owner": row.owner,
                "create_time": row.create_time.strftime("%Y-%m-%d %H:%M:%S"),
                "update_time": row.update_time.strftime("%Y-%m-%d %H:%M:%S"),
            }
        finally:
            session.close()

    def upload_submission(
        self,
        project_id: str,
        file_obj,
        material_category: str = "other",
        section_id: str = "",
        upload_mode: str = "",
    ) -> Tuple[bool, str, Optional[Dict[str, Any]]]: 
        print("[DEBUG] enter PreReviewService.upload_submission | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        if file_obj is None or not getattr(file_obj, "filename", ""):
            return False, "file is required", None

        session = self.db_conn.get_session()
        try:
            project = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if project is None:
                return False, "project not found", None

            display_name = self._safe_display_name(file_obj.filename)
            if not display_name:
                return False, "invalid file name", None
            section_meta = self.ctd_sections.get_section(section_id) if section_id else None
            normalized_mode = str(upload_mode or "").strip().lower()
            if normalized_mode not in {"", "zip", "section", "single"}:
                return False, f"unsupported upload mode: {upload_mode}", None
            upload_items: List[Dict[str, Any]]
            raw_bytes = file_obj.read()
            file_obj.stream.seek(0)
            is_zip_file = self._is_zip_file_name(display_name)
            if normalized_mode == "zip":
                if not is_zip_file:
                    return False, "zip mode requires a .zip file", None
                upload_items = self._extract_zip_submission_items(raw_bytes)
                if not upload_items:
                    return False, "zip contains no supported submission files", None
            else:
                if is_zip_file:
                    return False, "single/section mode does not accept zip files", None
                if normalized_mode == "section" and not str(section_id or "").strip():
                    return False, "section mode requires section_id", None
                upload_items = [
                    {
                        "display_name": display_name,
                        "path": display_name,
                        "file_bytes": raw_bytes,
                        "section_meta": section_meta or {},
                        "explicit_section_id": section_id or "",
                    }
                ]

            saved_items: List[Dict[str, Any]] = []
            for item in upload_items:
                saved = self._create_submission_row(
                    session=session,
                    project_id=project_id,
                    display_name=str(item.get("display_name", "")),
                    file_bytes=item.get("file_bytes", b""),
                    material_category=material_category,
                    section_meta=item.get("section_meta", {}) if isinstance(item.get("section_meta"), dict) else {},
                    relative_path=str(item.get("path", "") or item.get("display_name", "")),
                    explicit_section_id=str(item.get("explicit_section_id", "") or ""),
                    strict_ctd_mapping=self._is_ctd_structure_project(project),
                )
                saved_items.append(saved)
            session.commit()

            for item in saved_items:
                self._load_doc_chunks(project_id=project_id, doc_id=str(item.get("doc_id", "")))

            if len(saved_items) == 1:
                single = saved_items[0]
                return True, "submission uploaded", {
                    **single,
                    "is_chunked": True,
                    "upload_mode": normalized_mode or ("zip" if is_zip_file else "single"),
                }
            return True, "submission uploaded", {
                "items": saved_items,
                "count": len(saved_items),
                "upload_mode": normalized_mode or ("zip" if is_zip_file else "single"),
            }
        except Exception as e:
            session.rollback()
            return False, f"upload submission failed: {str(e)}", None
        finally:
            session.close()

    def list_submissions(self, project_id: str, page: int = 1, page_size: int = 20) -> Dict[str, Any]:
        print("[DEBUG] enter PreReviewService.list_submissions | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            query = session.query(PreReviewSubmissionFile).filter(
                and_(
                    PreReviewSubmissionFile.project_id == project_id,
                    PreReviewSubmissionFile.is_deleted == 0,
                )
            )
            total = query.count()
            page = max(1, int(page))
            page_size = max(1, int(page_size))
            rows = (
                query.order_by(desc(PreReviewSubmissionFile.id))
                .offset((page - 1) * page_size)
                .limit(page_size)
                .all()
            )
            return {
                "list": [
                    {
                        "doc_id": r.doc_id,
                        "project_id": r.project_id,
                        "file_name": r.file_name,
                        "file_type": r.file_type,
                        "material_category": getattr(r, "material_category", "other") or "other",
                        "section_id": getattr(r, "section_id", "") or "",
                        "section_code": getattr(r, "section_code", "") or "",
                        "section_name": getattr(r, "section_name", "") or "",
                        "section_path": self._parse_json_list(getattr(r, "section_path", "") or ""),
                        "is_chunked": bool(r.is_chunked),
                        "chunk_size": r.chunk_size or 0,
                        "create_time": r.create_time.strftime("%Y-%m-%d %H:%M:%S"),
                    }
                    for r in rows
                ],
                "page": page,
                "page_size": page_size,
                "total": total,
                "total_pages": (total + page_size - 1) // page_size,
            }
        finally:
            session.close()

    def get_submission_content(self, project_id: str, doc_id: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.get_submission_content | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            project = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if project is not None and self._is_ctd_structure_project(project):
                payload = self._build_structured_project_payload(session, project_id, source_doc_id=doc_id)
                section_items = payload.get("sections", []) if isinstance(payload.get("sections", []), list) else []
                display_lines = [
                    str(x.get("content", "")).strip()
                    for x in section_items
                    if isinstance(x, dict) and str(x.get("content", "")).strip()
                ]
                return True, "success", {
                    "doc_id": doc_id,
                    "project_id": project_id,
                    "content": "\n\n".join(
                        [str(x.get("text", "")).strip() for x in payload.get("review_units", []) if str(x.get("text", "")).strip()]
                    ),
                    "display_content": "\n\n".join(display_lines),
                    "chunk_size": len(payload.get("review_units", [])),
                    "review_unit_count": len(payload.get("review_units", [])),
                    "section_count": len(section_items),
                    "chapter_structure": payload.get("chapter_structure", []),
                    "sections": section_items,
                    "leaf_sibling_groups": payload.get("leaf_sibling_groups", []),
                    "content_source": "parsed_content",
                }
        finally:
            session.close()

        ok, msg, chunks = self._load_doc_chunks(project_id=project_id, doc_id=doc_id)
        if not ok:
            return False, msg, None

        edited = self._load_submission_edit(doc_id)
        lines = [edited] if edited.strip() else [str(c.get("text", "")).strip() for c in chunks if str(c.get("text", "")).strip()]
        ok_payload, _, payload = self._load_submission_parsed_payload(project_id=project_id, doc_id=doc_id)
        chapter_structure = []
        section_list = []
        leaf_sibling_groups = []
        if ok_payload and isinstance(payload, dict):
            chapter_structure = payload.get("chapter_structure") or []
            section_list = payload.get("sections") or []
            leaf_sibling_groups = payload.get("leaf_sibling_groups") or []
        if isinstance(section_list, list):
            normalized_sections = []
            for item in section_list:
                if not isinstance(item, dict):
                    continue
                raw_content = str(item.get("content", "") or "").strip()
                normalized_sections.append(
                    {
                        **item,
                        "raw_content": raw_content,
                        "cleaned_markdown": "",
                        "display_content": raw_content,
                    }
                )
            section_list = normalized_sections

        display_lines = []
        if isinstance(section_list, list) and section_list:
            for item in section_list:
                if not isinstance(item, dict):
                    continue
                display_value = str(item.get("display_content", "") or "").strip()
                if display_value:
                    display_lines.append(display_value)
        if not display_lines:
            display_lines = lines

        return True, "success", {
            "doc_id": doc_id,
            "project_id": project_id,
            "content": "\n\n".join(lines),
            "display_content": "\n\n".join(display_lines),
            "chunk_size": len(chunks),
            "review_unit_count": len(chunks),
            "section_count": len(section_list) if isinstance(section_list, list) else 0,
            "chapter_structure": chapter_structure if isinstance(chapter_structure, list) else [],
            "sections": section_list if isinstance(section_list, list) else [],
            "leaf_sibling_groups": leaf_sibling_groups if isinstance(leaf_sibling_groups, list) else [],
            "content_source": "parsed_content",
        }

    def get_submission_sections(self, project_id: str, doc_id: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.get_submission_sections | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            project = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if project is not None and self._is_ctd_structure_project(project):
                payload = self._build_structured_project_payload(session, project_id, source_doc_id=doc_id)
                return True, "success", {
                    "doc_id": doc_id,
                    "project_id": project_id,
                    "sections": payload.get("sections", []),
                    "chapter_structure": payload.get("chapter_structure", []),
                    "leaf_sibling_groups": payload.get("leaf_sibling_groups", []),
                    "review_units": payload.get("review_units", []),
                    "statistics": payload.get("statistics", {}),
                }
        finally:
            session.close()

        ok, msg, payload = self._load_submission_parsed_payload(project_id=project_id, doc_id=doc_id)
        if not ok:
            return False, msg, None
        if isinstance(payload, list):
            # Backward compatibility with old list-only parsed format
            return True, "success", {
                "doc_id": doc_id,
                "project_id": project_id,
                "sections": [],
                "chapter_structure": [],
                "leaf_sibling_groups": [],
                "review_units": payload,
            }
        if not isinstance(payload, dict):
            return True, "success", {
                "doc_id": doc_id,
                "project_id": project_id,
                "sections": [],
                "chapter_structure": [],
                "leaf_sibling_groups": [],
                "review_units": [],
            }
        return True, "success", {
            "doc_id": doc_id,
            "project_id": project_id,
            "sections": payload.get("sections", []),
            "chapter_structure": payload.get("chapter_structure", []),
            "leaf_sibling_groups": payload.get("leaf_sibling_groups", []),
            "review_units": payload.get("review_units", []),
            "statistics": payload.get("statistics", {}),
        }

    def get_submission_file_info(self, project_id: str, doc_id: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.get_submission_file_info | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            row = (
                session.query(PreReviewSubmissionFile)
                .filter(
                    and_(
                        PreReviewSubmissionFile.project_id == project_id,
                        PreReviewSubmissionFile.doc_id == doc_id,
                        PreReviewSubmissionFile.is_deleted == 0,
                    )
                )
                .first()
            )
            if row is None:
                return False, "submission file not found", None
            if not os.path.exists(row.file_path):
                return False, "submission file path not exists", None
            return True, "success", {
                "doc_id": row.doc_id,
                "project_id": row.project_id,
                "file_name": row.file_name,
                "file_type": row.file_type,
                "file_path": row.file_path,
                "section_id": getattr(row, "section_id", "") or "",
                "section_code": getattr(row, "section_code", "") or "",
                "section_name": getattr(row, "section_name", "") or "",
                "section_path": self._parse_json_list(getattr(row, "section_path", "") or ""),
            }
        finally:
            session.close()

    def save_submission_content(self, project_id: str, doc_id: str, content: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.save_submission_content | core:", {"project_id": project_id, "doc_id": doc_id})
        ok, msg, data = self.get_submission_file_info(project_id=project_id, doc_id=doc_id)
        if not ok:
            return False, msg, None
        path = self._save_submission_edit(doc_id=doc_id, content=content)
        return True, "content saved", {
            "doc_id": doc_id,
            "project_id": project_id,
            "edit_path": path,
            "char_count": len(str(content or "")),
            "file_name": data.get("file_name", "") if isinstance(data, dict) else "",
        }


    def update_section_concerns(
        self,
        project_id: str,
        section_id: str,
        concern_points: List[Any],
    ) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.update_section_concerns | core:", {"project_id": project_id, "section_id": section_id})
        section = self.ctd_sections.get_section(section_id)
        if not section:
            return False, "section not found", None
        normalized = self._dedupe_text_list(concern_points)
        session = self.db_conn.get_session()
        try:
            row = (
                session.query(PreReviewSectionConcern)
                .filter(
                    and_(
                        PreReviewSectionConcern.project_id == project_id,
                        PreReviewSectionConcern.section_id == section_id,
                    )
                )
                .first()
            )
            now = self._now()
            payload = self._safe_json_list(normalized)
            if row is None:
                row = PreReviewSectionConcern(
                    project_id=project_id,
                    section_id=section_id,
                    section_code=str(section.get("section_code", section_id)),
                    section_name=str(section.get("section_name", section_id)),
                    concern_points=payload,
                    create_time=now,
                    update_time=now,
                )
                session.add(row)
            else:
                row.section_code = str(section.get("section_code", section_id))
                row.section_name = str(section.get("section_name", section_id))
                row.concern_points = payload
                row.update_time = now
            session.commit()
        finally:
            session.close()
        return True, "section concerns updated", {
            "project_id": project_id,
            "section_id": section_id,
            "concern_points": normalized,
        }

    def _load_doc_chunks(self, project_id: str, doc_id: str) -> Tuple[bool, str, List[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService._load_doc_chunks | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            project = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if project is not None and self._is_ctd_structure_project(project):
                payload = self._build_structured_project_payload(session, project_id, source_doc_id=doc_id)
                units = payload.get("review_units", []) if isinstance(payload, dict) else []
                return True, "ok", units if isinstance(units, list) else []
        finally:
            session.close()

        ok, msg, payload = self._load_submission_parsed_payload(project_id=project_id, doc_id=doc_id)
        if not ok:
            return False, msg, []
        if isinstance(payload, dict):
            units = payload.get("review_units")
            if isinstance(units, list):
                return True, "ok", units
            sections = payload.get("sections")
            if isinstance(sections, list):
                fallback = []
                for idx, s in enumerate(sections, start=1):
                    if not isinstance(s, dict):
                        continue
                    text = str(s.get("content", "")).strip()
                    if not text:
                        continue
                    fallback.append(
                        {
                            "chunk_id": str(s.get("section_id", f"sec_{idx}")),
                            "section_id": str(s.get("section_id", f"sec_{idx}")),
                            "section_code": str(s.get("code", "")),
                            "section_name": str(s.get("title", "")),
                            "page": s.get("page_start"),
                            "page_start": s.get("page_start"),
                            "page_end": s.get("page_end"),
                            "text": text,
                        }
                    )
                return True, "ok", fallback
        if isinstance(payload, list):
            return True, "ok", payload
        return True, "ok", []

    def _load_submission_parsed_payload(self, project_id: str, doc_id: str) -> Tuple[bool, str, Any]:
        print("[DEBUG] enter PreReviewService._load_submission_parsed_payload | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        cache_key = self._submission_payload_cache_key(project_id, doc_id)
        if cache_key in self._submission_payload_cache:
            return True, "ok", deepcopy(self._submission_payload_cache[cache_key])

        parsed_path = os.path.join(SUBMISSION_PARSED_DIR, f"{doc_id}.json")
        cached_payload: Any = None
        if os.path.exists(parsed_path):
            try:
                with open(parsed_path, "r", encoding="utf-8") as f:
                    cached_payload = json.load(f)
                self._submission_payload_cache[cache_key] = deepcopy(cached_payload)
            except Exception:
                cached_payload = None

        session = self.db_conn.get_session()
        try:
            project = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            file_row = (
                session.query(PreReviewSubmissionFile)
                .filter(
                    and_(
                        PreReviewSubmissionFile.doc_id == doc_id,
                        PreReviewSubmissionFile.project_id == project_id,
                        PreReviewSubmissionFile.is_deleted == 0,
                    )
                )
                .first()
            )
            if file_row is None:
                return False, "source doc not found", []
            ext_hint = (file_row.file_type or "").strip()
            if not ext_hint and file_row.file_name and "." in file_row.file_name:
                ext_hint = file_row.file_name.rsplit(".", 1)[-1].lower()

            payload: Any = cached_payload
            parsed_now = False
            if payload is None:
                if project is not None and self._is_ctd_structure_project(project):
                    branch_root = self._branch_root_from_section_id(str(file_row.section_id or ""))
                    if branch_root in {"3.2.S", "3.2.P"} and ext_hint.lower() == "pdf" and str(file_row.section_id or "").strip() == branch_root:
                        payload = parse_coarse_ctd_submission(
                            file_path=file_row.file_path,
                            catalog=self.ctd_sections.get_catalog(),
                            root_section_id=branch_root,
                        )
                    elif branch_root == "3.2.S" and ext_hint.lower() == "pdf":
                        payload = parse_ctd_api_pdf_to_payload(
                            file_path=file_row.file_path,
                            catalog=self.ctd_sections.get_catalog(),
                            root_section_id=branch_root,
                        )
                    elif branch_root in {"3.2.S", "3.2.P"}:
                        payload = parse_strict_ctd_submission(
                            file_path=file_row.file_path,
                            catalog=self.ctd_sections.get_catalog(),
                            root_section_id=branch_root,
                            ext_hint=ext_hint,
                        )
                    elif ext_hint.lower() == "pdf":
                        payload = parse_submission_pdf_to_payload(file_path=file_row.file_path)
                    else:
                        chunks = self._parse_submission_file(file_path=file_row.file_path, ext_hint=ext_hint)
                        payload = {"review_units": chunks}
                elif ext_hint.lower() == "pdf":
                    payload = parse_submission_pdf_to_payload(file_path=file_row.file_path)
                    if not isinstance(payload, dict):
                        payload = {"review_units": []}
                else:
                    chunks = self._parse_submission_file(file_path=file_row.file_path, ext_hint=ext_hint)
                    payload = {"review_units": chunks}

                with open(parsed_path, "w", encoding="utf-8") as f:
                    json.dump(payload, f, ensure_ascii=False, indent=2)
                parsed_now = True

            if not isinstance(payload, (list, dict)):
                payload = {"review_units": []}

            self._sync_submission_section_content_rows(session, file_row, payload)
            index_ok, index_msg, indexed_count = self._index_submission_payload(
                file_row=file_row,
                payload=payload,
                force_reindex=parsed_now,
            )

            units = payload.get("review_units", []) if isinstance(payload, dict) else []
            if not isinstance(units, list):
                units = []
            if not units:
                units = [
                    {
                        "chunk_id": item.get("section_id"),
                        "section_id": item.get("section_id"),
                        "section_code": item.get("section_code"),
                        "section_name": item.get("section_name"),
                        "text": item.get("content"),
                    }
                    for item in self._extract_section_content_rows_from_payload(
                        payload,
                        fallback_section_id=str(file_row.section_id or ""),
                        fallback_section_name=str(file_row.section_name or ""),
                    )
                ]

            chunk_ids = [str(c.get("chunk_id", idx + 1)) for idx, c in enumerate(units) if isinstance(c, dict)]
            file_row.is_chunked = True
            file_row.chunk_ids = ";".join(chunk_ids)
            file_row.chunk_size = len(units)
            session.commit()
            if not index_ok:
                print(f"[SubmissionIndexWarn] doc_id={file_row.doc_id} message={index_msg}")
            elif indexed_count:
                print(f"[SubmissionIndex] doc_id={file_row.doc_id} indexed_count={indexed_count}")
            self._submission_payload_cache[cache_key] = deepcopy(payload)
            return True, "ok", deepcopy(payload)
        except Exception as e:
            session.rollback()
            return False, f"load doc chunks failed: {str(e)}", {}
        finally:
            session.close()

    def _parse_submission_file(self, file_path: str, ext_hint: str = "") -> List[Dict[str, Any]]:
        """
        Submission parsing pipeline for pre-review only.
        Keep it isolated from knowledge-base semantic indexing pipeline.
        """
        print("[DEBUG] enter PreReviewService._parse_submission_file | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        ext = (ext_hint or "").strip().lower()
        if not ext and "." in os.path.basename(file_path):
            ext = os.path.basename(file_path).rsplit(".", 1)[-1].lower()
        # For PDF submissions, use heading-based markdown parser.
        if ext == "pdf":
            parsed = parse_submission_pdf_to_payload(file_path=file_path)
            return parsed.get("review_units", []) if isinstance(parsed, dict) else []

        raw_chunks = ParserManager.parse(file_path, ext_hint=ext_hint)
        out: List[Dict[str, Any]] = []
        for idx, item in enumerate(raw_chunks, start=1):
            if not isinstance(item, dict):
                continue
            text = str(item.get("text", "")).strip()
            if not text:
                continue
            out.append(
                {
                    "chunk_id": str(item.get("chunk_id", f"sub_{idx}")),
                    "section_id": str(item.get("chunk_id", f"sub_{idx}")),
                    "section_code": str(item.get("chunk_id", f"sub_{idx}")),
                    "section_name": str(item.get("section_name") or item.get("title") or f"section-{idx}"),
                    "page": item.get("page"),
                    "text": text,
                    "pipeline": "submission_pre_review",
                }
            )
        return out

    def _next_version(self, session, project_id: str) -> int:
        print("[DEBUG] enter PreReviewService._next_version | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        max_ver = session.query(func.max(PreReviewRun.version_no)).filter(PreReviewRun.project_id == project_id).scalar()
        return int(max_ver or 0) + 1

    def _mark_project_failed(self, project_id: str, run_id: str = "") -> None:
        print("[DEBUG] enter PreReviewService._mark_project_failed | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            row = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if row is not None:
                self._set_project_runtime_status(session, row, status="failed", progress=0.0)
            target_run = None
            run_key = str(run_id or "").strip()
            if run_key:
                target_run = (
                    session.query(PreReviewRun)
                    .filter(PreReviewRun.run_id == run_key)
                    .first()
                )
            elif row is not None:
                target_run = (
                    session.query(PreReviewRun)
                    .filter(
                        and_(
                            PreReviewRun.project_id == project_id,
                            PreReviewRun.finish_time.is_(None),
                        )
                    )
                    .order_by(desc(PreReviewRun.create_time))
                    .first()
                )
            if target_run is not None and target_run.finish_time is None:
                target_run.finish_time = self._now()
                session.commit()
        except Exception:
            session.rollback()
        finally:
            session.close()

    def _set_project_runtime_status(
        self,
        session,
        project: Optional[PreReviewProject],
        status: str,
        progress: Optional[float] = None,
    ) -> None:
        if project is None:
            return
        project.status = str(status or "").strip() or project.status
        if progress is not None:
            project.progress = float(progress)
        project.update_time = self._now()

    def _recover_project_running_state(self, session, project: Optional[PreReviewProject], stale_after_seconds: int = 900) -> bool:
        if project is None or str(project.status or "").strip() != "running":
            return False
        latest_run = (
            session.query(PreReviewRun)
            .filter(PreReviewRun.project_id == project.project_id)
            .order_by(desc(PreReviewRun.create_time))
            .first()
        )
        unfinished_run = (
            session.query(PreReviewRun)
            .filter(
                and_(
                    PreReviewRun.project_id == project.project_id,
                    PreReviewRun.finish_time.is_(None),
                )
            )
            .order_by(desc(PreReviewRun.create_time))
            .first()
        )
        now = self._now()
        if unfinished_run is None:
            recovered_status = "completed" if latest_run is not None and latest_run.finish_time is not None else "created"
            self._set_project_runtime_status(
                session,
                project,
                status=recovered_status,
                progress=1.0 if recovered_status == "completed" else 0.0,
            )
            session.commit()
            return True
        reference_time = getattr(project, "update_time", None) or getattr(unfinished_run, "create_time", None)
        if reference_time is None:
            return False
        age_seconds = max(0.0, (now - reference_time).total_seconds())
        if age_seconds < float(stale_after_seconds):
            return False
        unfinished_run.finish_time = now
        summary_text = str(getattr(unfinished_run, "summary", "") or "").strip()
        stale_note = json.dumps(
            {
                "recovered_as": "failed",
                "reason": "stale_running_state",
                "recovered_at": now.strftime("%Y-%m-%d %H:%M:%S"),
            },
            ensure_ascii=False,
        )
        unfinished_run.summary = f"{summary_text}\n{stale_note}".strip() if summary_text else stale_note
        self._set_project_runtime_status(session, project, status="failed", progress=0.0)
        session.commit()
        return True

    def run_pre_review(
        self,
        project_id: str,
        source_doc_id: str,
        run_config: Optional[Dict[str, Any]] = None,
        progress_callback=None,
    ) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.run_pre_review | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        run_id = ""
        try:
            run_config = self._apply_active_prompt_version(run_config)
            workflow_mode = self._resolve_pre_review_mode(run_config)
            self._emit_progress(
                progress_callback,
                "run_start",
                "开始预审任务。",
                project_id=project_id,
                source_doc_id=source_doc_id,
                workflow_mode=workflow_mode,
                strategy=str(run_config.get("strategy", "") or ""),
            )
            project = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
            if project is None:
                return False, "project not found", None
            self._recover_project_running_state(session, project)
            if str(project.status or "").strip() == "running":
                return False, "project is already running", None

            ok, msg, chunks = self._load_doc_chunks(project_id=project_id, doc_id=source_doc_id)
            if not ok:
                return False, msg, None
            if not chunks:
                return False, "no chunk data for pre-review", None
            self._emit_progress(
                progress_callback,
                "chunks_loaded",
                f"已加载 {len(chunks)} 个章节单元。",
                project_id=project_id,
                source_doc_id=source_doc_id,
                chunk_count=len(chunks),
            )

            run_id = self._new_run_id()
            version = self._next_version(session, project_id)

            run_row = PreReviewRun(
                run_id=run_id,
                project_id=project_id,
                version_no=version,
                source_doc_id=source_doc_id,
                strategy=str(run_config.get("strategy", "multi_agent_plan_and_solve+reflection") or "multi_agent_plan_and_solve+reflection"),
                accuracy=None,
                summary="",
                create_time=self._now(),
                finish_time=None,
            )
            session.add(run_row)
            self._emit_progress(
                progress_callback,
                "run_created",
                f"已创建运行记录 {run_id}。",
                project_id=project_id,
                source_doc_id=source_doc_id,
                run_id=run_id,
                version_no=version,
            )

            self._set_project_runtime_status(session, project, status="running", progress=0.1)
            session.flush()
            session.commit()

            # Seed memory before section review:
            # 1) historical feedback memory (episodic)
            self._seed_historical_feedback_memory(project_id=project_id)
            # 2) submission structure memory (semantic)
            ordered_chunks = self._order_review_units(chunks)
            self._seed_submission_structure_memory(
                project_id=project_id,
                source_doc_id=source_doc_id,
                review_units=ordered_chunks,
            )
            self._emit_progress(
                progress_callback,
                "section_queue",
                f"待审章节数：{len(ordered_chunks)}。",
                project_id=project_id,
                source_doc_id=source_doc_id,
                run_id=run_id,
                section_count=len(ordered_chunks),
                section_ids=[str((chunk or {}).get("section_id", "") or "") for chunk in ordered_chunks if isinstance(chunk, dict)],
            )

            conclusions = []
            section_results: List[Dict[str, Any]] = []
            skipped_empty = 0
            previous_section_meta: Dict[str, Any] = {}
            for chunk in ordered_chunks:
                section_result = self._review_single_chunk(
                    project=project,
                    project_id=project_id,
                    run_id=run_id,
                    source_doc_id=source_doc_id,
                    chunk=chunk,
                    previous_section_meta=previous_section_meta,
                    run_config=run_config,
                    progress_callback=progress_callback,
                )
                if not section_result.get("success"):
                    skipped_empty += 1
                    self._emit_progress(
                        progress_callback,
                        "section_skipped",
                        "章节无可用内容，已跳过。",
                        run_id=run_id,
                        section_id=str((chunk or {}).get("section_id", "") or ""),
                        section_name=str((chunk or {}).get("section_name", "") or ""),
                    )
                    continue
                conclusions.append(section_result["conclusion_row"])
                session.add(
                    PreReviewSectionOutput(
                        run_id=run_id,
                        section_id=str(section_result.get("section_id", "") or ""),
                        section_name=str(section_result.get("section_meta", {}).get("section_name", "") or ""),
                        schema_version="chapter_review_v1",
                        output_json=json.dumps(section_result.get("review_result", {}), ensure_ascii=False),
                        create_time=self._now(),
                    )
                )
                session.add(section_result["trace_row"])
                section_results.append(section_result)
                previous_section_meta = section_result.get("previous_section_meta", previous_section_meta)
                self._emit_progress(
                    progress_callback,
                    "section_done",
                    "章节预审完成。",
                    run_id=run_id,
                    section_id=str(section_result.get("section_id", "") or ""),
                    section_name=str(section_result.get("section_meta", {}).get("section_name", "") or ""),
                    conclusion=str((section_result.get("review_result", {}) or {}).get("pre_review_conclusion", "") or ""),
                )

            if not conclusions:
                raise ValueError("no non-empty sections after parsing")

            session.bulk_save_objects(conclusions)
            if workflow_mode == SINGLE_SECTION_PRE_REVIEW_V2:
                post_review = {
                    "mode": SINGLE_SECTION_PRE_REVIEW_V2,
                    "consistency_result": {},
                    "run_qa_result": {},
                    "lead_result": {},
                    "trace": [],
                    "skipped": ["consistency", "run_qa", "lead_summary"],
                }
            else:
                post_review = self._run_post_review_agents(
                    project=project,
                    run_id=run_id,
                    source_doc_id=source_doc_id,
                    section_results=section_results,
                    run_config=run_config,
                )
            run_row.summary = json.dumps(
                {
                    "mode": workflow_mode,
                    "run_config_summary": {
                        "domain": str(run_config.get("domain", "") or ""),
                        "branch": str(run_config.get("branch", "") or ""),
                        "strategy": str(run_config.get("strategy", "") or ""),
                        "workflow_mode": str(workflow_mode or ""),
                        "feedback_loop_mode": str(run_config.get("feedback_loop_mode", "") or "feedback_optimize"),
                        "enable_feedback_optimize": bool(run_config.get("enable_feedback_optimize", True)),
                        "prompt_config": dict(run_config.get("prompt_config", {}) or {}) if isinstance(run_config.get("prompt_config", {}), dict) else {},
                    },
                    "section_conclusion_count": len(conclusions),
                    "skipped_empty_count": skipped_empty,
                    "consistency_result": post_review.get("consistency_result", {}),
                    "run_qa_result": post_review.get("run_qa_result", {}),
                    "lead_result": post_review.get("lead_result", {}),
                    "metrics": {},
                    "metrics_schema": "run_metrics_v1",
                },
                ensure_ascii=False,
            )
            self._refresh_run_metrics_summary(session, run_id)
            run_row.finish_time = self._now()

            self._set_project_runtime_status(session, project, status="completed", progress=1.0)

            session.commit()
            result = {
                "run_id": run_id,
                "version_no": version,
                "conclusion_count": len(conclusions),
                "section_output_count": len(section_results),
                "section_output_section_ids": [str(item.get("section_id", "") or "") for item in section_results],
                "schema_version": "chapter_review_v1",
                "workflow_mode": workflow_mode,
                "skipped_empty_count": skipped_empty,
                "post_review": post_review,
            }
            self._emit_progress(
                progress_callback,
                "run_done",
                "预审任务完成。",
                project_id=project_id,
                source_doc_id=source_doc_id,
                run_id=run_id,
                section_output_count=len(section_results),
                skipped_empty_count=skipped_empty,
                workflow_mode=workflow_mode,
            )
            return True, "pre-review completed", result
        except Exception as e:
            session.rollback()
            self._mark_project_failed(project_id, run_id=run_id)
            self._emit_progress(
                progress_callback,
                "run_failed",
                f"预审任务失败：{str(e)}",
                project_id=project_id,
                source_doc_id=source_doc_id,
                error=str(e),
            )
            return False, f"run pre-review failed: {str(e)}", None
        finally:
            session.close()

    def _build_structured_project_payload(self, session, project_id: str, source_doc_id: str = "") -> Dict[str, Any]:
        source_doc_key = str(source_doc_id or "").strip()
        catalog = self._merge_catalog_with_manual_concerns(
            self.ctd_sections.get_catalog(),
            self._load_manual_concern_map(session, project_id),
        )
        chapter_structure = deepcopy(catalog.get("chapter_structure", []))
        flat_sections = catalog.get("flat_sections", [])
        section_map = catalog.get("section_map", {})
        rows = (
            session.query(PreReviewSubmissionFile)
            .filter(
                and_(
                    PreReviewSubmissionFile.project_id == project_id,
                    PreReviewSubmissionFile.is_deleted == 0,
                    PreReviewSubmissionFile.section_id.isnot(None),
                )
            )
            .order_by(PreReviewSubmissionFile.id.asc())
            .all()
        )
        if source_doc_key:
            rows = [row for row in rows if str(getattr(row, "doc_id", "") or "").strip() == source_doc_key]
        section_content_rows = (
            session.query(PreReviewSubmissionSectionContent)
            .filter(PreReviewSubmissionSectionContent.project_id == project_id)
            .order_by(PreReviewSubmissionSectionContent.id.asc())
            .all()
        )
        if source_doc_key:
            section_content_rows = [
                row for row in section_content_rows if str(getattr(row, "doc_id", "") or "").strip() == source_doc_key
            ]
        file_by_doc_id = {str(row.doc_id): row for row in rows}
        doc_ids_with_section_content = {
            str(getattr(item, "doc_id", "") or "").strip()
            for item in section_content_rows
            if str(getattr(item, "doc_id", "") or "").strip()
        }
        attachments_by_section: Dict[str, List[Dict[str, Any]]] = {}
        if section_content_rows:
            for content_row in section_content_rows:
                sid = str(getattr(content_row, "section_id", "") or "").strip()
                if not sid:
                    continue
                file_row = file_by_doc_id.get(str(getattr(content_row, "doc_id", "") or "").strip())
                if file_row is None:
                    continue
                raw_text = str(getattr(content_row, "content", "") or "").strip()
                attachments_by_section.setdefault(sid, []).append(
                    {
                        "doc_id": file_row.doc_id,
                        "file_name": file_row.file_name,
                        "file_type": file_row.file_type,
                        "material_category": getattr(file_row, "material_category", "other") or "other",
                        "section_id": sid,
                        "section_code": getattr(content_row, "section_code", "") or sid,
                        "section_name": getattr(content_row, "section_name", "") or "",
                        "section_path": self._parse_json_list(getattr(file_row, "section_path", "") or ""),
                        "text": raw_text,
                        "raw_text": raw_text,
                        "cleaned_markdown": "",
                        "display_text": raw_text,
                    }
                )
        for row in rows:
            sid = str(getattr(row, "section_id", "") or "").strip()
            if not sid or sid in attachments_by_section or str(row.doc_id) in doc_ids_with_section_content:
                continue
            text_blocks: List[str] = []
            edited_text = self._load_submission_edit(row.doc_id)
            if edited_text.strip():
                text_blocks.append(edited_text.strip())
            ok_payload, _, payload = self._load_submission_parsed_payload(project_id=project_id, doc_id=row.doc_id)
            if ok_payload and not text_blocks:
                units: List[Dict[str, Any]] = []
                if isinstance(payload, dict):
                    if isinstance(payload.get("review_units"), list):
                        units = payload.get("review_units") or []
                    elif isinstance(payload.get("sections"), list):
                        units = payload.get("sections") or []
                elif isinstance(payload, list):
                    units = payload
                for unit in units:
                    if not isinstance(unit, dict):
                        continue
                    text_value = str(unit.get("text") or unit.get("content") or "").strip()
                    if text_value:
                        text_blocks.append(text_value)
            raw_text = self._preview_text_blocks(text_blocks)
            attachments_by_section.setdefault(sid, []).append(
                {
                    "doc_id": row.doc_id,
                    "file_name": row.file_name,
                    "file_type": row.file_type,
                    "material_category": getattr(row, "material_category", "other") or "other",
                    "section_id": sid,
                    "section_code": getattr(row, "section_code", "") or sid,
                    "section_name": getattr(row, "section_name", "") or "",
                    "section_path": self._parse_json_list(getattr(row, "section_path", "") or ""),
                    "text": raw_text,
                    "raw_text": raw_text,
                    "cleaned_markdown": "",
                    "display_text": raw_text,
                }
            )

        sections: List[Dict[str, Any]] = []
        review_units: List[Dict[str, Any]] = []
        for index, base in enumerate(flat_sections, start=1):
            sid = str(base.get("section_id", "")).strip()
            attached_files = attachments_by_section.get(sid, [])
            concern_points = list(base.get("concern_points") or [])
            merged_raw_text = self._preview_text_blocks(
                [f"[{item['file_name']}]\n{item.get('raw_text', '')}" for item in attached_files if item.get("raw_text", "")]
            )
            merged_display_text = self._preview_text_blocks(
                [f"[{item['file_name']}]\n{item.get('display_text', '')}" for item in attached_files if item.get("display_text", "")]
            )
            section_item = {
                "section_id": sid,
                "code": base.get("section_code", sid),
                "title": base.get("section_name", sid),
                "section_name": base.get("section_name", sid),
                "title_path": list(base.get("title_path") or []),
                "parent_section_id": base.get("parent_section_id", ""),
                "concern_points": concern_points,
                "attached_files": attached_files,
                "file_count": len(attached_files),
                "content": merged_raw_text,
                "raw_content": merged_raw_text,
                "cleaned_markdown": "",
                "display_content": merged_raw_text,
                "content_preview": self._preview(merged_raw_text, 320),
                "char_count": len(merged_raw_text),
                "paragraph_blocks": self._build_paragraph_anchors(
                    text=merged_raw_text,
                    section_id=sid,
                    section_code=str(base.get("section_code", sid)),
                ),
            }
            sections.append(section_item)
            if merged_raw_text:
                review_units.append(
                    {
                        "chunk_id": sid,
                        "section_id": sid,
                        "section_code": base.get("section_code", sid),
                        "section_name": base.get("section_name", sid),
                        "parent_section_id": base.get("parent_section_id", ""),
                        "parent_code": section_map.get(base.get("parent_section_id", ""), {}).get("section_code", ""),
                        "page": None,
                        "page_start": None,
                        "page_end": None,
                        "text": merged_raw_text,
                        "title_path": list(base.get("title_path") or []),
                        "unit_order": index,
                        "unit_type": "ctd_section_bundle",
                        "attached_files": attached_files,
                        "concern_points": concern_points,
                        "paragraph_blocks": section_item["paragraph_blocks"],
                    }
                )

        section_content_map = {str(item.get("section_id", "")): item for item in sections if str(item.get("section_id", ""))}

        def attach_content(nodes: List[Dict[str, Any]]) -> None:
            for node in nodes or []:
                sid = str(node.get("section_id", "")).strip()
                matched = section_content_map.get(sid, {})
                node["content"] = str(matched.get("content", "") or "")
                node["content_preview"] = str(matched.get("content_preview", "") or "")
                node["char_count"] = int(matched.get("char_count", 0) or 0)
                attach_content(node.get("children_sections") or [])

        attach_content(chapter_structure)

        return {
            "chapter_structure": chapter_structure,
            "sections": sections,
            "leaf_sibling_groups": [],
            "review_units": review_units,
            "statistics": {
                "section_total": len(sections),
                "review_unit_total": len(review_units),
                "attached_file_total": sum(len(v) for v in attachments_by_section.values()),
            },
        }

    def get_ctd_section_catalog(self, project_id: str = "") -> Dict[str, Any]:
        print("[DEBUG] enter PreReviewService.get_ctd_section_catalog | core:", {"project_id": project_id})
        catalog = self.ctd_sections.get_catalog()
        if project_id:
            session = self.db_conn.get_session()
            try:
                manual_map = self._load_manual_concern_map(session, project_id)
            finally:
                session.close()
            catalog = self._merge_catalog_with_manual_concerns(catalog, manual_map)
        return {
            "chapter_structure": catalog.get("chapter_structure", []),
            "flat_sections": catalog.get("flat_sections", []),
        }

    def _create_submission_row(
        self,
        session,
        project_id: str,
        display_name: str,
        file_bytes: bytes,
        material_category: str,
        section_meta: Optional[Dict[str, Any]] = None,
        relative_path: str = "",
        explicit_section_id: str = "",
        strict_ctd_mapping: bool = False,
    ) -> Dict[str, Any]:
        section_meta = section_meta or {}
        file_type = display_name.rsplit(".", 1)[-1].lower() if "." in display_name else ""
        if not ParserManager.is_supported(file_type):
            raise ValueError(f"unsupported file type: {file_type or 'unknown'}")

        doc_id = self._new_submission_doc_id()
        storage_name = self._submission_storage_name(doc_id, display_name)
        file_path = os.path.join(SUBMISSION_UPLOAD_DIR, storage_name)
        with open(file_path, "wb") as fh:
            fh.write(file_bytes)

        if strict_ctd_mapping:
            catalog = self.ctd_sections.get_catalog()
            branch_root = self._infer_submission_branch_root(
                material_category=material_category,
                explicit_section_id=explicit_section_id,
                section_meta=section_meta,
                display_name=display_name,
                relative_path=relative_path or display_name,
            )
            if branch_root in {"3.2.S", "3.2.P"} and (
                not explicit_section_id or str(explicit_section_id).strip() == branch_root
            ):
                section_meta = catalog.get("section_map", {}).get(branch_root, {})
            else:
                parsed_payload: Optional[Dict[str, Any]] = None
                try:
                    parsed_payload = parse_submission_pdf_to_payload(file_path=file_path)
                except Exception:
                    parsed_payload = None
                mapped = self.submission_parser_agent.map_file(
                    display_name=display_name,
                    relative_path=relative_path or display_name,
                    parsed_payload=parsed_payload,
                    catalog=catalog,
                    explicit_section_id=explicit_section_id or str(section_meta.get("section_id", "") or ""),
                )
                section_meta = mapped.get("section_meta", {}) if isinstance(mapped.get("section_meta", {}), dict) else {}
                if not str(section_meta.get("section_id", "") or "").strip():
                    raise ValueError(
                        f"strict ctd mapping failed for submission file: {display_name}, path={relative_path or display_name}"
                    )

        section_id = str(section_meta.get("section_id", "")).strip() or None
        section_code = str(section_meta.get("section_code", "")).strip() or section_id
        section_name = str(section_meta.get("section_name", "")).strip() or None
        section_path = self._safe_json_list(section_meta.get("title_path", [])) if section_id else None

        row = PreReviewSubmissionFile(
            doc_id=doc_id,
            project_id=project_id,
            file_name=display_name,
            file_path=file_path,
            file_type=file_type,
            material_category=str(material_category or "other").strip() or "other",
            section_id=section_id,
            section_code=section_code,
            section_name=section_name,
            section_path=section_path,
            is_chunked=False,
            chunk_ids="",
            chunk_size=0,
            is_deleted=False,
            create_time=self._now(),
        )
        session.add(row)
        shadow = session.query(FileInfo).filter(FileInfo.doc_id == doc_id).first()
        if shadow is None:
            session.add(
                FileInfo(
                    doc_id=doc_id,
                    file_name=display_name,
                    file_path=file_path,
                    file_type=file_type,
                    classification="submission_material",
                    affect_range="pre_review",
                    is_chunked=0,
                    chunk_ids="",
                    chunk_size=0,
                    is_deleted=1,
                    create_time=self._now().strftime("%Y-%m-%d %H:%M:%S"),
                    review_status=0,
                    review_time=None,
                )
            )
        return {
            "doc_id": doc_id,
            "project_id": project_id,
            "file_name": display_name,
            "file_type": file_type,
            "material_category": str(material_category or "other").strip() or "other",
            "section_id": section_id or "",
            "section_code": section_code or "",
            "section_name": section_name or "",
            "section_path": json.loads(section_path) if section_path else [],
            "file_path": file_path,
        }

    def _extract_zip_submission_items(self, zip_bytes: bytes) -> List[Dict[str, Any]]:
        items: List[Dict[str, Any]] = []
        with zipfile.ZipFile(BytesIO(zip_bytes)) as archive:
            for member in archive.infolist():
                if member.is_dir():
                    continue
                raw_name = self._decode_zip_name(member.filename)
                safe_name = self._sanitize_member_name(raw_name)
                if not safe_name:
                    continue
                display_name = os.path.basename(safe_name)
                if not display_name:
                    continue
                suffix = display_name.rsplit(".", 1)[-1].lower() if "." in display_name else ""
                if not ParserManager.is_supported(suffix):
                    continue
                section_id = self.ctd_sections.infer_section_id_from_path(safe_name)
                section_meta = self.ctd_sections.get_section(section_id) if section_id else None
                payload = archive.read(member)
                items.append(
                    {
                        "display_name": display_name,
                        "path": safe_name,
                        "file_bytes": payload,
                        "section_meta": section_meta or {},
                        "explicit_section_id": "",
                    }
                )
        return items

    def run_section_replay(
        self,
        project_id: str,
        source_doc_id: str,
        section_id: str,
        run_config: Optional[Dict[str, Any]] = None,
    ) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.run_section_replay | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        run_config = self._apply_active_prompt_version(run_config)
        session = self.db_conn.get_session()
        try:
            project = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.project_id == project_id, PreReviewProject.is_deleted == 0))
                .first()
            )
        finally:
            session.close()
        if project is None:
            return False, "project not found", None
        ok, msg, chunks = self._load_doc_chunks(project_id=project_id, doc_id=source_doc_id)
        if not ok:
            return False, msg, None
        ordered_chunks = self._order_review_units(chunks)
        target_section_id = str(section_id or "").strip()
        target_chunk: Optional[Dict[str, Any]] = None
        previous_section_meta: Dict[str, Any] = {}
        for chunk in ordered_chunks:
            current_section_id = str(chunk.get("section_id") or chunk.get("chunk_id", "")).strip()
            if current_section_id == target_section_id:
                target_chunk = chunk
                break
            text = str(chunk.get("text", "")).strip()
            if text:
                section_name = str(chunk.get("section_name") or chunk.get("title") or current_section_id)
                previous_section_meta = {
                    "section_id": current_section_id,
                    "conclusion_preview": self._preview(f"{current_section_id} {section_name}", max_len=180),
                }
        if target_chunk is None:
            return False, f"section not found: {target_section_id}", None

        replay_run_id = str(run_config.get("replay_run_id", "") or self._new_run_id())
        self._seed_historical_feedback_memory(project_id=project_id)
        self._seed_submission_structure_memory(
            project_id=project_id,
            source_doc_id=source_doc_id,
            review_units=ordered_chunks,
        )
        section_result = self._review_single_chunk(
            project=project,
            project_id=project_id,
            run_id=replay_run_id,
            source_doc_id=source_doc_id,
            chunk=target_chunk,
            previous_section_meta=previous_section_meta,
            run_config=run_config,
        )
        if not section_result.get("success"):
            return False, str(section_result.get("message", "section replay failed")), None

        session = self.db_conn.get_session()
        conclusion_record = {}
        trace_record = {}
        try:
            version = self._next_version(session, project_id)
            replay_run = PreReviewRun(
                run_id=replay_run_id,
                project_id=project_id,
                version_no=version,
                source_doc_id=source_doc_id,
                strategy=str(run_config.get("strategy", "section_replay") or "section_replay"),
                accuracy=None,
                summary="",
                create_time=self._now(),
                finish_time=None,
            )
            session.add(replay_run)
            session.add(section_result["conclusion_row"])
            session.add(
                PreReviewSectionOutput(
                    run_id=replay_run_id,
                    section_id=str(section_result.get("section_id", "") or ""),
                    section_name=str(section_result.get("section_meta", {}).get("section_name", "") or ""),
                    schema_version="chapter_review_v1",
                    output_json=json.dumps(section_result.get("review_result", {}), ensure_ascii=False),
                    create_time=self._now(),
                )
            )
            session.add(section_result["trace_row"])
            replay_run.summary = json.dumps(
                {
                    "mode": "section_replay",
                    "run_config_summary": {
                        "domain": str(run_config.get("domain", "") or ""),
                        "branch": str(run_config.get("branch", "") or ""),
                        "strategy": str(run_config.get("strategy", "section_replay") or "section_replay"),
                        "workflow_mode": "section_replay",
                        "feedback_loop_mode": str(run_config.get("feedback_loop_mode", "") or "feedback_optimize"),
                        "enable_feedback_optimize": bool(run_config.get("enable_feedback_optimize", True)),
                        "prompt_config": dict(run_config.get("prompt_config", {}) or {}) if isinstance(run_config.get("prompt_config", {}), dict) else {},
                    },
                    "section_conclusion_count": 1,
                    "skipped_empty_count": 0,
                    "metrics": {},
                    "metrics_schema": "run_metrics_v1",
                },
                ensure_ascii=False,
            )
            self._refresh_run_metrics_summary(session, replay_run_id)
            replay_run.finish_time = self._now()
            session.flush()
            conclusion_record = SectionConclusionRecord.from_entity(section_result["conclusion_row"]).to_dict()
            trace_record = SectionTraceRecord.from_entity(section_result["trace_row"]).to_dict()
            session.commit()
        except Exception as exc:
            session.rollback()
            return False, f"section replay persist failed: {str(exc)}", None
        finally:
            session.close()

        return True, "section replay completed", {
            "run_id": replay_run_id,
            "project_id": project_id,
            "source_doc_id": source_doc_id,
            "section_id": target_section_id,
            "strategy": str(run_config.get("strategy", "section_replay") or "section_replay"),
            "section_meta": section_result.get("section_meta", {}),
            "conclusion": conclusion_record,
            "trace": trace_record,
        }

    def _seed_historical_feedback_memory(self, project_id: str) -> None:
        """
        Load previous feedback in same project into episodic memory for context injection.
        """
        print("[DEBUG] enter PreReviewService._seed_historical_feedback_memory | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            rows = (
                session.query(PreReviewFeedback, PreReviewRun)
                .join(PreReviewRun, PreReviewFeedback.run_id == PreReviewRun.run_id)
                .filter(PreReviewRun.project_id == project_id)
                .order_by(desc(PreReviewFeedback.id))
                .limit(200)
                .all()
            )
            for fb, run in rows:
                key = f"fb:{run.run_id}:{fb.id}"
                value = f"{fb.feedback_type} {fb.feedback_text or ''} {fb.suggestion or ''}".strip()
                self.memory_tool.remember(
                    key=key,
                    value=value,
                    memory_type="episodic",
                    metadata={
                        "source": "historical_feedback",
                        "project_id": project_id,
                        "run_id": run.run_id,
                        "section_id": fb.section_id or "",
                        "operator": fb.operator or "",
                    },
                )
        finally:
            session.close()

    def get_run_history(self, project_id: str) -> List[Dict[str, Any]]:
        print("[DEBUG] enter PreReviewService.get_run_history | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            rows = (
                session.query(PreReviewRun)
                .filter(PreReviewRun.project_id == project_id)
                .order_by(desc(PreReviewRun.id))
                .all()
            )
            out = []
            for r in rows:
                try:
                    summary_payload = json.loads(r.summary) if r.summary else {}
                except Exception:
                    summary_payload = {}
                feedback_count = session.query(PreReviewFeedback).filter(PreReviewFeedback.run_id == r.run_id).count()
                sub = (
                    session.query(PreReviewSubmissionFile)
                    .filter(PreReviewSubmissionFile.doc_id == r.source_doc_id)
                    .first()
                )
                out.append(
                    {
                        "run_id": r.run_id,
                        "project_id": r.project_id,
                        "version_no": r.version_no,
                        "source_doc_id": r.source_doc_id,
                        "source_file_name": sub.file_name if sub else "",
                        "strategy": r.strategy,
                        "accuracy": r.accuracy,
                        "metrics": (summary_payload.get("metrics", {}) if isinstance(summary_payload.get("metrics", {}), dict) else {}),
                        "feedback_loop_mode": str(
                            (
                                summary_payload.get("run_config_summary", {})
                                if isinstance(summary_payload.get("run_config_summary", {}), dict)
                                else {}
                            ).get("feedback_loop_mode", "")
                            or "feedback_optimize"
                        ),
                        "summary": r.summary,
                        "summary_payload": summary_payload if isinstance(summary_payload, dict) else {},
                        "feedback_count": feedback_count,
                        "create_time": r.create_time.strftime("%Y-%m-%d %H:%M:%S"),
                        "finish_time": r.finish_time.strftime("%Y-%m-%d %H:%M:%S") if r.finish_time else None,
                    }
                )
            return out
        finally:
            session.close()

    def get_section_conclusions(self, run_id: str, section_id: str = "") -> List[Dict[str, Any]]:
        print("[DEBUG] enter PreReviewService.get_section_conclusions | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            def _safe_json_load(raw: Optional[str]) -> List[Any]:
                if not raw:
                    return []
                try:
                    data = json.loads(raw)
                    return data if isinstance(data, list) else []
                except Exception:
                    return []

            run_row = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
            query = session.query(PreReviewSectionConclusion).filter(PreReviewSectionConclusion.run_id == run_id)
            if section_id:
                query = query.filter(
                    and_(
                        PreReviewSectionConclusion.run_id == run_id,
                        (
                            (PreReviewSectionConclusion.section_id == section_id)
                            | (PreReviewSectionConclusion.section_name.like(f"{section_id}%"))
                        ),
                    )
                )
            rows = query.order_by(PreReviewSectionConclusion.id.asc()).all()
            output_rows = (
                session.query(PreReviewSectionOutput)
                .filter(PreReviewSectionOutput.run_id == run_id)
                .all()
            )
            output_by_section: Dict[str, Dict[str, Any]] = {}
            for row in output_rows:
                try:
                    output_by_section[str(row.section_id)] = json.loads(row.output_json or "{}")
                except Exception:
                    output_by_section[str(row.section_id)] = {}
            out = [
                {
                    **SectionConclusionRecord.from_entity(r).to_dict(),
                    "standard_output": output_by_section.get(str(r.section_id), {}),
                }
                for r in rows
            ]
            if (not section_id) and run_row is not None:
                ok, _, payload = self.get_submission_sections(
                    project_id=run_row.project_id,
                    doc_id=run_row.source_doc_id,
                )
                if ok and isinstance(payload, dict):
                    units = payload.get("review_units", [])
                    if isinstance(units, list):
                        order_map: Dict[str, int] = {}
                        for i, u in enumerate(units, start=1):
                            if not isinstance(u, dict):
                                continue
                            sid = str(u.get("section_id") or u.get("chunk_id") or "").strip()
                            if sid and sid not in order_map:
                                order_map[sid] = i
                        if order_map:
                            out = sorted(
                                out,
                                key=lambda x: (
                                    int(order_map.get(str(x.get("section_id", "")), 10**9)),
                                    str(x.get("section_id", "")),
                                ),
                            )
            return out
        finally:
            session.close()

    def get_standardized_section_outputs(self, run_id: str, section_id: str = "") -> List[Dict[str, Any]]:
        session = self.db_conn.get_session()
        try:
            query = session.query(PreReviewSectionOutput).filter(PreReviewSectionOutput.run_id == run_id)
            if section_id:
                query = query.filter(PreReviewSectionOutput.section_id == section_id)
            rows = query.order_by(PreReviewSectionOutput.id.asc()).all()
            out: List[Dict[str, Any]] = []
            for row in rows:
                try:
                    payload = json.loads(row.output_json or "{}")
                except Exception:
                    payload = {}
                out.append(
                    {
                        "run_id": row.run_id,
                        "section_id": row.section_id,
                        "section_name": row.section_name,
                        "schema_version": row.schema_version,
                        "output": payload if isinstance(payload, dict) else {},
                        "create_time": row.create_time.strftime("%Y-%m-%d %H:%M:%S") if row.create_time else "",
                    }
                )
            return out
        finally:
            session.close()

    def get_run_section_overview(self, run_id: str) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.get_run_section_overview | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
            if run is None:
                return False, "run not found", None
            try:
                run_summary = json.loads(run.summary) if run.summary else {}
            except Exception:
                run_summary = {}

            ok, msg, section_payload = self.get_submission_sections(
                project_id=run.project_id,
                doc_id=run.source_doc_id,
            )
            if not ok:
                return False, msg, None

            conclusions = self.get_section_conclusions(run_id=run_id, section_id="")
            section_outputs = self.get_standardized_section_outputs(run_id=run_id, section_id="")
            conclusion_by_section = {str(x.get("section_id")): x for x in conclusions}
            section_output_by_section = {
                str(x.get("section_id")): (x.get("standard_output", {}) if isinstance(x.get("standard_output", {}), dict) else {})
                for x in conclusions
            }
            standardized_output_by_section = {
                str(x.get("section_id")): (x.get("output", {}) if isinstance(x.get("output", {}), dict) else {})
                for x in section_outputs
            }
            traces = self.get_section_traces(run_id=run_id, section_id="")
            trace_digest_by_section = {}
            for t in traces:
                sid = str(t.get("section_id", ""))
                if not sid:
                    continue
                coordination = t.get("coordination", {}) if isinstance(t.get("coordination"), dict) else {}
                retrieval = coordination.get("retrieval", {}) if isinstance(coordination.get("retrieval"), dict) else {}
                agent_meta = t.get("agent", {}) if isinstance(t.get("agent"), dict) else {}
                memory_meta = t.get("memory", {}) if isinstance(t.get("memory"), dict) else {}
                trace_digest_by_section[sid] = {
                    "trace_schema": t.get("trace_schema", "legacy"),
                    "retrieval_hit_count": retrieval.get("hit_count", 0),
                    "grouped_doc_count": retrieval.get("grouped_doc_count", 0),
                    "memory_hit_count": memory_meta.get("hit_count", 0),
                    "findings_count": agent_meta.get("findings_count", 0),
                    "score": agent_meta.get("score", 0.0),
                    "source_breakdown": t.get("retrieval_detail", {}).get("source_breakdown", {}) if isinstance(t.get("retrieval_detail", {}), dict) else {},
                    "retrieval_metrics": t.get("retrieval_detail", {}).get("metrics", {}) if isinstance(t.get("retrieval_detail", {}), dict) else {},
                    "retrieval_error_breakdown": t.get("retrieval_detail", {}).get("error_breakdown", {}) if isinstance(t.get("retrieval_detail", {}), dict) else {},
                }

            return True, "success", {
                "run_id": run_id,
                "project_id": run.project_id,
                "source_doc_id": run.source_doc_id,
                "strategy": str(run.strategy or ""),
                "workflow_mode": str((run_summary if isinstance(run_summary, dict) else {}).get("mode", "") or self._resolve_pre_review_mode({"strategy": str(run.strategy or "")})),
                "run_config": {
                    "strategy": str(
                        (
                            (run_summary.get("run_config_summary", {}) if isinstance(run_summary.get("run_config_summary", {}), dict) else {})
                            .get("strategy", "")
                        )
                        or run.strategy
                        or ""
                    ),
                    "workflow_mode": str(
                        (
                            (run_summary.get("run_config_summary", {}) if isinstance(run_summary.get("run_config_summary", {}), dict) else {})
                            .get("workflow_mode", "")
                        )
                        or (run_summary if isinstance(run_summary, dict) else {}).get("mode", "")
                        or self._resolve_pre_review_mode({"strategy": str(run.strategy or "")})
                    ),
                    "domain": str(
                        (
                            (run_summary.get("run_config_summary", {}) if isinstance(run_summary.get("run_config_summary", {}), dict) else {})
                            .get("domain", "")
                        )
                        or ""
                    ),
                    "branch": str(
                        (
                            (run_summary.get("run_config_summary", {}) if isinstance(run_summary.get("run_config_summary", {}), dict) else {})
                            .get("branch", "")
                        )
                        or ""
                    ),
                    "feedback_loop_mode": str(
                        (
                            (run_summary.get("run_config_summary", {}) if isinstance(run_summary.get("run_config_summary", {}), dict) else {})
                            .get("feedback_loop_mode", "")
                        )
                        or "feedback_optimize"
                    ),
                    "enable_feedback_optimize": bool(
                        (
                            (run_summary.get("run_config_summary", {}) if isinstance(run_summary.get("run_config_summary", {}), dict) else {})
                            .get("enable_feedback_optimize", True)
                        )
                    ),
                    "prompt_config": dict(
                        (
                            (run_summary.get("run_config_summary", {}) if isinstance(run_summary.get("run_config_summary", {}), dict) else {})
                            .get("prompt_config", {})
                        )
                        or {}
                    ),
                },
                "chapter_structure": section_payload.get("chapter_structure", []),
                "sections": section_payload.get("sections", []),
                "leaf_sibling_groups": section_payload.get("leaf_sibling_groups", []),
                "review_units": section_payload.get("review_units", []),
                "conclusion_by_section_id": conclusion_by_section,
                "section_output_by_section_id": section_output_by_section,
                "standardized_output_by_section_id": standardized_output_by_section,
                "section_outputs": section_outputs,
                "schema_version": "chapter_review_v1",
                "trace_digest_by_section_id": trace_digest_by_section,
                "conclusions": conclusions,
                "run_summary": run_summary if isinstance(run_summary, dict) else {},
            }
        finally:
            session.close()

    def get_section_traces(self, run_id: str, section_id: str = "") -> List[Dict[str, Any]]:
        print("[DEBUG] enter PreReviewService.get_section_traces | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            query = session.query(PreReviewSectionTrace).filter(PreReviewSectionTrace.run_id == run_id)
            if section_id:
                query = query.filter(PreReviewSectionTrace.section_id == section_id)
            rows = query.order_by(PreReviewSectionTrace.id.asc()).all()

            out: List[Dict[str, Any]] = []
            for r in rows:
                raw_trace = PreReviewRepository.load_trace_payload(r.trace_json)
                if isinstance(raw_trace, dict) and str(raw_trace.get("trace_schema", "") or "").strip() == "chapter_review_v1":
                    trace_schema = "chapter_review_v1"
                    coordination = raw_trace.get("coordination", {}) if isinstance(raw_trace.get("coordination"), dict) else {}
                    memory = raw_trace.get("memory", {}) if isinstance(raw_trace.get("memory"), dict) else {}
                    trace = raw_trace.get("trace", {}) if isinstance(raw_trace.get("trace"), dict) else {}
                    agent = raw_trace.get("agent", {}) if isinstance(raw_trace.get("agent"), dict) else {}
                elif isinstance(raw_trace, dict) and "trace" in raw_trace:
                    trace_schema = "coordination_v1"
                    coordination = raw_trace.get("coordination", {})
                    memory = raw_trace.get("memory", {})
                    trace = raw_trace.get("trace", {})
                    agent = raw_trace.get("agent", {})
                else:
                    trace_schema = "legacy"
                    coordination = {}
                    memory = {}
                    trace = raw_trace if isinstance(raw_trace, dict) else {}
                    agent = {}
                retrieved_materials = raw_trace.get("retrieved_materials", []) if isinstance(raw_trace.get("retrieved_materials", []), list) else []
                retrieval_detail = self._compute_retrieval_feedback_detail(session, run_id=run_id, section_id=r.section_id)
                retrieval_detail["source_breakdown"] = self._summarize_source_breakdown(retrieved_materials)
                out.append(
                    {
                        "run_id": r.run_id,
                        "section_id": r.section_id,
                        "trace_schema": trace_schema,
                        "coordination": coordination,
                        "memory": memory,
                        "agent": agent,
                        "planner_result": raw_trace.get("planner_result", {}) if isinstance(raw_trace.get("planner_result", {}), dict) else {},
                        "retrieved_materials": retrieved_materials,
                        "retrieval_detail": retrieval_detail,
                        "prompt_rules": raw_trace.get("prompt_rules", {}) if isinstance(raw_trace.get("prompt_rules", {}), dict) else {},
                        "standardized_output": raw_trace.get("standardized_output", {}) if isinstance(raw_trace.get("standardized_output", {}), dict) else {},
                        "section_packet": raw_trace.get("section_packet", {}) if isinstance(raw_trace.get("section_packet", {}), dict) else {},
                        "trace": trace,
                        "create_time": r.create_time.strftime("%Y-%m-%d %H:%M:%S"),
                    }
                )
            return out
        finally:
            session.close()

    def get_section_patch_candidates(self, run_id: str, section_id: str = "") -> Tuple[bool, str, Optional[List[Dict[str, Any]]]]:
        print("[DEBUG] enter PreReviewService.get_section_patch_candidates | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
            if run is None:
                return False, "run not found", None

            related_run_ids = [
                str(item.run_id or "").strip()
                for item in session.query(PreReviewRun.run_id).filter(PreReviewRun.project_id == run.project_id).all()
                if str(item.run_id or "").strip()
            ]
            if not related_run_ids:
                related_run_ids = [str(run_id)]

            query = session.query(PreReviewPatchRegistry).filter(
                PreReviewPatchRegistry.run_id.in_(related_run_ids)
            )
            if str(section_id or "").strip():
                query = query.filter(PreReviewPatchRegistry.section_id == str(section_id).strip())

            rows = query.order_by(PreReviewPatchRegistry.update_time.desc(), PreReviewPatchRegistry.id.desc()).all()
            out: List[Dict[str, Any]] = []
            for row in rows:
                try:
                    payload = json.loads(row.payload_json) if row.payload_json else {}
                except Exception:
                    payload = {}
                out.append(
                    {
                        "patch_id": str(row.patch_id or ""),
                        "run_id": str(row.run_id or ""),
                        "section_id": str(row.section_id or ""),
                        "patch_type": str(row.patch_type or ""),
                        "target_agent": str(row.target_agent or ""),
                        "target_scope": str(row.target_scope or ""),
                        "trigger_condition": str(row.trigger_condition or ""),
                        "patch_content": str(row.patch_content or ""),
                        "source_feedback_key": str(row.source_feedback_key or ""),
                        "status": str(row.status or ""),
                        "version": int(row.version or 1),
                        "payload": payload if isinstance(payload, dict) else {},
                        "create_time": row.create_time.strftime("%Y-%m-%d %H:%M:%S") if row.create_time else "",
                        "update_time": row.update_time.strftime("%Y-%m-%d %H:%M:%S") if row.update_time else "",
                    }
                )
            return True, "success", out
        finally:
            session.close()

    def get_dashboard_summary(self) -> Dict[str, Any]:
        print("[DEBUG] enter PreReviewService.get_dashboard_summary | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            project_total = session.query(PreReviewProject).filter(PreReviewProject.is_deleted == 0).count()
            running_projects = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.is_deleted == 0, PreReviewProject.status == "running"))
                .count()
            )
            completed_projects = (
                session.query(PreReviewProject)
                .filter(and_(PreReviewProject.is_deleted == 0, PreReviewProject.status == "completed"))
                .count()
            )
            run_total = session.query(PreReviewRun).count()
            feedback_total = session.query(PreReviewFeedback).count()
            avg_acc = session.query(func.avg(PreReviewRun.accuracy)).scalar()

            recent_runs = (
                session.query(PreReviewRun)
                .order_by(desc(PreReviewRun.id))
                .limit(5)
                .all()
            )
            return {
                "project_total": project_total,
                "running_projects": running_projects,
                "completed_projects": completed_projects,
                "run_total": run_total,
                "feedback_total": feedback_total,
                "avg_accuracy": round(float(avg_acc), 4) if avg_acc is not None else None,
                "recent_runs": [
                    {
                        "run_id": r.run_id,
                        "project_id": r.project_id,
                        "source_doc_id": r.source_doc_id,
                        "accuracy": r.accuracy,
                        "create_time": r.create_time.strftime("%Y-%m-%d %H:%M:%S"),
                    }
                    for r in recent_runs
                ],
            }
        finally:
            session.close()


    def export_report_word(self, run_id: str) -> Tuple[bool, str, Optional[str]]:
        print("[DEBUG] enter PreReviewService.export_report_word | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
            if run is None:
                return False, "run not found", None

            sections = (
                session.query(PreReviewSectionConclusion)
                .filter(PreReviewSectionConclusion.run_id == run_id)
                .order_by(PreReviewSectionConclusion.id.asc())
                .all()
            )
            if not sections:
                return False, "no conclusions found", None

            doc = Document()
            doc.add_heading(f"Pre-review Report - {run.run_id}", 0)
            doc.add_paragraph(f"Project ID: {run.project_id}")
            doc.add_paragraph(f"Version: v{run.version_no}")
            doc.add_paragraph(f"Strategy: {run.strategy}")
            doc.add_paragraph(f"Generated At: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            for sec in sections:
                doc.add_heading(f"{sec.section_name} ({sec.section_id})", level=2)
                doc.add_paragraph(f"Risk Level: {sec.risk_level}")
                doc.add_paragraph(f"Conclusion: {sec.conclusion}")
                issues = PreReviewRepository.load_findings(sec.highlighted_issues)
                rules = PreReviewRepository.load_string_list(sec.linked_rules)
                issue_titles = self._findings_to_titles(issues)
                doc.add_paragraph(f"Highlighted Issues: {', '.join(issue_titles) if issue_titles else 'None'}")
                doc.add_paragraph(f"Linked Rules: {', '.join(rules) if rules else 'None'}")

            output_path = os.path.join(REPORT_DIR, f"{run_id}.docx")
            doc.save(output_path)
            return True, "report exported", output_path
        except Exception as e:
            return False, f"export report failed: {str(e)}", None
        finally:
            session.close()

    def _compute_feedback_stats(self, session, run_id: str, section_id: str = "") -> Dict[str, Any]:
        print("[DEBUG] enter PreReviewService._compute_feedback_stats | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        query = session.query(PreReviewFeedback).filter(PreReviewFeedback.run_id == run_id)
        if section_id:
            query = query.filter(PreReviewFeedback.section_id == section_id)
        rows = query.all()

        metrics = feedback_metrics([x.feedback_type for x in rows])
        return {
            "run_id": run_id,
            "section_id": section_id or "",
            "feedback_total": int(metrics["feedback_total"]),
            "valid_count": int(metrics["tp_valid"]),
            "false_positive_count": int(metrics["fp_false_positive"]),
            "missed_count": int(metrics["fn_missed"]),
            "accuracy": metrics["accuracy"],
            "precision": metrics["precision"],
            "recall": metrics["recall"],
            "f1": metrics["f1"],
            "reward_score": metrics["reward_score"],
        }

    @staticmethod
    def _feedback_event_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
        event_payload = dict(payload if isinstance(payload, dict) else {})
        event_payload["analysis_kind"] = str(event_payload.get("analysis_kind", "") or "feedback_event")
        return event_payload

    @staticmethod
    def _extract_feedback_meta(payload: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(payload, dict):
            return {}
        meta = payload.get("feedback_meta", {}) if isinstance(payload.get("feedback_meta", {}), dict) else {}
        chain_mode = str(meta.get("chain_mode", "") or payload.get("chain_mode", "") or "").strip()
        feedback_type = str(meta.get("feedback_type", "") or payload.get("feedback_type", "") or "").strip()
        decision = str(meta.get("decision", "") or payload.get("decision", "") or "").strip()
        return {
            "analysis_kind": str(payload.get("analysis_kind", "") or "feedback_optimize"),
            "chain_mode": chain_mode or "feedback_optimize",
            "feedback_type": feedback_type,
            "decision": decision,
            "manual_modified": bool(meta.get("manual_modified", False)),
            "diff_changed": bool(meta.get("diff_changed", False)),
            "feedback_optimize_status": str(meta.get("feedback_optimize_status", "") or payload.get("feedback_optimize_status", "") or "").strip(),
            "candidate_register_status": str(meta.get("candidate_register_status", "") or payload.get("candidate_register_status", "") or "").strip(),
            "replay_status": str(meta.get("replay_status", "") or payload.get("replay_status", "") or "").strip(),
            "error_message": str(meta.get("error_message", "") or payload.get("error_message", "") or "").strip(),
        }

    def record_feedback_pipeline_event(
        self,
        run_id: str,
        section_id: str,
        payload: Dict[str, Any],
    ) -> Tuple[bool, str]:
        session = self.db_conn.get_session()
        try:
            event_payload = self._feedback_event_payload(payload)
            feedback_key = str(
                event_payload.get("feedback_key", "") or f"feedback_pipeline:{run_id}:{section_id or 'global'}:{uuid.uuid4().hex[:8]}"
            )
            session.add(
                PreReviewFeedbackAnalysisResult(
                    feedback_key=feedback_key,
                    run_id=run_id,
                    section_id=section_id or None,
                    analysis_json=json.dumps(event_payload, ensure_ascii=False),
                    create_time=self._now(),
                )
            )
            self._refresh_accuracy(session, run_id)
            session.commit()
            return True, "feedback pipeline event recorded"
        except Exception as e:
            session.rollback()
            return False, f"record feedback pipeline event failed: {str(e)}"
        finally:
            session.close()

    def _load_feedback_analysis_payloads(self, session, run_id: str, section_id: str = "") -> List[Dict[str, Any]]:
        query = session.query(PreReviewFeedbackAnalysisResult).filter(PreReviewFeedbackAnalysisResult.run_id == run_id)
        if section_id:
            query = query.filter(PreReviewFeedbackAnalysisResult.section_id == section_id)
        rows = query.order_by(PreReviewFeedbackAnalysisResult.id.asc()).all()
        payloads: List[Dict[str, Any]] = []
        for row in rows:
            try:
                data = json.loads(row.analysis_json or "{}")
            except Exception:
                data = {}
            if isinstance(data, dict):
                payloads.append(data)
        return payloads

    def _compute_run_metrics(self, session, run_id: str, section_id: str = "") -> Dict[str, Any]:
        base_stats = self._compute_feedback_stats(session, run_id=run_id, section_id=section_id)
        analysis_payloads = self._load_feedback_analysis_payloads(session, run_id=run_id, section_id=section_id)

        feedback_only_count = 0
        feedback_optimize_count = 0
        feedback_optimize_success_count = 0
        feedback_optimize_failed_count = 0
        candidate_register_success_count = 0
        candidate_register_failed_count = 0
        replay_pass_count = 0
        replay_failed_count = 0
        manual_total = 0
        manual_changed = 0
        retrieval_tp = 0
        retrieval_fp = 0
        retrieval_fn = 0

        for payload in analysis_payloads:
            meta = self._extract_feedback_meta(payload)
            chain_mode = meta.get("chain_mode", "")
            if chain_mode == "feedback_only":
                feedback_only_count += 1
            elif chain_mode == "feedback_optimize":
                feedback_optimize_count += 1
            if meta.get("feedback_optimize_status") == "completed":
                feedback_optimize_success_count += 1
            elif meta.get("feedback_optimize_status") == "failed":
                feedback_optimize_failed_count += 1
            if meta.get("candidate_register_status") == "completed":
                candidate_register_success_count += 1
            elif meta.get("candidate_register_status") == "failed":
                candidate_register_failed_count += 1
            if meta.get("replay_status") == "passed":
                replay_pass_count += 1
            elif meta.get("replay_status") == "failed":
                replay_failed_count += 1

            if meta.get("manual_modified") or meta.get("diff_changed"):
                manual_changed += 1
            if meta.get("analysis_kind") in {"feedback_event", "feedback_optimize"}:
                manual_total += 1

            if meta.get("analysis_kind") != "feedback_optimize":
                continue
            error_types = {
                str(x).strip()
                for x in payload.get("error_types", [])
                if str(x).strip()
            }
            primary_error_type = str(payload.get("primary_error_type", "") or "").strip()
            if primary_error_type:
                error_types.add(primary_error_type)
            if error_types & RETRIEVAL_FALSE_NEGATIVE_ERRORS:
                retrieval_fn += 1
            elif error_types & RETRIEVAL_FALSE_POSITIVE_ERRORS:
                retrieval_fp += 1
            else:
                retrieval_tp += 1

        retrieval_total = retrieval_tp + retrieval_fp + retrieval_fn
        retrieval_precision = _safe_ratio(retrieval_tp, retrieval_tp + retrieval_fp)
        retrieval_recall = _safe_ratio(retrieval_tp, retrieval_tp + retrieval_fn)
        retrieval_f1 = _safe_f1(retrieval_precision, retrieval_recall)
        retrieval_accuracy = _safe_ratio(retrieval_tp, retrieval_total)

        feedback_acceptance_rate = _safe_ratio(base_stats["valid_count"], base_stats["feedback_total"])
        manual_modification_rate = _safe_ratio(manual_changed, manual_total)

        run_row = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
        trajectory: Dict[str, Any] = {
            "previous_run_id": None,
            "previous_false_positive_count": None,
            "false_positive_delta": None,
            "false_positive_reduction_rate": None,
        }
        if run_row is not None:
            previous_run = (
                session.query(PreReviewRun)
                .filter(
                    and_(
                        PreReviewRun.project_id == run_row.project_id,
                        PreReviewRun.source_doc_id == run_row.source_doc_id,
                        PreReviewRun.id < run_row.id,
                    )
                )
                .order_by(desc(PreReviewRun.id))
                .first()
            )
            if previous_run is not None:
                prev_summary = {}
                try:
                    prev_summary = json.loads(previous_run.summary) if previous_run.summary else {}
                except Exception:
                    prev_summary = {}
                prev_metrics = prev_summary.get("metrics", {}) if isinstance(prev_summary.get("metrics", {}), dict) else {}
                prev_review = prev_metrics.get("review", {}) if isinstance(prev_metrics.get("review", {}), dict) else {}
                previous_fp = prev_review.get("false_positive_count")
                if previous_fp is None:
                    previous_fp = self._compute_feedback_stats(session, previous_run.run_id).get("false_positive_count", 0)
                if previous_fp is not None:
                    previous_fp = int(previous_fp)
                    fp_delta = previous_fp - int(base_stats["false_positive_count"])
                    trajectory = {
                        "previous_run_id": previous_run.run_id,
                        "previous_false_positive_count": previous_fp,
                        "false_positive_delta": fp_delta,
                        "false_positive_reduction_rate": _safe_ratio(max(fp_delta, 0), previous_fp) if previous_fp > 0 else None,
                    }

        return {
            "run_id": run_id,
            "section_id": section_id or "",
            "review": {
                "feedback_total": int(base_stats["feedback_total"]),
                "valid_count": int(base_stats["valid_count"]),
                "false_positive_count": int(base_stats["false_positive_count"]),
                "missed_count": int(base_stats["missed_count"]),
                "accuracy": float(base_stats["accuracy"]),
                "precision": float(base_stats["precision"]),
                "recall": float(base_stats["recall"]),
                "f1": float(base_stats["f1"]),
                "reward_score": float(base_stats["reward_score"]),
            },
            "retrieval": {
                "evaluated_feedback_count": int(retrieval_total),
                "true_positive_count": int(retrieval_tp),
                "false_positive_count": int(retrieval_fp),
                "false_negative_count": int(retrieval_fn),
                "accuracy": float(retrieval_accuracy),
                "precision": float(retrieval_precision),
                "recall": float(retrieval_recall),
                "f1": float(retrieval_f1),
            },
            "feedback": {
                "feedback_only_count": int(feedback_only_count),
                "feedback_optimize_count": int(feedback_optimize_count),
                "feedback_optimize_success_count": int(feedback_optimize_success_count),
                "feedback_optimize_failed_count": int(feedback_optimize_failed_count),
                "candidate_register_success_count": int(candidate_register_success_count),
                "candidate_register_failed_count": int(candidate_register_failed_count),
                "replay_pass_count": int(replay_pass_count),
                "replay_failed_count": int(replay_failed_count),
                "feedback_acceptance_rate": float(feedback_acceptance_rate),
                "manual_modification_count": int(manual_changed),
                "manual_modification_total": int(manual_total),
                "manual_modification_rate": float(manual_modification_rate),
            },
            "trajectory": trajectory,
        }

    def _compute_retrieval_feedback_detail(self, session, run_id: str, section_id: str = "") -> Dict[str, Any]:
        analysis_payloads = self._load_feedback_analysis_payloads(session, run_id=run_id, section_id=section_id)
        retrieval_tp = 0
        retrieval_fp = 0
        retrieval_fn = 0
        breakdown: Dict[str, int] = {}
        for payload in analysis_payloads:
            meta = self._extract_feedback_meta(payload)
            if meta.get("analysis_kind") != "feedback_optimize":
                continue
            error_types = {
                str(x).strip()
                for x in payload.get("error_types", [])
                if str(x).strip()
            }
            primary_error_type = str(payload.get("primary_error_type", "") or "").strip()
            if primary_error_type:
                error_types.add(primary_error_type)
            if error_types & RETRIEVAL_FALSE_NEGATIVE_ERRORS:
                retrieval_fn += 1
            elif error_types & RETRIEVAL_FALSE_POSITIVE_ERRORS:
                retrieval_fp += 1
            else:
                retrieval_tp += 1
            for item in error_types:
                breakdown[item] = int(breakdown.get(item, 0)) + 1
        precision = _safe_ratio(retrieval_tp, retrieval_tp + retrieval_fp)
        recall = _safe_ratio(retrieval_tp, retrieval_tp + retrieval_fn)
        return {
            "metrics": {
                "evaluated_feedback_count": int(retrieval_tp + retrieval_fp + retrieval_fn),
                "true_positive_count": int(retrieval_tp),
                "false_positive_count": int(retrieval_fp),
                "false_negative_count": int(retrieval_fn),
                "accuracy": float(_safe_ratio(retrieval_tp, retrieval_tp + retrieval_fp + retrieval_fn)),
                "precision": float(precision),
                "recall": float(recall),
                "f1": float(_safe_f1(precision, recall)),
            },
            "error_breakdown": breakdown,
        }

    @staticmethod
    def _summarize_source_breakdown(retrieved_materials: List[Dict[str, Any]]) -> Dict[str, int]:
        breakdown: Dict[str, int] = {}
        for item in retrieved_materials:
            if not isinstance(item, dict):
                continue
            source_type = str(item.get("source_type", "") or "").strip() or "鏈煡鏉ユ簮"
            breakdown[source_type] = int(breakdown.get(source_type, 0)) + 1
        return breakdown

    def _refresh_run_metrics_summary(self, session, run_id: str) -> Dict[str, Any]:
        metrics = self._compute_run_metrics(session, run_id=run_id, section_id="")
        run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
        if run is None:
            return metrics
        try:
            summary_payload = json.loads(run.summary) if run.summary else {}
        except Exception:
            summary_payload = {}
        if not isinstance(summary_payload, dict):
            summary_payload = {}
        summary_payload["metrics"] = metrics
        summary_payload["metrics_schema"] = "run_metrics_v1"
        summary_payload["last_metrics_refresh_time"] = self._now().strftime("%Y-%m-%d %H:%M:%S")
        run.summary = json.dumps(summary_payload, ensure_ascii=False)
        return metrics

    def _refresh_accuracy(self, session, run_id: str) -> Dict[str, Any]:
        print("[DEBUG] enter PreReviewService._refresh_accuracy | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        stats = self._compute_feedback_stats(session, run_id)
        run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
        if run is not None:
            run.accuracy = stats["accuracy"] if stats["feedback_total"] > 0 else None
        self._refresh_run_metrics_summary(session, run_id)
        return stats

    @staticmethod
    def _decision_to_feedback_type(decision: str, labels: Optional[List[Any]] = None) -> str:
        normalized = str(decision or "").strip().lower()
        label_set = {str(x).strip().lower() for x in (labels or []) if str(x).strip()}
        if normalized in {"false_positive", "invalid", "rejected"}:
            return "false_positive"
        if normalized in {"missed", "missing_risk"} or {"missed", "missing_risk"} & label_set:
            return "missed"
        return "valid"

    def _optimize_from_feedback_memory(
        self,
        run_id: str,
        section_id: str,
        feedback_type: str,
        feedback_text: str,
        suggestion: str,
        operator: str,
    ) -> None:
        print("[DEBUG] enter PreReviewService._optimize_from_feedback_memory | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        note = f"{feedback_type} | {feedback_text or ''} | {suggestion or ''}".strip(" |")
        if not note:
            return

        base_meta = {
            "source": "online_feedback",
            "run_id": run_id,
            "section_id": section_id or "",
            "operator": operator or "",
        }
        memory_type = "episodic"
        if feedback_type == "missed":
            memory_type = "semantic"
        elif feedback_type == "false_positive":
            memory_type = "working"

        self.memory_tool.remember(
            key=f"feedback:{run_id}:{section_id or 'global'}:{uuid.uuid4().hex[:8]}",
            value=note,
            memory_type=memory_type,
            metadata=base_meta,
        )

    def add_feedback(
        self,
        run_id: str,
        section_id: str,
        feedback_type: str,
        feedback_text: str = "",
        suggestion: str = "",
        operator: str = "",
        feedback_meta: Optional[Dict[str, Any]] = None,
    ) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.add_feedback | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        if feedback_type not in {"valid", "false_positive", "missed"}:
            return False, "feedback_type must be one of valid/false_positive/missed", None

        session = self.db_conn.get_session()
        try:
            run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
            if run is None:
                return False, "run not found", None

            if section_id:
                section_exists = (
                    session.query(PreReviewSectionConclusion)
                    .filter(
                        and_(
                            PreReviewSectionConclusion.run_id == run_id,
                            PreReviewSectionConclusion.section_id == section_id,
                        )
                    )
                    .first()
                )
                if section_exists is None:
                    return False, "section_id does not belong to this run", None

            fb = PreReviewFeedback(
                run_id=run_id,
                section_id=section_id or None,
                feedback_type=feedback_type,
                feedback_text=feedback_text,
                suggestion=suggestion,
                operator=operator,
                create_time=self._now(),
            )
            session.add(fb)
            if isinstance(feedback_meta, dict) and feedback_meta.get("persist_event", False):
                feedback_key = str(feedback_meta.get("feedback_key", "") or f"feedback_event:{run_id}:{section_id or 'global'}:{uuid.uuid4().hex[:8]}")
                event_payload = self._feedback_event_payload(
                    {
                        "chain_mode": str(feedback_meta.get("chain_mode", "") or "feedback_only"),
                        "feedback_type": feedback_type,
                        "feedback_meta": {
                            "chain_mode": str(feedback_meta.get("chain_mode", "") or "feedback_only"),
                            "feedback_type": feedback_type,
                            "decision": str(feedback_meta.get("decision", "") or ""),
                            "manual_modified": bool(feedback_meta.get("manual_modified", False)),
                            "diff_changed": bool(feedback_meta.get("manual_modified", False)),
                        },
                    }
                )
                session.add(
                    PreReviewFeedbackAnalysisResult(
                        feedback_key=feedback_key,
                        run_id=run_id,
                        section_id=section_id or None,
                        analysis_json=json.dumps(event_payload, ensure_ascii=False),
                        create_time=self._now(),
                    )
                )
            session.flush()
            stats = self._refresh_accuracy(session, run_id)
            session.commit()
            self._optimize_from_feedback_memory(
                run_id=run_id,
                section_id=section_id,
                feedback_type=feedback_type,
                feedback_text=feedback_text,
                suggestion=suggestion,
                operator=operator,
            )
            return True, "feedback added", stats
        except Exception as e:
            session.rollback()
            return False, f"add feedback failed: {str(e)}", None
        finally:
            session.close()

    def _load_current_template_text(self, template_name: str) -> str:
        template_path = Path(__file__).resolve().parents[1] / "prompts" / "pre_review_agent_prompt" / template_name
        if not template_path.exists():
            return ""
        try:
            return template_path.read_text(encoding="utf-8")
        except Exception:
            return ""

    def _load_feedback_section_context(self, session, run_id: str, section_id: str) -> Dict[str, Any]:
        run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
        if run is None:
            return {}
        ok, _, section_payload = self.get_submission_sections(project_id=run.project_id, doc_id=run.source_doc_id)
        sections = section_payload.get("sections", []) if ok and isinstance(section_payload, dict) else []
        section_item = {}
        for item in sections:
            if not isinstance(item, dict):
                continue
            if str(item.get("section_id", "") or "") == section_id:
                section_item = item
                break
        output_row = (
            session.query(PreReviewSectionOutput)
            .filter(
                and_(
                    PreReviewSectionOutput.run_id == run_id,
                    PreReviewSectionOutput.section_id == section_id,
                )
            )
            .order_by(PreReviewSectionOutput.id.desc())
            .first()
        )
        standardized_output = {}
        if output_row is not None:
            try:
                standardized_output = json.loads(output_row.output_json or "{}")
            except Exception:
                standardized_output = {}
        trace_row = (
            session.query(PreReviewSectionTrace)
            .filter(
                and_(
                    PreReviewSectionTrace.run_id == run_id,
                    PreReviewSectionTrace.section_id == section_id,
                )
            )
            .order_by(PreReviewSectionTrace.id.desc())
            .first()
        )
        trace_payload = {}
        if trace_row is not None:
            trace_payload = PreReviewRepository.load_trace_payload(trace_row.trace_json)
        product_type = self._infer_product_type(
            session.query(PreReviewProject).filter(PreReviewProject.project_id == run.project_id).first(),
            section_name=str(section_item.get("section_name", "") or ""),
        )
        historical_experience = self._load_historical_experience(
            session=session,
            section_id=section_id,
            product_type=product_type,
        )
        return {
            "run": run,
            "section_item": section_item,
            "standardized_output": standardized_output if isinstance(standardized_output, dict) else {},
            "trace_payload": trace_payload if isinstance(trace_payload, dict) else {},
            "historical_experience": historical_experience,
        }

    def process_feedback_closed_loop(
        self,
        feedback_record: Dict[str, Any],
        run_trace: Dict[str, Any],
        run_context: Dict[str, Any],
    ) -> Dict[str, Any]:
        session = self.db_conn.get_session()
        try:
            run_id = str(feedback_record.get("run_id", "") or "")
            section_id = str(feedback_record.get("section_id", "") or "")
            context = self._load_feedback_section_context(session, run_id=run_id, section_id=section_id)
            if not context:
                return {}
            run = context.get("run")
            if run is None:
                return {}
            section_item = context.get("section_item", {}) if isinstance(context.get("section_item", {}), dict) else {}
            standardized_output = context.get("standardized_output", {}) if isinstance(context.get("standardized_output", {}), dict) else {}
            trace_payload = context.get("trace_payload", {}) if isinstance(context.get("trace_payload", {}), dict) else {}
            historical_experience = context.get("historical_experience", []) if isinstance(context.get("historical_experience", []), list) else []
            feedback_key = f"{run_id}:{section_id}:{uuid.uuid4().hex[:8]}"
            analyzer_input = {
                "task_id": feedback_key,
                "section_id": section_id,
                "raw_text": str(section_item.get("content", "") or trace_payload.get("section_packet", {}).get("raw_text", "") or ""),
                "focus_points": list(section_item.get("concern_points", []) or trace_payload.get("section_packet", {}).get("focus_points", []) or []),
                "system_output": standardized_output,
                "user_feedback_text": str(feedback_record.get("feedback_text", "") or ""),
                "decision": str(feedback_record.get("decision", "") or ""),
                "labels": list(feedback_record.get("labels", []) or []),
                "reference_inputs": {
                    "retrieved_materials": trace_payload.get("retrieved_materials", []) if isinstance(trace_payload.get("retrieved_materials", []), list) else [],
                    "historical_experience": historical_experience,
                },
            }
            analyzer_prompt_config = self._compose_runtime_prompt_config(
                project_id=run.project_id,
                task_type="feedback_analyzer",
                base_prompt_config=(run_context.get("prompt_config", {}) if isinstance(run_context.get("prompt_config", {}), dict) else {}),
                section_id=section_id,
                section_name=str(section_item.get("section_name", "") or ""),
                review_domain=self._infer_review_domain(
                    session.query(PreReviewProject).filter(PreReviewProject.project_id == run.project_id).first()
                ),
                product_type=product_type,
                registration_class=str((run_context.get("run_config", {}) if isinstance(run_context.get("run_config", {}), dict) else {}).get("registration_class", "") or ""),
            )
            analysis_result = self.feedback_agent.analyze_feedback(analyzer_input, prompt_config=analyzer_prompt_config)
            optimizer_input = {
                "task_id": feedback_key,
                "section_id": section_id,
                "raw_text": analyzer_input["raw_text"],
                "focus_points": analyzer_input["focus_points"],
                "user_feedback_text": analyzer_input["user_feedback_text"],
                "feedback_analysis_result": analysis_result,
                "current_templates": {
                    "planner_prompt": self._load_current_template_text("chapter_planner.j2"),
                    "pre_review_prompt": self._load_current_template_text("chapter_reviewer.j2"),
                    "feedback_analyzer_prompt": self._load_current_template_text("feedback_analyzer.j2"),
                    "feedback_optimizer_prompt": self._load_current_template_text("feedback_optimizer.j2"),
                },
            }
            optimizer_prompt_config = self._compose_runtime_prompt_config(
                project_id=run.project_id,
                task_type="feedback_optimizer",
                base_prompt_config=(run_context.get("prompt_config", {}) if isinstance(run_context.get("prompt_config", {}), dict) else {}),
                section_id=section_id,
                section_name=str(section_item.get("section_name", "") or ""),
                review_domain=self._infer_review_domain(
                    session.query(PreReviewProject).filter(PreReviewProject.project_id == run.project_id).first()
                ),
                product_type=product_type,
                registration_class=str((run_context.get("run_config", {}) if isinstance(run_context.get("run_config", {}), dict) else {}).get("registration_class", "") or ""),
            )
            patch_result = self.feedback_agent.propose_patch(optimizer_input, prompt_config=optimizer_prompt_config)
            pipeline_trace = {
                "feedback_optimize_status": "completed",
                "candidate_register_status": "pending",
                "replay_status": "pending",
                "error_message": "",
            }
            persisted_analysis = dict(analysis_result if isinstance(analysis_result, dict) else {})
            persisted_analysis["analysis_kind"] = "feedback_optimize"
            persisted_analysis["feedback_type"] = self._decision_to_feedback_type(
                decision=str(feedback_record.get("decision", "") or ""),
                labels=feedback_record.get("labels", []) if isinstance(feedback_record.get("labels", []), list) else [],
            )
            persisted_analysis["feedback_meta"] = {
                "chain_mode": str(feedback_record.get("chain_mode", "") or "feedback_optimize"),
                "decision": str(feedback_record.get("decision", "") or ""),
                "manual_modified": bool(feedback_record.get("manual_modified", False)),
                "diff_changed": bool(
                    (
                        (feedback_record.get("diff_result", {}) if isinstance(feedback_record.get("diff_result", {}), dict) else {})
                        .get("changed", False)
                    )
                ),
                "feedback_type": persisted_analysis["feedback_type"],
                "feedback_optimize_status": pipeline_trace["feedback_optimize_status"],
                "candidate_register_status": pipeline_trace["candidate_register_status"],
                "replay_status": pipeline_trace["replay_status"],
                "error_message": pipeline_trace["error_message"],
            }
            persisted_analysis["feedback_optimize_status"] = pipeline_trace["feedback_optimize_status"]
            persisted_analysis["candidate_register_status"] = pipeline_trace["candidate_register_status"]
            persisted_analysis["replay_status"] = pipeline_trace["replay_status"]
            persisted_analysis["error_message"] = pipeline_trace["error_message"]
            session.add(
                PreReviewFeedbackAnalysisResult(
                    feedback_key=feedback_key,
                    run_id=run_id,
                    section_id=section_id or None,
                    analysis_json=json.dumps(persisted_analysis, ensure_ascii=False),
                    create_time=self._now(),
                )
            )
            for exp in analysis_result.get("new_experience", []) if isinstance(analysis_result.get("new_experience", []), list) else []:
                if not isinstance(exp, dict):
                    continue
                applicable_scope = str(exp.get("applicable_scope", "") or section_id).strip() or section_id
                scope_type = "section" if applicable_scope.startswith("3.2.") else "product_type"
                session.add(
                    PreReviewExperienceMemory(
                        experience_id=f"exp_{uuid.uuid4().hex[:12]}",
                        scope_type=scope_type,
                        scope_key=applicable_scope,
                        experience_type=str(exp.get("experience_type", "") or "review_rule"),
                        content=str(exp.get("content", "") or "").strip(),
                        source_feedback_ids=json.dumps([feedback_key], ensure_ascii=False),
                        trigger_conditions=json.dumps([f"section_id == '{section_id}'"], ensure_ascii=False),
                        usage_count=0,
                        success_count=0,
                        status="active",
                        payload_json=json.dumps(exp, ensure_ascii=False),
                        create_time=self._now(),
                        update_time=self._now(),
                    )
                )
            patch_version = 1
            for patch in patch_result.get("patches", []) if isinstance(patch_result.get("patches", []), list) else []:
                if not isinstance(patch, dict):
                    continue
                session.add(
                    PreReviewPatchRegistry(
                        patch_id=str(patch.get("patch_id", "") or f"patch_{uuid.uuid4().hex[:12]}"),
                        run_id=run_id,
                        section_id=section_id or None,
                        patch_type=str(patch.get("patch_type", "") or ""),
                        target_agent=str(patch.get("target_agent", "") or ""),
                        target_scope=str(patch.get("target_scope", "") or section_id),
                        trigger_condition=str(patch.get("trigger_condition", "") or ""),
                        patch_content=str(patch.get("patch_content", "") or "").strip(),
                        source_feedback_key=feedback_key,
                        status=str(patch.get("status", "candidate") or "candidate"),
                        version=patch_version,
                        payload_json=json.dumps(patch, ensure_ascii=False),
                        create_time=self._now(),
                        update_time=self._now(),
                    )
                )
                patch_version += 1
            self._refresh_accuracy(session, run_id)
            session.commit()
            return {
                "success": True,
                "feedback_key": feedback_key,
                "analysis_result": analysis_result,
                "patch_result": patch_result,
                "active_rules": {
                    "feedback_analyzer": ((analyzer_prompt_config.get("prompt_bundle", {}) if isinstance(analyzer_prompt_config.get("prompt_bundle", {}), dict) else {}).get("active_rules", {}) or {}).get("feedback_analyzer", []),
                    "feedback_optimizer": ((optimizer_prompt_config.get("prompt_bundle", {}) if isinstance(optimizer_prompt_config.get("prompt_bundle", {}), dict) else {}).get("active_rules", {}) or {}).get("feedback_optimizer", []),
                },
                "feedback_optimize_status": pipeline_trace["feedback_optimize_status"],
                "candidate_register_status": pipeline_trace["candidate_register_status"],
                "replay_status": pipeline_trace["replay_status"],
                "error_message": pipeline_trace["error_message"],
            }
        except Exception as exc:
            session.rollback()
            return {
                "success": False,
                "feedback_key": "",
                "analysis_result": None,
                "patch_result": None,
                "feedback_optimize_status": "failed",
                "candidate_register_status": "skipped",
                "replay_status": "skipped",
                "error_message": str(exc),
            }
        finally:
            session.close()

    def get_feedback_stats(self, run_id: str, section_id: str = "") -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        print("[DEBUG] enter PreReviewService.get_feedback_stats | core:", {k: ((v[:100] + "...") if isinstance(v, str) and len(v) > 100 else v) for k, v in locals().items() if k in ("project_id", "doc_id", "file_path", "file_name", "file_type", "classification", "section_id", "run_id", "query", "page", "page_size", "status", "chunk_id")})
        session = self.db_conn.get_session()
        try:
            run = session.query(PreReviewRun).filter(PreReviewRun.run_id == run_id).first()
            if run is None:
                return False, "run not found", None

            stats = self._compute_run_metrics(session, run_id=run_id, section_id=section_id)
            return True, "success", stats
        except Exception as e:
            return False, f"get feedback stats failed: {str(e)}", None
        finally:
            session.close()


